import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
import xgboost
import shap
from docx import Document
from docx.shared import Inches
import random
import os
from pathlib import Path
from docx.oxml.ns import qn
import glob


class Constants:
    SEED_VALUE = 42
    LSTM_BATCH_SIZE = 1024
    LSTM_EPOCHS = 50
    TIME_SERIES_SPLITS = 5
    SEQ_LENGTH = 24
    PRED_LENGTH = 24
    MISSING_VALUE_THRESHOLD = 0.5
    FILE_NAME_MAPPING = {
        "(37t)ศาลหลักเมือง(ปิดสถานี)": "San_Lak_Mueang_Closed_Station",
        "(37t)สถานีอุตุนิยมวิทยาลำปาง": "Lampang_Meteorological_Station",
        "(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด": "Ban_Sopad_Hospital",
        "(39t)โรงพยาบาลส่งเสริมสุขภาพตำบลท่าสี": "Tha_Si_Hospital",
        "(40t)การประปาส่วนภูมิภาคแม่เมาะ": "Mae_Mo_Waterworks",
    }
    OUTPUT_DIR_NAME = "outputs"


class DataProcessor:

    def __init__(self, directory):
        self.directory = directory
        self.constants = Constants()  # Access constants through an instance
        self.set_random_seeds()
        self.configure_gpu()

    def set_random_seeds(self):
        """Sets random seeds for reproducibility."""
        random.seed(self.constants.SEED_VALUE)
        np.random.seed(self.constants.SEED_VALUE)
        tf.random.set_seed(self.constants.SEED_VALUE)

    def configure_gpu(self):
        """Configures GPU memory growth for LSTM."""
        gpus = tf.config.experimental.list_physical_devices('GPU')
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
                print("GPU is set up and ready for LSTM!")
            except RuntimeError as e:
                print(e)
        print("Num GPUs Available: ", len(tf.config.list_physical_devices('GPU')))


    def load_data(self, file_path):
        """Loads data from a CSV file."""
        print(f"Loading data from {file_path}...")
        df = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
        return df.copy()  # Return a copy to avoid modifying the original DataFrame


    def process_files(self):
        """Processes all CSV files in the specified directory."""
        all_files = glob.glob(os.path.join(self.directory, "*.csv"))

        if not all_files:
            raise ValueError(f"No CSV files found in the directory: {self.directory}")

        for file_path in all_files:
            try:
                self.process_single_file(file_path)
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")


    def process_single_file(self, file_path):
        """Processes a single CSV file."""

        df_raw = self.load_data(file_path)

        # Dynamically generate report name from file path
        file_name = os.path.basename(file_path).replace(".csv", "")
        english_file_name = self.constants.FILE_NAME_MAPPING.get(file_name, file_name)
        if all([char >= 'ก' and char <= '๙' for char in file_name]):
             raise ValueError("Invalid filename: Must be in Thai characters [ก-๙]")
        output_dir = Path.cwd() / self.constants.OUTPUT_DIR_NAME
        output_dir.mkdir(exist_ok=True)

        doc = Document()
        doc.add_heading(f"LSTM Model Report for SO2 Prediction with XGBoost SHAP FI - {file_name}", 0)
        doc.add_paragraph(f"Dataset: {file_path} (Limited to 1000 rows)")

        df = self.eda_and_preprocessing(df_raw, doc, file_name, english_file_name, output_dir)
        if df is None:
             return
        X_full, y_full, X_full_flat, y_full_flat, features, scaler_target = self.prepare_data_for_modeling(df, doc, file_name)

        if X_full is None:
            return # Stop if data preparation failed
        self.train_lstm_model(X_full, y_full, features, scaler_target, doc, file_name, english_file_name, output_dir)
        self.shap_analysis(X_full_flat, y_full_flat, features, doc, file_name, english_file_name, output_dir)
        self.final_lstm_prediction(df, X_full, features, scaler_target, doc, file_name, english_file_name, output_dir)

        self.save_report(doc, output_dir, file_name)


    def eda_and_preprocessing(self, df, doc, file_name, english_file_name, output_dir):
        """Performs EDA and preprocessing steps."""
        # Show Data Before Any Processing
        print("Step 5: Showing Data Before Any Processing...")
        doc.add_heading(f"Data Before Any Processing ({file_name})", 1)
        doc.add_paragraph(f"Shape Before: {df.shape}")
        doc.add_paragraph(f"Missing Values Total Before: {df.isnull().sum().sum()}")

        # Columns Description (List of Features)
        print("Step 5.1: Listing Columns (Before)...")
        doc.add_heading(f"Columns List (Before) - {file_name}", 2)
        doc.add_paragraph(f"List of columns in the dataset before processing ({file_name}):")
        for col in df.columns:
            doc.add_paragraph(f"- {col}")

        # Data Sample (Head(10) Before)
        print("Step 5.2: Showing Data Sample (Before)...")
        df_head_before = df.head(10)
        self.add_table(doc, df_head_before, f"Sample of First 10 Rows (Before) - {file_name}", f"ตารางนี้แสดงตัวอย่าง 10 แถวแรกของข้อมูลดิบก่อนดำเนินการ ({file_name})")

        # Statistical Summary using df.info (Before)
        print("Step 5.3: Computing Data Information (Before)...")
        doc.add_heading(f"Data Information (Before) - {file_name}", 2)
        # Capture df.info() output as a string
        info_str = []
        df.info(buf=pd.io.common.StringIO(), verbose=True, memory_usage=True, show_counts=True)
        info_str = df.info(verbose=True, buf=pd.io.common.StringIO(), memory_usage=True, show_counts=True)
        doc.add_paragraph(str(info_str))

        # EDA (Before Removing Missing > 50%, Trim, Trim SO2)
        print("Step 6: Performing EDA (Before Any Processing)...")
        doc.add_heading(f"Exploratory Data Analysis (Before Any Processing) - {file_name}", 1)
        doc.add_paragraph(f"Shape: {df.shape}")
        doc.add_paragraph(f"Missing Values Total: {df.isnull().sum().sum()}")

        # Missing Data Heatmap (Before Any Processing)
        plt.figure(figsize=(12, 6))
        heatmap = sns.heatmap(df.isnull(), cbar=True, cmap='Blues', annot=False)
        plt.title(f"Missing Data Heatmap (Before Any Processing) - {english_file_name}")
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(output_dir / f"eda_heatmap_before_all_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()

        nan_count_before_all = df.isnull().sum().sum()
        nan_max_col_before_all = df.isnull().sum().idxmax() if nan_count_before_all > 0 else "None"
        nan_max_val_before_all = df.isnull().sum().max() if nan_count_before_all > 0 else 0
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) ก่อนดำเนินการใด ๆ รวมทั้งหมด {nan_count_before_all} ค่า ({file_name}) "
        if nan_count_before_all > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_before_all}' ({nan_max_val_before_all} ค่า) "
            if nan_max_val_before_all > df.shape[0] * 0.5:
                desc += " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่ "
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น "
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วง 1000 แถวนี้ "
        desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}"
        self.add_figure(doc, output_dir / f"eda_heatmap_before_all_{english_file_name}.png", f"Missing Data Heatmap (Before Any Processing) - {file_name}", desc)

        # Correlation Matrix (Before Any Processing)
        corr_matrix_before = df.corr()
        correlation_df_before = pd.DataFrame({
            "feature": corr_matrix_before.index,
            "correlation": corr_matrix_before['SO2'],
            "abs_correlation": corr_matrix_before['SO2'].abs(),
        })
        correlation_df_before = correlation_df_before.sort_values(by="abs_correlation", ascending=False)
        self.add_table(doc, correlation_df_before, f"Sorted Correlation Matrix with SO2 (Before) - {file_name}", f"ตารางนี้แสดงค่าสหสัมพันธ์กับ SO2 เรียงลำดับตามค่าสัมประสิทธิ์สัมพันธ์สัมบูรณ์ก่อนดำเนินการ ({file_name})")

        plt.figure(figsize=(10, 8))
        heatmap = sns.heatmap(corr_matrix_before, annot=True, cmap='RdBu', center=0, vmin=-1, vmax=1, fmt='.2f')
        plt.title(f"Correlation Matrix (Before Any Processing) - {english_file_name}")
        plt.xticks(rotation=45, ha='right')
        plt.yticks(rotation=0)
        plt.savefig(output_dir / f"eda_corr_before_all_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()

        corr_max_before = df.corr()['SO2'].drop('SO2').abs().max()
        corr_max_col_before = df.corr()['SO2'].drop('SO2').abs().idxmax()
        corr_min_before = df.corr()['SO2'].drop('SO2').abs().min()
        corr_min_col_before = df.corr()['SO2'].drop('SO2').abs().idxmin()
        desc = f"กราฟนี้แสดงความสัมพันธ์ระหว่างฟีเจอร์ก่อนดำเนินการ ค่าสหสัมพันธ์สูงสุดกับ SO2 คือ {corr_max_before:.2f} จาก '{corr_max_col_before}' ({file_name}) "
        desc += f"ค่าสหสัมพันธ์ต่ำสุดกับ SO2 คือ {corr_min_before:.2f} จาก '{corr_min_col_before}' "
        if corr_max_before > 0.7:
            desc += "แสดงถึงความสัมพันธ์ที่สูงมาก "
        elif corr_max_before > 0.3:
            desc += "แสดงถึงความสัมพันธ์ปานกลาง "
        else:
            desc += "แสดงถึงความสัมพันธ์ที่ต่ำ "
        desc += "การกระจายของค่าสหสัมพันธ์มีแนวโน้มสมมาตรตามแนวทแยง"
        self.add_figure(doc, output_dir / f"eda_corr_before_all_{english_file_name}.png", f"Correlation Matrix (Before Any Processing) - {file_name}", desc)

        # Feature Outlier Distribution (Before Any Processing)
        plt.figure(figsize=(15, 6))
        boxplot = df.boxplot(figsize=(15, 6))
        plt.xticks(rotation=45, ha='right')
        plt.title(f"Feature Outlier Distribution (Before Any Processing) - {english_file_name}")
        plt.savefig(output_dir / f"eda_boxplot_before_all_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()

        outliers_before = {}
        for col in df.columns:
            q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
            iqr = q3 - q1
            outlier_count = ((df[col] < (q1 - 1.5 * iqr)) | (df[col] > (q3 + 1.5 * iqr))).sum()
            outliers_before[col] = outlier_count
        max_outlier_col_before = max(outliers_before, key=outliers_before.get)
        min_outlier_col_before = min(outliers_before, key=outliers_before.get)
        desc = f"กราฟนี้แสดงการกระจายและ outlier ของฟีเจอร์ก่อนดำเนินการ คอลัมน์ที่มี outlier มากที่สุดคือ '{max_outlier_col_before}' ({outliers_before[max_outlier_col_before]} ค่า) ({file_name}) "
        desc += f"คอลัมน์ที่มี outlier น้อยที่สุดคือ '{min_outlier_col_before}' ({outliers_before[min_outlier_col_before]} ค่า) "
        if outliers_before[max_outlier_col_before] > df.shape[0] * 0.1:
            desc += "ซึ่งเกิน 10% ของข้อมูล แสดงถึงการกระจายที่มีความผันผวนสูง "
        else:
            desc += "ซึ่งอยู่ในระดับต่ำ ไม่มีผลกระทบมาก "
        desc += f"ค่าสถิติของ '{max_outlier_col_before}': Mean={df[max_outlier_col_before].mean():.2f}, Std={df[max_outlier_col_before].std():.2f}, "
        desc += f"Min={df[max_outlier_col_before].min():.2f}, Max={df[max_outlier_col_before].max():.2f}"
        self.add_figure(doc, output_dir / f"eda_boxplot_before_all_{english_file_name}.png", f"Feature Outlier Distribution (Before Any Processing) - {file_name}", desc)

        # Feature Histograms (Before Any Processing)
        n_cols = 4
        n_rows = (len(df.columns) + n_cols - 1) // n_cols
        plt.figure(figsize=(15, n_rows * 3))
        for i, col in enumerate(df.columns):
            plt.subplot(n_rows, n_cols, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.hist(data, bins=50, edgecolor='black')
            else:
                plt.text(0.5, 0.5, 'No valid data (all NaN)', ha='center', va='center')
            plt.title(col)
        plt.tight_layout()
        plt.savefig(output_dir / f"eda_histogram_before_all_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()

        desc = f"กราฟนี้แสดงการกระจายของแต่ละฟีเจอร์ก่อนดำเนินการในรูปแบบฮิสโตแกรม ({file_name}) "
        skew_max_before = df.skew().abs().max()
        skew_max_col_before = df.skew().abs().idxmax()
        skew_min_before = df.skew().abs().min()
        skew_min_col_before = df.skew().abs().idxmin()
        desc += f"ฟีเจอร์ที่มี skewness สูงสุดคือ '{skew_max_col_before}' ({skew_max_before:.2f}) "
        desc += f"ฟีเจอร์ที่มี skewness ต่ำสุดคือ '{skew_min_col_before}' ({skew_min_before:.2f}) "
        if skew_max_before > 1:
            desc += "แสดงถึงการกระจายที่เบ้มาก (skewed distribution) "
        elif skew_max_before > 0.5:
            desc += "แสดงถึงการกระจายที่เบ้ปานกลาง "
        else:
            desc += "แสดงถึงการกระจายที่ใกล้เคียงปกติ "
        desc += f"ค่าสถิติของ '{skew_max_col_before}': Min={df[skew_max_col_before].min():.2f}, Max={df[skew_max_col_before].max():.2f}, "
        desc += f"Mean={df[skew_max_col_before].mean():.2f}, Std={df[skew_max_col_before].std():.2f}"
        self.add_figure(doc, output_dir / f"eda_histogram_before_all_{english_file_name}.png", f"Feature Histograms (Before Any Processing) - {file_name}", desc)

        # Timeseries Plot for All Features (Before Any Processing)
        print("Step 6.3: Creating Timeseries Plot for All Features (Before Any Processing)...")
        n_features = len(df.columns)
        plt.figure(figsize=(15, n_features * 4))  # Dynamic height based on number of features
        for i, col in enumerate(df.columns):
            plt.subplot(n_features, 1, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.plot(df.index[df[col].notna()], data, label=col)
            else:
                plt.text(0.5, 0.5, 'No valid data (all NaN)', ha='center', va='center')
            plt.title(f"Timeseries of {col} (Yearly Labels) - {english_file_name}")
            plt.legend()
        plt.tight_layout()
        plt.savefig(output_dir / f"timeseries_plot_before_all_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()


        desc = f"กราฟนี้แสดงข้อมูล Timeseries ของทุกฟีเจอร์ก่อนดำเนินการ โดยใช้ Subplot และแสดง Label แบบ Yearly (เช่น 2010, 2011, ...) ({file_name}) "
        desc += "ความสูงของภาพปรับตามจำนวนฟีเจอร์ (4 นิ้วต่อฟีเจอร์)"
        self.add_figure(doc, output_dir / f"timeseries_plot_before_all_{english_file_name}.png", f"Timeseries Plot for All Features (Before Any Processing) - {file_name}", desc)

        # Remove Features with Missing > 50% (Before Trim, Exclude SO2)
        print("Step 6.4: Removing Features with Missing > 50% (Excluding SO2)...")
        missing_pct = df.isnull().mean()
        features_to_drop = missing_pct[missing_pct > self.constants.MISSING_VALUE_THRESHOLD].index
        features_to_drop = [col for col in features_to_drop if col != 'SO2']  # Exclude SO2
        if features_to_drop:
            df = df.drop(columns=features_to_drop)
            doc.add_paragraph(f"Dropped features (>50% missing, excluding SO2) - {file_name}: {features_to_drop}")
        else:
            doc.add_paragraph(f"No features with missing > 50% (excluding SO2) to drop - {file_name}.")

        # Trim Head/Tails Based on SO2
        print("Step 6.5: Trimming Head/Tails Based on SO2...")
        if 'SO2' in df.columns and df['SO2'].notna().any():
            valid_so2 = df['SO2'].notna()
            if valid_so2.sum() > 0:
                df = df[valid_so2]
                doc.add_paragraph(f"Trimmed head/tails based on non-NaN SO2 values - {file_name}. New shape: {df.shape}")
            else:
                doc.add_paragraph(f"No valid SO2 values found after removing missing > 50% - {file_name}. Using all available data.")
        else:
            doc.add_paragraph(f"SO2 column not found or all NaN after removing missing > 50% - {file_name}. Using all available data.")

        if len(df) == 0:
            doc.add_paragraph(f"No data remaining after preprocessing steps. Exiting. - {file_name}")
            self.set_thsarabupsk_font(doc)
            doc.save(output_dir / f"report_{file_name}_empty.docx")
            return None


        # Fill missing dates and interpolate
        all_dates = pd.date_range(start=df.index.min(), end=df.index.max(), freq='H')
        df = df.reindex(all_dates).interpolate(method='linear')

        # Drop any remaining NaN rows after interpolation
        df_trimmed = df.dropna()
        
        if 'SO2' not in df_trimmed.columns:
             doc.add_paragraph("SO2 IS NOT FOUND")
             self.set_thsarabupsk_font(doc)
             doc.save(output_dir / f"report_{file_name}_SO2_NOT_FOUND.docx")
             return None
        
        return df_trimmed

    def prepare_data_for_modeling(self, df, doc, file_name):
         """Prepares data for LSTM and XGBoost."""
         # Data Preprocessing

         print("Step 7: Preparing Data for LSTM...")
         doc.add_heading(f"Data Preprocessing for LSTM - {file_name}", 1)
         if 'SO2' not in df.columns:
              doc.add_paragraph(f"SO2 column missing from the data. Cannot proceed. - {file_name}")
              self.set_thsarabupsk_font(doc)
              doc.save(output_dir/f"report_{file_name}_error_SO2missing.docx")
              return None, None, None, None, None, None
         features = [col for col in df.columns if col != 'SO2']  # Exclude SO2 from features
         if not features:
             doc.add_paragraph("No features other than SO2 found.  Cannot proceed.")
             return None, None, None, None, None, None

         scaler_features = MinMaxScaler()
         scaler_target = MinMaxScaler()
         df[features] = scaler_features.fit_transform(df[features])
         df['SO2'] = scaler_target.fit_transform(df[['SO2']])


         # Prepare sequences for LSTM
         print("Step 7.1: Preparing Sequences for LSTM...")
         X_full, y_full = self.prepare_sequences(df, features, self.constants.SEQ_LENGTH, self.constants.PRED_LENGTH)
         print(f"X_full shape: {X_full.shape}, y_full shape: {y_full.shape}")

         # Prepare flat data for XGBoost SHAP
         X_full_flat = df[features].values
         y_full_flat = df['SO2'].values
         print(f"X_full_flat shape: {X_full_flat.shape}, y_full_flat shape: {y_full_flat.shape}")
         return X_full, y_full, X_full_flat, y_full_flat, features, scaler_target


    def train_lstm_model(self, X_full, y_full, features, scaler_target, doc, file_name, english_file_name, output_dir):
        """Trains the LSTM model."""
        print("Step 8: Training LSTM Model with All Features...")
        doc.add_heading(f"Model Training with LSTM - {file_name}", 1)
        doc.add_paragraph(
            f"Model: LSTM\nBatch Size: {self.constants.LSTM_BATCH_SIZE}\nEpochs: {self.constants.LSTM_EPOCHS}\nSplits: {self.constants.TIME_SERIES_SPLITS}\nFeatures Used: {list(features)} ({file_name})")

        tscv = TimeSeriesSplit(n_splits=self.constants.TIME_SERIES_SPLITS)
        batch_size = self.constants.LSTM_BATCH_SIZE
        scores = {"mse": [], "mae": [], "r2": []}

        for fold, (train_idx, test_idx) in enumerate(tscv.split(X_full)):
            print(f"Step 8.{fold + 1}: Training Fold {fold + 1} - {file_name}...")
            X_train, X_test = X_full[train_idx], X_full[test_idx]
            y_train, y_test = y_full[train_idx], y_full[test_idx]

            # Create tf.data.Dataset for batching
            train_dataset = tf.data.Dataset.from_tensor_slices((X_train, y_train)).batch(batch_size)
            test_dataset = tf.data.Dataset.from_tensor_slices((X_test, y_test)).batch(batch_size)

            # Clear Keras backend
            tf.keras.backend.clear_session()

            model = Sequential([
                LSTM(128, input_shape=(self.constants.SEQ_LENGTH, len(features)), return_sequences=True),
                Dropout(0.2),
                LSTM(64),
                Dropout(0.2),
                Dense(32, activation='relu'),
                Dense(self.constants.PRED_LENGTH)
            ])
            model.compile(optimizer='adam', loss='mse')

            # Fit model
            history = model.fit(train_dataset, epochs=self.constants.LSTM_EPOCHS, validation_data=test_dataset, verbose=1)

            # Predict and handle NaN
            y_pred = model.predict(test_dataset)

            y_test_inv = scaler_target.inverse_transform(y_test)
            y_pred_inv = scaler_target.inverse_transform(y_pred)

            # Check for NaN in y_test_inv and y_pred_inv before metrics
            if np.any(np.isnan(y_test_inv)) or np.any(np.isnan(y_pred_inv)):
                doc.add_paragraph(
                    f"Warning: y_test_inv or y_pred_inv in Fold {fold + 1} contains NaN values after inverse transform ({file_name})")
                # Remove or replace NaN before calculating metrics
                mask = ~np.isnan(y_test_inv) & ~np.isnan(y_pred_inv)
                if np.any(mask):  # If there are valid data points
                    y_test_inv_clean = y_test_inv[mask]
                    y_pred_inv_clean = y_pred_inv[mask]
                else:
                    doc.add_paragraph(f"No valid data in Fold {fold + 1}. Skipping metrics calculation ({file_name})")
                    continue  # Skip this fold if no valid data
            else:
                y_test_inv_clean = y_test_inv
                y_pred_inv_clean = y_pred_inv

            mse = mean_squared_error(y_test_inv_clean, y_pred_inv_clean)
            mae = mean_absolute_error(y_test_inv_clean, y_pred_inv_clean)
            r2 = r2_score(y_test_inv_clean.flatten(), y_pred_inv_clean.flatten())
            scores["mse"].append(mse)
            scores["mae"].append(mae)
            scores["r2"].append(r2)

            # Loss plot
            plt.figure(figsize=(10, 6))
            plt.plot(history.history['loss'], label='Training Loss')
            plt.plot(history.history['val_loss'], label='Validation Loss')
            plt.title(f"Fold {fold + 1} Loss - {english_file_name}")
            plt.legend()
            plt.savefig(output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png", dpi=300, bbox_inches='tight')
            plt.close()

            train_loss_mean = np.mean(history.history['loss'])
            train_loss_std = np.std(history.history['loss'])
            val_loss_mean = np.mean(history.history['val_loss'])
            val_loss_std = np.std(history.history['val_loss'])
            desc = f"กราฟนี้แสดง Loss จากการฝึก Fold {fold + 1} ({file_name}) "
            desc += f"ค่าเฉลี่ย Loss การฝึก: {train_loss_mean:.4f} (±{train_loss_std:.4f}), การทดสอบ: {val_loss_mean:.4f} (±{val_loss_std:.4f}) "
            if val_loss_mean > train_loss_mean * 1.2:
                desc += "แสดงถึงภาวะ overfitting ที่อาจเกิดขึ้นในช่วงท้ายการฝึก เนื่องจาก Loss ทดสอบสูงกว่าการฝึกอย่างมีนัยสำคัญ "
            elif val_loss_mean < train_loss_mean * 0.8:
                desc += "แสดงถึงภาวะ underfitting หรือการฝึกที่ยังไม่สมบูรณ์ เนื่องจาก Loss ทดสอบต่ำกว่าการฝึกอย่างมาก "
            else:
                desc += "แสดงถึงการฝึกที่สมดุลดี ระหว่างการฝึกและทดสอบ "
            desc += f"จำนวน epochs ที่ใช้: 50, Batch Size: {batch_size}"
            self.add_figure(doc, output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png",
                       f"Fold {fold + 1} Loss - {file_name}", desc)

            # Actual vs Predicted
            plt.figure(figsize=(12, 6))
            plt.plot(y_test_inv_clean.flatten()[:720], label='Actual')
            plt.plot(y_pred_inv_clean.flatten()[:720], label='Predicted', linestyle='--')
            plt.title(f"Fold {fold + 1} Actual vs Predicted (First 30 Days) - {english_file_name}")
            plt.legend()
            plt.savefig(output_dir / f"fold_{fold + 1}_pred_{english_file_name}.png", dpi=300, bbox_inches='tight')
            plt.close()

            pred_diff = np.abs(y_test_inv_clean.flatten()[:720] - y_pred_inv_clean.flatten()[:720])
            max_diff_idx = np.argmax(pred_diff)
            max_diff_date = None
            test_start_idx = test_idx[0]
            date_idx = test_start_idx + self.constants.SEQ_LENGTH + max_diff_idx
            if date_idx < len(df.index):
                max_diff_date = df.index[date_idx].strftime('%Y-%m-%d %H:%M')
            actual_mean = y_test_inv_clean.mean()
            pred_mean = y_pred_inv_clean.mean()
            actual_std = y_test_inv_clean.std()
            pred_std = y_pred_inv_clean.std()
            desc = f"กราฟนี้แสดงการเปรียบเทียบระหว่างค่าจจริงและค่าพยากรณ์ Fold {fold + 1} (30 วันแรก) ({file_name}) "
            desc += f"ค่าเฉลี่ย SO2 จจริง: {actual_mean:.2f} (±{actual_std:.2f}), ค่าพยากรณ์: {pred_mean:.2f} (±{pred_std:.2f}) "
            if pred_diff.max() > actual_std * 2:
                if max_diff_date:
                    desc += f"จุดที่มีความคลาดเคลื่อนสูงสุด ({pred_diff.max():.2f}) เกิดขึ้นวันที่ {max_diff_date} "
                else:
                    desc += f"จุดที่มีความคลาดเคลื่อนสูงสุด ({pred_diff.max():.2f}) อยู่นอกขอบเขตวันที่ที่มีข้อมูล "
                desc += "แสดงถึงความคลาดเคลื่อนที่สูงผิดปกติ อาจเกิดจาก noise หรือ outlier ในข้อมูล "
            else:
                desc += "การพยากรณ์ส่วนใหญ่ใกล้เคียงกับค่าจริง "
            desc += f"วันที่เริ่มต้นของชุดทดสอบ: {df.index[test_start_idx + self.constants.SEQ_LENGTH].strftime('%Y-%m-%d %H:%M')}, "
            desc += f"วันที่สิ้นสุด: {df.index[min(test_start_idx + self.constants.SEQ_LENGTH + 719, len(df.index) - 1)].strftime('%Y-%m-%d %H:%M')}"
            self.add_figure(doc, output_dir / f"fold_{fold + 1}_pred_{english_file_name}.png",
                       f"Fold {fold + 1} Predictions - {file_name}", desc)

            doc.add_heading(f"Fold {fold+1} Metrics - {file_name}", 2)
            doc.add_paragraph(f"MSE: {mse:.4f}\nMAE: {mae:.4f}\nR²: {r2:.4f}")

        # Average metrics
        if scores["mse"]:  # Check if there are any scores to average
            avg_metrics = pd.DataFrame([
                {"Metric": "MSE", "Value": np.mean(scores["mse"])},
                {"Metric": "MAE", "Value": np.mean(scores["mae"])},
                {"Metric": "R²", "Value": np.mean(scores["r2"])}
            ])
            desc = f"ตารางนี้แสดงค่าเฉลี่ยของเมตริกทั้ง {len(scores['mse'])} folds ที่มีข้อมูลสมบูรณ์ ({file_name}) "
            r2_avg = avg_metrics['Value'][avg_metrics['Metric'] == 'R²'].values[0]
            mse_avg = avg_metrics['Value'][avg_metrics['Metric'] == 'MSE'].values[0]
            mae_avg = avg_metrics['Value'][avg_metrics['Metric'] == 'MAE'].values[0]
            desc += f"R²: {r2_avg:.4f}, MSE: {mse_avg:.4f}, MAE: {mae_avg:.4f} "
            if r2_avg > 0.8:
                desc += "R² สูงกว่า 0.8 แสดงถึงความแม่นยำสูงมากของโมเดล "
            elif r2_avg < 0.5:
                desc += "R² ต่ำกว่า 0.5 แสดงถึงความแม่นยำที่อาจไม่เพียงพอ "
            else:
                desc += "R² อยู่ในระดับปานกลาง "
            desc += "MSE และ MAE สะท้อนระดับของความคลาดเคลื่อนในการพยากรณ์"
            self.add_table(doc, avg_metrics, f"Average Metrics Across Folds - {file_name}", desc)
        else:
            doc.add_paragraph(f"Cannot calculate average metrics due to no complete data in all folds ({file_name})")


    def shap_analysis(self, X_full_flat, y_full_flat, features, doc, file_name, english_file_name, output_dir):
        """Performs SHAP analysis using XGBoost."""

        print("Step 9: Computing SHAP Tree Importance with XGBoost...")
        doc.add_heading(f"Feature Importance with SHAP Tree via XGBoost (Post-Training) - {file_name}", 1)

        # Train XGBoost model for SHAP analysis
        xgb_model = xgboost.XGBRegressor(
            objective='reg:squarederror',
            tree_method='hist',  # Updated per XGBoost 2.0+ recommendation
            device='cuda',      # Use CUDA for GPU acceleration
            random_state=self.constants.SEED_VALUE
        )
        xgb_model.fit(X_full_flat, y_full_flat)

        # Use Tree Explainer
        X_sample_flat = X_full_flat[-100:]  # Sample for SHAP
        explainer = shap.TreeExplainer(xgb_model)
        shap_values = explainer.shap_values(X_sample_flat)

        # Convert shap_values to Explanation object
        shap_explanation = shap.Explanation(
            values=shap_values,                                   # SHAP values
            base_values=explainer.expected_value,                 # Base value
            data=X_sample_flat,                                  # Input data
            feature_names=features                                # Feature names
        )

        # Calculate mean absolute SHAP values for feature importance
        shap_importance = np.abs(shap_values).mean(axis=0)
        feature_importance = pd.DataFrame({
            'Feature': features,
            'Importance': shap_importance,
            'Min_SHAP': np.min(shap_values, axis=0),
            'Max_SHAP': np.max(shap_values, axis=0),
            'Mean_SHAP': np.mean(shap_values, axis=0),
            'Std_SHAP': np.std(shap_values, axis=0)
        }).sort_values('Importance', ascending=False)

        # Save SHAP summary plots - Both Bar and Beeswarm
        # Bar plot (Absolute)
        plt.figure(figsize=(10, 6))
        shap.plots.bar(shap_explanation, max_display=len(features), show=False)
        plt.title(f"SHAP Feature Importance (Bar) - {english_file_name}")
        plt.savefig(output_dir / f"shap_bar_{english_file_name}.png", dpi=300, bbox_inches='tight', format='png')
        plt.close()

        # Beeswarm plot
        plt.figure(figsize=(12, 8))
        shap.plots.beeswarm(shap_explanation, max_display=len(features), show=False)
        plt.title(f"SHAP Feature Importance (Beeswarm) - {english_file_name}")
        plt.savefig(output_dir / f"shap_beeswarm_{english_file_name}.png", dpi=300, bbox_inches='tight', format='png')
        plt.close()

        # Ensure the file paths are correct and the files exist before adding to document
        bar_plot_path = output_dir / f"shap_bar_{english_file_name}.png"
        beeswarm_plot_path = output_dir / f"shap_beeswarm_{english_file_name}.png"

        max_fi_feature = feature_importance.iloc[0]['Feature']
        max_fi_value = feature_importance.iloc[0]['Importance']
        max_shap_min = feature_importance.iloc[0]['Min_SHAP']
        max_shap_max = feature_importance.iloc[0]['Max_SHAP']
        desc = f"This graph shows the feature importance from XGBoost SHAP Tree Explainer ({file_name}) "
        desc += f"The most important feature is '{max_fi_feature}' ({max_fi_value:.4f}) "
        desc += f"SHAP values of '{max_fi_feature}' range from {max_shap_min:.4f} to {max_shap_max:.4f} "
        if max_fi_value > shap_importance.mean() * 2:
            desc += "indicating a significantly higher influence compared to other features "
        else:
            desc += "indicating an influence similar to other features "
        desc += f"SHAP value distribution shows both positive and negative impacts on the prediction"
        self.add_figure(doc, bar_plot_path, f"SHAP Tree Summary Plot (Bar) - {file_name}", desc)
        self.add_figure(doc, beeswarm_plot_path, f"SHAP Tree Summary Plot (Beeswarm) - {file_name}", desc)

        # Add feature importance table
        desc = f"This table shows the ranking of feature importance from SHAP Tree. Average Importance: {shap_importance.mean():.4f}, "
        desc += f"Std Importance: {shap_importance.std():.4f}, Min Importance: {shap_importance.min():.4f}, Max Importance: {shap_importance.max():.4f} ({file_name}) "
        if feature_importance['Importance'].std() > shap_importance.mean() * 0.5:
            desc += "the distribution of importance is highly variable "
        else:
            desc += "the distribution of importance is relatively consistent "
        self.add_table(doc, feature_importance, f"Feature Importance Ranking (SHAP Tree via XGBoost) - {file_name}", desc)



    def final_lstm_prediction(self, df, X_full, features, scaler_target, doc, file_name, english_file_name, output_dir):
        """Generates the final LSTM prediction."""
        print("Step 10: Generating Final Prediction with LSTM...")
        doc.add_heading(f"Final Prediction with LSTM - {file_name}", 1)
        # Re-instantiate and train on the FULL dataset.  Very important for final prediction.
        tf.keras.backend.clear_session()
        model = Sequential([
                LSTM(128, input_shape=(self.constants.SEQ_LENGTH, len(features)), return_sequences=True),
                Dropout(0.2),
                LSTM(64),
                Dropout(0.2),
                Dense(32, activation='relu'),
                Dense(self.constants.PRED_LENGTH)
            ])
        model.compile(optimizer='adam', loss='mse')
        model.fit(X_full, y_full, epochs=self.constants.LSTM_EPOCHS, batch_size=self.constants.LSTM_BATCH_SIZE, verbose=1) # Train on the *entire* dataset.


        last_data = X_full[-1:]
        pred_next_24 = model.predict(last_data)
        pred_next_24_inv = scaler_target.inverse_transform(pred_next_24)
        future_timestamps = pd.date_range(start=df.index[-1], periods=25, freq='H')[1:]

        # Check for NaN in actual data for the last 24 hours
        actual_last_24 = df['SO2'][-24:].values.reshape(-1, 1)
        actual_last_24_inv = scaler_target.inverse_transform(actual_last_24)

        plt.figure(figsize=(12, 6))
        plt.plot(df.index[-24:], actual_last_24_inv, label='Actual')
        plt.plot(future_timestamps, pred_next_24_inv[0], label='Predicted', linestyle='--')
        plt.title(f"Next 24-Hour SO2 Prediction (LSTM) - {english_file_name}")
        plt.legend()
        plt.savefig(output_dir / f"final_pred_{english_file_name}.png", dpi=300, bbox_inches='tight')
        plt.close()

        actual_last_mean = actual_last_24_inv.mean()
        actual_last_std = actual_last_24_inv.std()
        actual_last_min = actual_last_24_inv.min()
        actual_last_max = actual_last_24_inv.max()
        pred_mean = pred_next_24_inv[0].mean()
        pred_std = pred_next_24_inv[0].std()
        pred_min = pred_next_24_inv[0].min()
        pred_max = pred_next_24_inv[0].max()
        desc = f"กราฟนี้แสดงการพยากรณ์ SO2 24 ชั่วโมงสุดท้าย ({file_name}) "
        desc += f"ค่าเฉลี่ยจริง: {actual_last_mean:.2f} (±{actual_last_std:.2f}), ค่าพยากรณ์: {pred_mean:.2f} (±{pred_std:.2f} "
        desc += f"Min จริง: {actual_last_min:.2f}, Max จริง: {actual_last_max:.2f}, "
        desc += f"Min พยากรณ์: {pred_min:.2f}, Max พยากรณ์: {pred_max:.2f} "
        if abs(pred_mean - actual_last_mean) > actual_last_mean * 0.2:
            desc += "การพยากรณ์มีความคลาดเคลื่อนมากกว่า 20% จากค่าจริง "
        else:
            desc += "การพยากรณ์ใกล้เคียงกับค่าจจริงในระดับที่น่าพอใจ "
        desc += f"วันที่เริ่มต้นพยากรณ์: {future_timestamps[0].strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {future_timestamps[-1].strftime('%Y-%m-%d %H:%M')}"
        self.add_figure(doc, output_dir / f"final_pred_{english_file_name}.png", f"Next 24-Hour Prediction (LSTM) - {file_name}", desc)

    def save_report(self, doc, output_dir, file_name):
        """Saves the report to a .docx file."""
        print("Step 11: Saving Report...")
        self.set_thsarabupsk_font(doc)
        doc.save(output_dir / f"report_{file_name}.docx")
        print(f"Step 12: Process Completed for {file_name}!")


    def prepare_sequences(self, data, features, seq_length, pred_length):
        """Prepares sequences for LSTM."""
        X, y = [], []
        for i in range(len(data) - seq_length - pred_length + 1):
            X.append(data[features].iloc[i:i + seq_length].values)
            y.append(data['SO2'].iloc[i + seq_length:i + seq_length + pred_length].values)
        return np.array(X), np.array(y)

    # Helper functions for .docx reporting (placed inside the class)
    def add_table(self, doc, df, title, description=""):
        """Adds a table to the document."""
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
        table.style.border_width = 1
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = col
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.iloc[i, j])
        if description:
            doc.add_paragraph(description)

    def add_figure(self, doc, filename, title, description=""):
        """Adds a figure to the document."""
        if isinstance(filename, Path):  # Check if filename is a Path object
            filename = str(filename)    # Convert Path to string
        if os.path.exists(filename):   # Check if the file exists
            doc.add_heading(title, level=2)
            doc.add_picture(filename, width=Inches(6))
            if description:
                doc.add_paragraph(description)
        else:
            doc.add_heading(title, level=2)
            doc.add_paragraph(f"Warning: Image file {filename} not found.")

    def set_thsarabupsk_font(self, doc, font_name="TH SarabunPSK"):
        """Sets the specified font for all text in a Word document."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name = font_name
                font._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                font._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                font._element.rPr.rFonts.set(qn('w:cs'), font_name)
                font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                font._element.rPr.rFonts.set(qn('w:complexScript'), font_name)

        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        font._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        font._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        font._element.rPr.rFonts.set(qn('w:cs'), font_name)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        font._element.rPr.rFonts.set(qn('w:complexScript'), font_name)

def main():
    """Main function to execute the data processing."""
    directory = input("Enter the directory containing the CSV files: ")  # Get directory from user
    processor = DataProcessor(directory)
    processor.process_files()


if __name__ == "__main__":
    main()