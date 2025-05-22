# Import Libraries
print("Step 1: Importing Libraries...")
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.linear_model import LinearRegression
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
from tensorflow.keras.callbacks import EarlyStopping
import xgboost
import shap
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import os
from pathlib import Path
import gc
from docx.oxml.ns import qn


# --- Helper Functions for Report Generation (Mostly Unchanged) ---
def describe_dataframe_skewness_for_report(df, doc, section_title, english_file_name):
    # This function is now specifically for adding to docx, not returning a string
    doc.add_heading(
        f"{section_title} - Skewness Analysis - {english_file_name}", level=3
    )
    numeric_df = df.select_dtypes(include=np.number)
    if numeric_df.empty:
        doc.add_paragraph("No numeric features to calculate skewness.")
        return

    try:
        skew_series = numeric_df.skew()
        if skew_series.empty or skew_series.isnull().all():
            doc.add_paragraph(
                "Cannot calculate meaningful skewness (e.g., all NaNs or insufficient data in numeric columns)."
            )
            return
    except Exception as e:
        doc.add_paragraph(f"Error calculating skewness: {e}")
        return

    abs_skew_series = skew_series.abs().dropna()
    # Drop NaNs from abs_skew_series before finding idxmax/min
    if abs_skew_series.empty:
        doc.add_paragraph("No valid (non-NaN) absolute skewness values found.")
        return

    skew_max_val = abs_skew_series.max()
    skew_max_col = abs_skew_series.idxmax()
    skew_min_val = abs_skew_series.min()
    skew_min_col = abs_skew_series.idxmin()

    desc_para = doc.add_paragraph()
    desc_para.add_run(f"Feature with highest absolute skewness: ").bold = True
    desc_para.add_run(f"'{skew_max_col}' ({skew_max_val:.2f}).\n")
    desc_para.add_run(f"Feature with lowest absolute skewness: ").bold = True
    desc_para.add_run(f"'{skew_min_col}' ({skew_min_val:.2f}).\n")

    if skew_max_val > 1:
        desc_para.add_run(
            "The highest skewness indicates a highly skewed distribution. "
        )
    elif skew_max_val > 0.5:
        desc_para.add_run(
            "The highest skewness indicates a moderately skewed distribution. "
        )
    else:
        desc_para.add_run(
            "The highest skewness indicates an approximately symmetric distribution. "
        )

    try:
        if skew_max_col in df.columns:
            stats_series = df[skew_max_col]
            desc_para.add_run(
                f"\nStats for '{skew_max_col}': Min={stats_series.min():.2f}, Max={stats_series.max():.2f}, Mean={stats_series.mean():.2f}, Std={stats_series.std():.2f}."
            )
        else:
            desc_para.add_run(
                f"\nCould not retrieve stats for '{skew_max_col}' (column not found)."
            )
    except Exception as e_stats:
        desc_para.add_run(f"\nError retrieving stats for '{skew_max_col}': {e_stats}.")


def set_default_font(doc, font_name="TH SarabunPSK", font_size=11):
    # Set font for the entire document body
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        for footer in section.footer.paragraphs:
            for run in footer.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    # Set font for all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # Set font for styles
    styles = doc.styles
    for style_name in [
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Caption",
        "List Paragraph",
    ]:
        try:
            style = styles[style_name]
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            # For headings, you might want different sizes
            if "Heading" in style_name:
                if style_name == "Heading 1":
                    font.size = Pt(16)
                    font.bold = True
                elif style_name == "Heading 2":
                    font.size = Pt(14)
                    font.bold = True
                elif style_name == "Heading 3":
                    font.size = Pt(12)
                    font.bold = True

            # Ensure complex script and East Asia fonts are also set if the font supports them
            # This is crucial for Thai characters to render correctly with TH SarabunPSK
            if hasattr(font._element, "rPr") and font._element.rPr is not None:
                # Check if rFonts exists, if not, create it
                rfonts = font._element.rPr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = type(
                        font._element.rPr
                    )._new_rFonts()  # Create new rFonts element
                    font._element.rPr.append(rfonts)  # Append it to rPr

                rfonts.set(qn("w:ascii"), font_name)
                rfonts.set(qn("w:hAnsi"), font_name)
                rfonts.set(qn("w:cs"), font_name)  # Critical for Thai
                rfonts.set(qn("w:eastAsia"), font_name)  # Also important for Thai
        except KeyError:
            print(
                f"Warning: Style '{style_name}' not found. Font not applied to this style."
            )


def add_table_to_doc(doc, df, title, description=""):  # Renamed to avoid conflict
    doc.add_heading(title, level=2)
    if not isinstance(df, pd.DataFrame) or df.empty:
        doc.add_paragraph(
            f"Cannot add table '{title}': DataFrame is not valid or empty."
        )
        if description:
            doc.add_paragraph(description)
        return

    try:
        # Convert all data to string to avoid type issues with docx table cells
        df_str = df.astype(str)
        table = doc.add_table(
            rows=df_str.shape[0] + 1, cols=df_str.shape[1], style="Table Grid"
        )

        for j, col_name in enumerate(df_str.columns):
            table.cell(0, j).text = str(col_name)
        for i in range(df_str.shape[0]):
            for j in range(df_str.shape[1]):
                table.cell(i + 1, j).text = df_str.iloc[i, j]
        if description:
            doc.add_paragraph(description)
    except Exception as e:
        doc.add_paragraph(f"Error adding table '{title}': {e}")
        print(f"Error adding table '{title}' to docx: {e}")


def add_figure_to_doc(doc, filename, title, description=""):  # Renamed
    if isinstance(filename, Path):
        filename = str(filename)
    doc.add_heading(title, level=2)
    if os.path.exists(filename):
        try:
            doc.add_picture(filename, width=Inches(6))
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Warning: Could not add image {filename}. Error: {e}")
    else:
        doc.add_paragraph(f"Warning: Image file {filename} not found.")
    if description:
        doc.add_paragraph(description)


def prepare_sequences(data, features, seq_length, pred_length):
    X, y = [], []
    if not features or "SO2" not in data.columns:
        print(
            "Warning (prepare_sequences): Features list is empty or SO2 column is missing."
        )
        return np.array(X), np.array(y)

    valid_features = [f for f in features if f in data.columns]
    if not valid_features:
        print(
            f"Warning (prepare_sequences): None of the provided features {features} found in data columns."
        )
        return np.array(X), np.array(y)

    if len(data) < seq_length + pred_length:  # Not enough data
        print(
            f"Warning (prepare_sequences): Data length ({len(data)}) is less than seq_length ({seq_length}) + pred_length ({pred_length})."
        )
        return np.array(X), np.array(y)

    for i in range(len(data) - seq_length - pred_length + 1):
        X.append(data[valid_features].iloc[i : i + seq_length].values)
        y.append(data["SO2"].iloc[i + seq_length : i + seq_length + pred_length].values)

    if not X or not y:
        return np.array(X), np.array(y)

    return np.array(X).astype(np.float32), np.array(y).astype(np.float32)


# --- Main Class ---
class DynamicInputDataset:
    class CONST:
        SEED_VALUE = 42
        BATCH_SIZE = 1024
        EPOCHS = 50  # Consider reducing for faster runs during dev, e.g., 10-20
        TEST_SPLIT_RATIO = 0.2
        SEQ_LENGTH = 24
        PRED_LENGTH = 24
        OUTPUT_FOLDER = "outputs_detailed_eda"
        NUM_TOP_FEATURES_FOR_SELECTION = 5
        SHAP_IMPORTANCE_THRESHOLD = 0.005  # Example threshold
        ITERATIVE_PRED_HORIZON_DAYS = 365  # For 1-year prediction

    file_name_mapping = {
        "(37t)ศาลหลักเมือง(ปิดสถานี)": "San_Lak_Mueang_Closed_Station",
        "(37t)สถานีอุตุนิยมวิทยาลำปาง": "Lampang_Meteorological_Station",
        "(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด": "Ban_Sopad_Hospital",
        "(39t)โรงพยาบาลส่งเสริมสุขภาพตำบลท่าสี": "Tha_Si_Hospital",
        "(40t)การประปาส่วนภูมิภาคแม่เมาะ": "Mae_Mo_Waterworks",
    }

    def __init__(self, input_dir):
        self.input_dir = Path(input_dir)
        self.csv_files = sorted(
            list(self.input_dir.glob("*.csv"))
        )  # Sort for consistent order
        self.output_dir = Path.cwd() / self.CONST.OUTPUT_FOLDER
        self.output_dir.mkdir(exist_ok=True)
        random.seed(self.CONST.SEED_VALUE)
        np.random.seed(self.CONST.SEED_VALUE)
        tf.random.set_seed(self.CONST.SEED_VALUE)
        gpus = tf.config.experimental.list_physical_devices("GPU")
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
                print("GPU memory growth enabled.")
            except RuntimeError as e:
                print(f"Error setting GPU memory growth: {e}")
        print(f"Num GPUs Available: {len(gpus)}")

    def _perform_full_eda(
        self, doc, df_input, eda_stage_name, english_file_name, file_name_base
    ):
        """Helper function to perform and report a full EDA stage."""
        doc.add_heading(
            f"Exploratory Data Analysis ({eda_stage_name}) - {file_name_base}", level=1
        )
        doc.add_paragraph(f"Current Data Shape: {df_input.shape}")
        doc.add_paragraph(f"Total Missing Values: {df_input.isnull().sum().sum()}")

        # 1. Data Preview
        doc.add_heading(f"Data Preview ({eda_stage_name})", level=2)
        if not df_input.empty:
            # Columns List
            doc.add_heading(f"Columns List ({eda_stage_name})", level=3)
            for col in df_input.columns:
                doc.add_paragraph(f"- {col}")
            # Sample of First 10 Rows
            add_table_to_doc(
                doc,
                df_input.head(10),
                f"Sample of First 10 Rows ({eda_stage_name})",
                f"Showing the first 10 rows of data at the '{eda_stage_name}' stage.",
            )
            # Data Information (df.info())
            doc.add_heading(
                f"Data Information (dtypes, non-null count) ({eda_stage_name})", level=3
            )

            # Capture df.info() output
            from io import StringIO

            buffer = StringIO()
            df_input.info(buf=buffer)
            info_str = buffer.getvalue()
            doc.add_paragraph(info_str)

        else:
            doc.add_paragraph(
                f"DataFrame is empty at '{eda_stage_name}' stage. Cannot show preview."
            )
            return  # Stop EDA for this stage if df is empty

        # EDA Plots
        plot_suffix = f"{eda_stage_name.lower().replace(' ', '_')}_{english_file_name}"

        # Missing Data Heatmap
        plt.figure(figsize=(12, 6))
        sns.heatmap(
            df_input.isnull(), cbar=True, cmap="viridis", annot=False
        )  # Changed cmap
        plt.title(f"Missing Data Heatmap ({eda_stage_name}) - {english_file_name}")
        heatmap_path = self.output_dir / f"heatmap_{plot_suffix}.png"
        plt.savefig(heatmap_path, dpi=300, bbox_inches="tight")
        plt.close()
        add_figure_to_doc(
            doc,
            heatmap_path,
            f"Missing Data Heatmap ({eda_stage_name})",
            f"Heatmap of missing values at the '{eda_stage_name}' stage.",
        )

        if "SO2" in df_input.columns and pd.api.types.is_numeric_dtype(df_input["SO2"]):
            numeric_df_for_corr = df_input.select_dtypes(include=np.number)
            if not numeric_df_for_corr.empty and "SO2" in numeric_df_for_corr.columns:
                corr_matrix = numeric_df_for_corr.corr()
                # Sorted Correlation Matrix with SO2
                if "SO2" in corr_matrix.columns:
                    so2_corr = corr_matrix[["SO2"]].sort_values(
                        by="SO2", ascending=False
                    )
                    add_table_to_doc(
                        doc,
                        so2_corr,
                        f"Sorted Correlation with SO2 ({eda_stage_name})",
                        f"Correlation of numeric features with SO2 at the '{eda_stage_name}' stage.",
                    )
                # Full Correlation Matrix Heatmap
                plt.figure(figsize=(12, 10))
                sns.heatmap(
                    corr_matrix, annot=True, cmap="coolwarm", fmt=".2f", linewidths=0.5
                )
                plt.title(
                    f"Correlation Matrix ({eda_stage_name}) - {english_file_name}"
                )
                plt.xticks(rotation=45, ha="right")
                plt.yticks(rotation=0)
                corr_heatmap_path = self.output_dir / f"corr_heatmap_{plot_suffix}.png"
                plt.savefig(corr_heatmap_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    corr_heatmap_path,
                    f"Correlation Matrix Heatmap ({eda_stage_name})",
                    f"Heatmap of correlations between numeric features at the '{eda_stage_name}' stage.",
                )
            else:
                doc.add_paragraph(
                    f"Not enough numeric data or SO2 column for correlation analysis at '{eda_stage_name}'."
                )
        else:
            doc.add_paragraph(
                f"SO2 column not found or not numeric for correlation analysis at '{eda_stage_name}'."
            )

        # Feature Outlier Distribution (Boxplot)
        numeric_cols_for_boxplot = df_input.select_dtypes(include=np.number).columns
        if not numeric_cols_for_boxplot.empty:
            plt.figure(
                figsize=(15, max(6, len(numeric_cols_for_boxplot) * 0.5))
            )  # Dynamic height
            df_input[numeric_cols_for_boxplot].boxplot(rot=45, grid=False)
            plt.title(
                f"Feature Outlier Distribution ({eda_stage_name}) - {english_file_name}"
            )
            plt.tight_layout()
            boxplot_path = self.output_dir / f"boxplot_{plot_suffix}.png"
            plt.savefig(boxplot_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                boxplot_path,
                f"Feature Outlier Distribution ({eda_stage_name})",
                f"Boxplots showing the distribution and outliers for numeric features at '{eda_stage_name}'.",
            )
        else:
            doc.add_paragraph(
                f"No numeric columns for boxplot analysis at '{eda_stage_name}'."
            )

        # Feature Histograms & Skewness
        numeric_cols_for_hist = df_input.select_dtypes(include=np.number).columns
        if not numeric_cols_for_hist.empty:
            n_cols_hist = min(4, len(numeric_cols_for_hist))
            n_rows_hist = (len(numeric_cols_for_hist) + n_cols_hist - 1) // n_cols_hist
            plt.figure(figsize=(15, n_rows_hist * 3.5))
            for i, col in enumerate(numeric_cols_for_hist):
                plt.subplot(n_rows_hist, n_cols_hist, i + 1)
                sns.histplot(df_input[col].dropna(), kde=True)
                plt.title(col)
            plt.tight_layout()
            hist_path = self.output_dir / f"histograms_{plot_suffix}.png"
            plt.savefig(hist_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                hist_path,
                f"Feature Histograms ({eda_stage_name})",
                f"Histograms showing the distribution of numeric features at '{eda_stage_name}'.",
            )
            describe_dataframe_skewness_for_report(
                df_input, doc, f"Skewness ({eda_stage_name})", english_file_name
            )  # Add skewness text to doc
        else:
            doc.add_paragraph(
                f"No numeric columns for histogram/skewness analysis at '{eda_stage_name}'."
            )

        # Timeseries Plot for All Features
        if isinstance(df_input.index, pd.DatetimeIndex) and not df_input.empty:
            # Plot numeric features only for timeseries plots to avoid errors with non-numeric types
            numeric_cols_for_ts = df_input.select_dtypes(include=np.number).columns
            if not numeric_cols_for_ts.empty:
                n_features_ts = len(numeric_cols_for_ts)
                plt.figure(
                    figsize=(15, n_features_ts * 2.5 if n_features_ts > 0 else 5)
                )  # Adjust height
                for i, col in enumerate(numeric_cols_for_ts):
                    plt.subplot(n_features_ts, 1, i + 1)
                    plt.plot(df_input.index, df_input[col], label=col)
                    plt.title(f"Timeseries of {col} - {english_file_name}")
                    plt.legend(loc="upper right")
                plt.tight_layout()
                ts_path = self.output_dir / f"timeseries_all_{plot_suffix}.png"
                plt.savefig(ts_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    ts_path,
                    f"Timeseries Plot for All Numeric Features ({eda_stage_name})",
                    f"Timeseries plots for all numeric features at '{eda_stage_name}'.",
                )
            else:
                doc.add_paragraph(
                    f"No numeric columns to plot as timeseries at '{eda_stage_name}'."
                )

        else:
            doc.add_paragraph(
                f"Index is not DatetimeIndex or DataFrame is empty. Skipping timeseries plots for '{eda_stage_name}'."
            )

    def _plot_train_test_split_bar_chart(
        self, train_size, test_size, model_suffix, english_file_name, doc
    ):
        # (Code from previous answer, seems fine)
        dataset_info = pd.DataFrame([{"Train": train_size, "Validation": test_size}])
        plt.figure(figsize=(6, 4))
        dataset_info.plot(kind="barh", stacked=True, ax=plt.gca())
        total_text = train_size + test_size
        train_percentage = (train_size / total_text * 100) if total_text > 0 else 0
        validation_percentage = (test_size / total_text * 100) if total_text > 0 else 0
        plt.text(
            total_text,
            0,
            f"{total_text}",
            ha="left",
            va="center",
            color="black",
            fontsize=9,
        )
        plt.text(
            train_size / 2,
            0,
            f"{train_size}\n({train_percentage:.1f}%)",
            ha="center",
            va="center",
            color="white",
            fontsize=9,
        )
        plt.text(
            train_size + test_size / 2,
            0,
            f"{test_size}\n({validation_percentage:.1f}%)",
            ha="center",
            va="center",
            color="black",
            fontsize=9,
        )
        plt.xlabel("Dataset Size (Sequences)")
        plt.ylabel("Split")
        plt.title(
            f"Train-Test Split ({model_suffix.replace('_', ' ').title()}) - {english_file_name}"
        )
        plt.yticks([])
        plt.legend(["Train", "Validation"], loc="lower right")
        chart_path = (
            self.output_dir / f"train_test_split_{model_suffix}_{english_file_name}.png"
        )
        plt.savefig(
            chart_path,
            dpi=300,
            bbox_inches="tight",
            facecolor="white",
            transparent=False,
        )
        plt.close()
        add_figure_to_doc(
            doc,
            chart_path,
            f"Train-Test Split ({model_suffix.replace('_', ' ').title()})",
            f"Bar chart showing the number of sequences in training and testing sets for the {model_suffix.replace('_', ' ')} model.",
        )

    def _train_evaluate_report_lstm(
        self,
        doc,
        df_scaled_input,
        features_to_use,
        scaler_target,
        model_suffix,
        english_file_name,
        file_name_base,
        original_df_for_time_ref,  # Unscaled df for time referencing
    ):
        report_title_suffix = model_suffix.replace("_", " ").title()
        doc.add_heading(
            f"LSTM Model Training & Evaluation ({report_title_suffix}) - {file_name_base}",
            1,
        )

        if not features_to_use:
            doc.add_paragraph(
                f"No features provided for '{report_title_suffix}' model. Skipping training."
            )
            return None, {}, None, None

        doc.add_paragraph(f"Features used: {', '.join(features_to_use)}")
        doc.add_paragraph(
            f"Sequence Length: {self.CONST.SEQ_LENGTH}, Prediction Length: {self.CONST.PRED_LENGTH}"
        )

        X_full, y_full = prepare_sequences(
            df_scaled_input,
            features_to_use,
            self.CONST.SEQ_LENGTH,
            self.CONST.PRED_LENGTH,
        )

        if X_full.size == 0 or y_full.size == 0:
            doc.add_paragraph(
                f"Not enough data for sequences ({report_title_suffix}). X shape: {X_full.shape}, y shape: {y_full.shape}. Skipping."
            )
            return None, {}, None, None

        X_train, X_test, y_train, y_test = train_test_split(
            X_full,
            y_full,
            test_size=self.CONST.TEST_SPLIT_RATIO,
            shuffle=False,
            random_state=self.CONST.SEED_VALUE,
        )
        test_start_offset_in_original_df = len(X_full) - len(X_test)

        self._plot_train_test_split_bar_chart(
            len(X_train), len(X_test), model_suffix, english_file_name, doc
        )

        model_filename_h5 = self.output_dir / f"{file_name_base}-lstm-{model_suffix}.h5"
        history_data = {"loss": [], "val_loss": []}
        is_model_trained_this_run = False
        model = None

        if model_filename_h5.exists():
            print(f"Loading pre-trained model: {model_filename_h5}")
            try:
                model = tf.keras.models.load_model(model_filename_h5, compile=True)
            except Exception as e:
                print(
                    f"Failed to load compiled model {model_filename_h5}, attempting compile=False. Error: {e}"
                )
                try:
                    model = tf.keras.models.load_model(model_filename_h5, compile=False)
                    model.compile(optimizer="adam", loss="mse")
                except Exception as e_recompile:
                    print(
                        f"Failed to load and recompile model {model_filename_h5}: Error: {e_recompile}"
                    )
                    doc.add_paragraph(
                        f"Error loading model {model_filename_h5}. Training new one if possible."
                    )
                    model = None  # Force retrain

        if model is None:  # Retrain if loading failed or model doesn't exist
            print(f"Training new model for {model_suffix}...")
            tf.keras.backend.clear_session()  # Clear previous session
            model = Sequential(
                [
                    LSTM(
                        128,
                        input_shape=(self.CONST.SEQ_LENGTH, X_train.shape[-1]),
                        return_sequences=True,
                    ),
                    Dropout(0.2),
                    LSTM(64, return_sequences=False),
                    Dropout(0.2),
                    Dense(32, activation="relu"),
                    Dense(self.CONST.PRED_LENGTH),
                ]
            )
            model.compile(
                optimizer=tf.keras.optimizers.Adam(learning_rate=0.001), loss="mse"
            )  # Added learning rate

            early_stopping = EarlyStopping(
                monitor="val_loss", patience=10, restore_best_weights=True, verbose=1
            )
            history = model.fit(
                X_train,
                y_train,
                epochs=self.CONST.EPOCHS,
                batch_size=self.CONST.BATCH_SIZE,
                validation_data=(
                    X_test,
                    y_test,
                ),  # X_test, y_test are already float32 from prepare_sequences
                verbose=1,
                callbacks=[early_stopping],
            )
            history_data = history.history
            model.save(model_filename_h5)
            is_model_trained_this_run = True

        model.summary(print_fn=lambda x: doc.add_paragraph(x))

        y_pred_scaled = model.predict(X_test)
        y_test_inv = scaler_target.inverse_transform(np.nan_to_num(y_test))
        y_pred_inv = scaler_target.inverse_transform(np.nan_to_num(y_pred_scaled))

        valid_mask = (
            ~np.isnan(y_test_inv).any(axis=1) & ~np.isnan(y_pred_inv).any(axis=1)
            if y_test_inv.ndim > 1
            else ~np.isnan(y_test_inv) & ~np.isnan(y_pred_inv)
        )
        y_test_inv_clean = y_test_inv[valid_mask]
        y_pred_inv_clean = y_pred_inv[valid_mask]

        scores = {"mse": np.nan, "mae": np.nan, "r2": np.nan}
        if y_test_inv_clean.size > 0 and y_pred_inv_clean.size > 0:
            scores["mse"] = mean_squared_error(y_test_inv_clean, y_pred_inv_clean)
            scores["mae"] = mean_absolute_error(y_test_inv_clean, y_pred_inv_clean)
            scores["r2"] = r2_score(y_test_inv_clean, y_pred_inv_clean)
        else:
            doc.add_paragraph(
                f"Not enough valid data after inverse transform for metrics ({report_title_suffix})."
            )

        # Reporting Plots
        if (
            is_model_trained_this_run
            and history_data.get("loss")
            and history_data.get("val_loss")
        ):
            if (
                history_data["loss"] and history_data["val_loss"]
            ):  # Check if lists are not empty
                plt.figure(figsize=(10, 6))
                plt.plot(history_data["loss"], label="Training Loss")
                plt.plot(history_data["val_loss"], label="Validation Loss")
                plt.title(f"Model Loss ({report_title_suffix}) - {english_file_name}")
                plt.xlabel("Epoch")
                plt.ylabel("Loss (MSE)")
                plt.legend()
                loss_plot_path = (
                    self.output_dir / f"loss_{model_suffix}_{english_file_name}.png"
                )
                plt.savefig(loss_plot_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    loss_plot_path,
                    f"Model Loss ({report_title_suffix})",
                    "Training and validation loss.",
                )
        else:
            doc.add_paragraph(
                f"Loss plot not generated for {report_title_suffix} (model loaded or no history)."
            )

        y_test_plot_main = (
            y_test_inv_clean[:, 0]
            if y_test_inv_clean.ndim > 1 and y_test_inv_clean.shape[1] > 0
            else y_test_inv_clean.flatten()
        )
        y_pred_plot_main = (
            y_pred_inv_clean[:, 0]
            if y_pred_inv_clean.ndim > 1 and y_pred_inv_clean.shape[1] > 0
            else y_pred_inv_clean.flatten()
        )

        # Actual vs. Predicted Time Series (24h from test set)
        num_points_24h = min(
            self.CONST.PRED_LENGTH * 1, len(y_test_plot_main)
        )  # Show first prediction window
        if num_points_24h > 0:
            plt.figure(figsize=(12, 6))
            plt.plot(
                y_test_plot_main[:num_points_24h],
                label="Actual SO2 (Test Set)",
                marker=".",
            )
            plt.plot(
                y_pred_plot_main[:num_points_24h],
                label="Predicted SO2 (Test Set)",
                linestyle="--",
                marker="x",
            )
            plt.title(
                f"Actual vs. Predicted SO2 ({report_title_suffix}, First {num_points_24h} Hours of Test) - {english_file_name}"
            )
            plt.xlabel(f"Time Step (Hours)")
            plt.ylabel("SO2 Value")
            plt.legend()
            ts_plot_path_24h = (
                self.output_dir
                / f"timeseries_pred_24h_{model_suffix}_{english_file_name}.png"
            )
            plt.savefig(ts_plot_path_24h, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                ts_plot_path_24h,
                f"Actual vs. Predicted (First {num_points_24h}h Test, {report_title_suffix})",
                f"Comparison for the first {num_points_24h} hours of the test set.",
            )

        # Scatter Plots
        time_frames_scatter = {
            "Test Set (All Cleaned Data)": (y_test_plot_main, y_pred_plot_main)
        }
        # Add 24h scatter if enough data
        if num_points_24h > 1:
            time_frames_scatter[f"Test Set (First {num_points_24h} Hours)"] = (
                y_test_plot_main[:num_points_24h],
                y_pred_plot_main[:num_points_24h],
            )

        for frame_name, (
            y_actual_scatter,
            y_pred_scatter,
        ) in time_frames_scatter.items():
            if len(y_actual_scatter) > 1:
                plt.figure(figsize=(8, 8))
                plt.scatter(
                    y_actual_scatter,
                    y_pred_scatter,
                    alpha=0.6,
                    label="Actual vs. Predicted",
                )
                min_val = (
                    min(y_actual_scatter.min(), y_pred_scatter.min())
                    if len(y_actual_scatter) > 0
                    else 0
                )
                max_val = (
                    max(y_actual_scatter.max(), y_pred_scatter.max())
                    if len(y_actual_scatter) > 0
                    else 1
                )
                plt.plot(
                    [min_val, max_val],
                    [min_val, max_val],
                    "k--",
                    lw=2,
                    label="Ideal y=x Line",
                )
                if len(np.unique(y_actual_scatter)) > 1:
                    reg = LinearRegression().fit(
                        y_actual_scatter.reshape(-1, 1), y_pred_scatter
                    )
                    r2_scat = reg.score(y_actual_scatter.reshape(-1, 1), y_pred_scatter)
                    plt.plot(
                        y_actual_scatter,
                        reg.predict(y_actual_scatter.reshape(-1, 1)),
                        color="red",
                        label=f"Fit (R²={r2_scat:.2f})",
                    )
                plt.xlabel("Actual SO2")
                plt.ylabel("Predicted SO2")
                plt.title(
                    f"Actual vs. Predicted Scatter ({frame_name}, {report_title_suffix}) - {english_file_name}"
                )
                plt.legend()
                plt.grid(True)
                scatter_fname_suffix = (
                    frame_name.lower()
                    .replace(" ", "_")
                    .replace("(", "")
                    .replace(")", "")
                )
                scatter_plot_path = (
                    self.output_dir
                    / f"scatter_{scatter_fname_suffix}_{model_suffix}_{english_file_name}.png"
                )
                plt.savefig(scatter_plot_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    scatter_plot_path,
                    f"Scatter Plot ({frame_name}, {report_title_suffix})",
                    f"Scatter plot for {frame_name}.",
                )
            else:
                doc.add_paragraph(
                    f"Not enough data for scatter plot ({frame_name}, {report_title_suffix})."
                )

        doc.add_heading(f"Performance Metrics ({report_title_suffix})", 2)
        if not all(np.isnan(v) for v in scores.values()):
            metrics_df = pd.DataFrame([scores]).round(4)
            add_table_to_doc(
                doc,
                metrics_df,
                f"Metrics Summary ({report_title_suffix})",
                "Performance metrics on the test set.",
            )
        else:
            doc.add_paragraph(
                f"Metrics could not be calculated for {report_title_suffix}."
            )

        return model, scores, y_test_inv_clean, y_pred_inv_clean

    def _perform_iterative_forecast(
        self,
        doc,
        model,
        df_scaled_input,
        features_to_use,
        scaler_target,
        model_suffix,
        english_file_name,
        file_name_base,
        original_df_for_time_ref,
        horizon_days,
    ):
        doc.add_heading(
            f"Iterative Forecast ({horizon_days} Days - {model_suffix}) - {file_name_base}",
            2,
        )
        if (
            model is None
            or df_scaled_input.empty
            or not features_to_use
            or scaler_target is None
        ):
            doc.add_paragraph(
                "Cannot perform iterative forecast due to missing model, data, features, or scaler."
            )
            return None, None

        # Get the last known sequence from the full scaled dataset
        last_known_sequence_scaled, _ = prepare_sequences(
            df_scaled_input, features_to_use, self.CONST.SEQ_LENGTH, 0
        )  # pred_length = 0
        if last_known_sequence_scaled.size == 0:
            doc.add_paragraph(
                "Could not get last known sequence for iterative forecast."
            )
            return None, None

        current_sequence = last_known_sequence_scaled[-1].reshape(
            1, self.CONST.SEQ_LENGTH, len(features_to_use)
        )

        all_predictions_scaled = []
        num_prediction_steps = (
            horizon_days * 24 // self.CONST.PRED_LENGTH
        )  # How many PRED_LENGTH blocks

        for _ in range(num_prediction_steps):
            pred_block_scaled = model.predict(
                current_sequence
            )  # Shape (1, PRED_LENGTH)
            all_predictions_scaled.append(
                pred_block_scaled[0]
            )  # Append the (PRED_LENGTH,) array

            # Prepare next input: roll sequence, append new *predicted* SO2 (and other features if multi-variate in a more complex setup)
            # For simplicity here, we assume other features remain constant or are also predicted (complex).
            # A common simplification for uni-variate target with exogenous vars is to repeat last known exogenous vars
            # or use their forecasted values if available.
            # Here, we only update the SO2 part if we were predicting multi-feature outputs.
            # Since our LSTM predicts only SO2 for PRED_LENGTH steps, we need to construct the *next* input sequence.
            # This is the hardest part of iterative forecasting.

            # Simple approach: take the *first* step of the predicted block to update the sequence if SEQ_LENGTH > PRED_LENGTH
            # If PRED_LENGTH is what we shift by, we'd use the whole block.
            # Assuming we shift by PRED_LENGTH (e.g., predict 24h, then use those 24h to predict next 24h)

            # This part needs careful thought based on how features are structured for input.
            # If features_to_use includes lagged SO2, it's more straightforward.
            # If features_to_use are purely exogenous, they need to be forecasted too or assumed.

            # Crude way: Shift current_sequence, and for the 'SO2' feature part, plug in the new prediction.
            # This assumes 'SO2' is one of the features_to_use or that we can construct a new feature vector.
            # This is highly complex if not just predicting SO2 based on past SO2.

            # Let's assume for now this iterative forecast is primarily illustrative and might not be robust
            # without a more sophisticated feature update mechanism.
            # We'll just append the predicted SO2 values.

            # If the model input features (X_train.shape[-1]) only contained lagged values of SO2 and other
            # features, and the output is just SO2, then constructing the next sequence is about taking
            # the last (SEQ_LENGTH - PRED_LENGTH) from the previous input and appending the PRED_LENGTH predictions.

            # If model was (past_features, past_SO2) -> future_SO2
            # And if future_features are needed, they must be known or forecasted.
            # For now, let's assume the features in `current_sequence` are sufficient and we just need to roll them.
            # This is a simplification.
            new_features_for_sequence = current_sequence[
                0, self.CONST.PRED_LENGTH :, :
            ].copy()  # Take last (SEQ_LENGTH - PRED_LENGTH) steps of features

            # We need to append PRED_LENGTH steps of *all* features.
            # The model only predicts SO2. So, we need to decide what the other features will be.
            # Simplest (often poor) assumption: repeat last known values of other features.

            new_so2_block_scaled = pred_block_scaled.reshape(
                self.CONST.PRED_LENGTH, 1
            )  # (24, 1)

            # Construct PRED_LENGTH new feature vectors
            appended_feature_vectors = []
            so2_feature_index = (
                -1
            )  # Assume SO2 is the last feature if not explicitly known
            if (
                "SO2_lagged" in features_to_use
            ):  # Example if SO2 itself (lagged) is a feature
                so2_feature_index = features_to_use.index(
                    "SO2_lagged"
                )  # Find actual index

            for i in range(self.CONST.PRED_LENGTH):
                new_vector = current_sequence[
                    0, -1, :
                ].copy()  # Start with the last known feature vector
                if so2_feature_index != -1:
                    new_vector[so2_feature_index] = new_so2_block_scaled[i, 0]
                # For other features, they'd ideally be forecasted or carried over.
                # If SO2 is the ONLY feature, this is simpler.
                # If features_to_use has only one element (e.g. 'SO2_scaled_lagged'):
                elif len(features_to_use) == 1:
                    new_vector[0] = new_so2_block_scaled[i, 0]

                appended_feature_vectors.append(new_vector)

            if not appended_feature_vectors:  # Should not happen if PRED_LENGTH > 0
                doc.add_paragraph(
                    "Iterative forecast: Failed to create new feature vectors. Stopping."
                )
                break

            current_sequence = np.vstack(
                [new_features_for_sequence, np.array(appended_feature_vectors)]
            )
            current_sequence = current_sequence.reshape(
                1, self.CONST.SEQ_LENGTH, len(features_to_use)
            )

        if not all_predictions_scaled:
            doc.add_paragraph("No predictions made in iterative forecast.")
            return None, None

        all_predictions_scaled_np = np.concatenate(
            all_predictions_scaled, axis=0
        )  # From list of (PRED_LENGTH,) to (TotalPredLen,)
        all_predictions_inv = scaler_target.inverse_transform(
            all_predictions_scaled_np.reshape(-1, 1)
        )  # Reshape for scaler
        all_predictions_inv = all_predictions_inv.flatten()  # Back to 1D

        # Time index for the forecast
        forecast_time_index = None
        if (
            isinstance(original_df_for_time_ref.index, pd.DatetimeIndex)
            and not original_df_for_time_ref.empty
        ):
            last_known_time = original_df_for_time_ref.index[-1]
            forecast_time_index = pd.date_range(
                start=last_known_time + pd.Timedelta(hours=1),
                periods=len(all_predictions_inv),
                freq="h",
            )
        else:  # Fallback numeric index
            forecast_time_index = np.arange(len(all_predictions_inv))

        # Plot
        plt.figure(figsize=(15, 7))
        # Plot some recent actual data if available
        if (
            isinstance(original_df_for_time_ref.index, pd.DatetimeIndex)
            and "SO2" in original_df_for_time_ref
        ):
            recent_actual_len = min(
                len(original_df_for_time_ref), 24 * 7
            )  # Last 7 days actual
            plt.plot(
                original_df_for_time_ref.index[-recent_actual_len:],
                original_df_for_time_ref["SO2"].iloc[-recent_actual_len:],
                label="Recent Actual SO2",
            )

        plt.plot(
            forecast_time_index,
            all_predictions_inv,
            label=f"Iterative Forecast ({horizon_days} Days)",
            linestyle="--",
            color="red",
        )
        plt.title(
            f"{horizon_days}-Day Iterative SO2 Forecast ({model_suffix}) - {english_file_name}"
        )
        plt.xlabel("Time")
        plt.ylabel("SO2 Value")
        plt.legend()
        plt.grid(True)
        if isinstance(forecast_time_index, pd.DatetimeIndex):
            plt.xticks(rotation=30, ha="right")
        iter_forecast_plot_path = (
            self.output_dir
            / f"iter_forecast_{horizon_days}d_{model_suffix}_{english_file_name}.png"
        )
        plt.savefig(iter_forecast_plot_path, dpi=300, bbox_inches="tight")
        plt.close()
        add_figure_to_doc(
            doc,
            iter_forecast_plot_path,
            f"{horizon_days}-Day Iterative Forecast ({model_suffix})",
            f"Iterative SO2 forecast for {horizon_days} days using the {model_suffix} model. Note: Long-term iterative forecasts can have high uncertainty.",
        )

        return forecast_time_index, all_predictions_inv

    def process_file(self, file_path):
        print(f"\n--- Processing file: {file_path.name} ---")
        file_name_base = file_path.stem
        english_file_name = self.file_name_mapping.get(file_name_base, file_name_base)

        doc = Document()
        set_default_font(doc)  # Apply default font to the new document
        doc.add_heading(f"LSTM Model Report for SO2 Prediction - {file_name_base}", 0)
        doc.add_paragraph(f"Dataset: {file_path.name}")

        try:
            df_original = pd.read_csv(
                file_path, parse_dates=["Datetime"], index_col="Datetime"
            )
        except Exception as e:
            doc.add_paragraph(f"Error loading CSV {file_path.name}: {e}")
            report_filename = (
                self.output_dir / f"report_{file_name_base}_load_error.docx"
            )
            doc.save(report_filename)
            print(f"Failed to load {file_path.name}, report saved: {report_filename}")
            return

        df = df_original.head(3000).copy()  # Limit rows if needed
        df = df_original.copy()  # Limit rows if needed
        if "SO2" in df.columns:
            df["SO2"] = df["SO2"].clip(lower=0)
        else:
            doc.add_paragraph(
                f"Critical: SO2 column not found in {file_path.name}. Processing cannot continue meaningfully."
            )
            report_filename = self.output_dir / f"report_{file_name_base}_no_SO2.docx"
            set_default_font(doc)
            doc.save(report_filename)
            return

        if "dev" in df.columns:
            df = df.drop(columns=["dev"])

        # --- EDA Stage 1: Before Any Processing ---
        self._perform_full_eda(
            doc, df.copy(), "Before Any Processing", english_file_name, file_name_base
        )

        df_after_initial_cleaning = df.copy()

        # --- Data Cleaning Step ---
        doc.add_heading("Data Cleaning Steps", level=1)
        # Remove features with > 50% missing (excluding SO2)
        missing_pct = df_after_initial_cleaning.isnull().mean()
        features_to_drop_missing = missing_pct[missing_pct > 0.5].index.tolist()
        features_to_drop_missing = [
            col for col in features_to_drop_missing if col != "SO2"
        ]  # Keep SO2

        if features_to_drop_missing:
            df_after_initial_cleaning = df_after_initial_cleaning.drop(
                columns=features_to_drop_missing
            )
            doc.add_paragraph(
                f"Dropped features with >50% missing (excluding SO2): {', '.join(features_to_drop_missing)}. New shape: {df_after_initial_cleaning.shape}"
            )
        else:
            doc.add_paragraph(
                "No features had >50% missing values (excluding SO2) to drop."
            )

        # --- EDA Stage 2: After Removing High Missing % Features ---
        self._perform_full_eda(
            doc,
            df_after_initial_cleaning.copy(),
            "After Removing High Missing Features",
            english_file_name,
            file_name_base,
        )

        # Trim head/tails based on SO2 first/last valid index
        df_after_so2_trim = df_after_initial_cleaning.copy()
        if (
            "SO2" in df_after_so2_trim.columns
            and df_after_so2_trim["SO2"].notna().any()
        ):
            first_valid_idx = df_after_so2_trim["SO2"].first_valid_index()
            last_valid_idx = df_after_so2_trim["SO2"].last_valid_index()
            if first_valid_idx is not None and last_valid_idx is not None:
                df_after_so2_trim = df_after_so2_trim.loc[
                    first_valid_idx:last_valid_idx
                ]
                doc.add_paragraph(
                    f"Trimmed data based on first/last valid SO2 datetime. New shape: {df_after_so2_trim.shape}"
                )
            else:
                doc.add_paragraph(
                    "SO2 column has no valid values; SO2-based trimming not performed."
                )
        else:
            doc.add_paragraph(
                "SO2 column missing or all NaN; SO2-based trimming not performed."
            )

        # --- EDA Stage 3: After Trimming by SO2 ---
        self._perform_full_eda(
            doc,
            df_after_so2_trim.copy(),
            "After Trimming by SO2",
            english_file_name,
            file_name_base,
        )

        # --- Statistical Summary Tables (After Trimming/Remove Missing) ---
        doc.add_heading("Data Summaries After Cleaning & Trimming", level=1)
        if not df_after_so2_trim.empty:
            # Data Information Table (like .describe() but more detailed for docx)
            doc.add_heading(
                f"Data Column Information (After Cleaning/Trimming)", level=2
            )
            summary_list = []
            for col in df_after_so2_trim.columns:
                col_data = df_after_so2_trim[col]
                is_numeric = pd.api.types.is_numeric_dtype(col_data)
                summary_list.append(
                    {
                        "Column": col,
                        "DataType": str(col_data.dtype),
                        "Count": col_data.count(),
                        "Unique": col_data.nunique(),
                        "Nulls": col_data.isnull().sum(),
                        "Mean": f"{col_data.mean():.2f}" if is_numeric else "N/A",
                        "Std": f"{col_data.std():.2f}" if is_numeric else "N/A",
                        "Min": f"{col_data.min():.2f}"
                        if is_numeric
                        else col_data.min()
                        if not col_data.empty and not is_numeric
                        else "N/A",  # Handle non-numeric min
                        "Max": f"{col_data.max():.2f}"
                        if is_numeric
                        else col_data.max()
                        if not col_data.empty and not is_numeric
                        else "N/A",  # Handle non-numeric max
                    }
                )
            summary_df = pd.DataFrame(summary_list)
            add_table_to_doc(
                doc,
                summary_df,
                "Detailed Column Information Table",
                "Summary statistics for each column after cleaning and SO2 trimming.",
            )

            # Standard Statistical Summary Table (numeric only, like .describe())
            numeric_df_final = df_after_so2_trim.select_dtypes(include=np.number)
            if not numeric_df_final.empty:
                stats_summary_table = numeric_df_final.describe().transpose().round(2)
                add_table_to_doc(
                    doc,
                    stats_summary_table,
                    "Statistical Summary of Numeric Features (After Cleaning/Trimming)",
                    "Descriptive statistics for numeric features after all cleaning.",
                )
            else:
                doc.add_paragraph(
                    "No numeric features remaining for statistical summary table after cleaning."
                )
        else:
            doc.add_paragraph(
                "DataFrame is empty after cleaning and trimming. No summary tables generated."
            )

        # --- Data Preprocessing for LSTM (Interpolation & Scaling) ---
        doc.add_heading(f"Data Preprocessing for LSTM - {file_name_base}", 1)
        df_to_process = df_after_so2_trim.copy()
        df_scaled = pd.DataFrame()
        scaler_features_global = None  # One scaler for all features
        scaler_target_global = None  # One scaler for the target
        features_all_scaled = []
        original_df_unscaled_for_lstm = (
            pd.DataFrame()
        )  # To store the version right before scaling

        if df_to_process.empty:
            doc.add_paragraph("Data is empty before interpolation. Skipping LSTM part.")
            set_default_font(doc)
            doc.save(self.output_dir / f"report_{file_name_base}_empty_pre_interp.docx")
            return

        if (
            not isinstance(df_to_process.index, pd.DatetimeIndex)
            or not df_to_process.index.is_monotonic_increasing
        ):
            df_to_process = df_to_process.sort_index()
            doc.add_paragraph("Ensured DataFrame index is sorted before interpolation.")

        # Interpolate
        try:
            if (
                isinstance(df_to_process.index, pd.DatetimeIndex)
                and df_to_process.index.nunique() > 1
            ):  # Ensure multiple unique time points for range
                all_dates = pd.date_range(
                    start=df_to_process.index.min(),
                    end=df_to_process.index.max(),
                    freq="h",
                )
                df_reindexed = df_to_process.reindex(all_dates)
                df_interpolated = df_reindexed.interpolate(
                    method="linear", limit_direction="both"
                )  # limit_direction helps with edges
                doc.add_paragraph(
                    f"Interpolated data to hourly frequency. Shape after reindex & interpolate: {df_interpolated.shape}"
                )
            else:  # Not datetime index or not enough unique points
                df_interpolated = df_to_process.interpolate(
                    method="linear", limit_direction="both"
                )
                doc.add_paragraph(
                    f"Interpolated data (no reindexing). Shape after interpolate: {df_interpolated.shape}"
                )
        except Exception as e_interp:
            doc.add_paragraph(
                f"Error during interpolation: {e_interp}. Using data before interpolation attempt."
            )
            df_interpolated = df_to_process.copy()

        df_processed_for_scaling = df_interpolated.dropna(
            how="all"
        )  # Drop rows if ALL are NaN
        doc.add_paragraph(
            f"Shape after dropping rows with all NaNs (post-interpolation): {df_processed_for_scaling.shape}"
        )

        # Critical check: if SO2 is all NaN or missing after this
        if (
            "SO2" not in df_processed_for_scaling.columns
            or df_processed_for_scaling["SO2"].isnull().all()
        ):
            doc.add_paragraph(
                f"Critical: SO2 column is missing or all NaN after interpolation and initial dropna. Cannot proceed with scaling/training for {file_name_base}."
            )
            set_default_font(doc)
            doc.save(
                self.output_dir / f"report_{file_name_base}_no_SO2_post_interp.docx"
            )
            return

        # Final attempt to fill any remaining NaNs before scaling (e.g., with 0 or ffill/bfill)
        # Using ffill then bfill is often better than just 0 for time series
        df_filled_for_scaling = df_processed_for_scaling.fillna(method="ffill").fillna(
            method="bfill"
        )
        # If any NaNs still remain (e.g., whole column was NaN), fill with 0
        if df_filled_for_scaling.isnull().any().any():
            df_filled_for_scaling = df_filled_for_scaling.fillna(0)
            doc.add_paragraph(
                "Warning: Some NaNs remained after ffill/bfill; these were filled with 0. This might affect model performance."
            )

        if df_filled_for_scaling.empty:
            doc.add_paragraph(
                f"Data became empty after all filling attempts before scaling for {file_name_base}. Skipping LSTM."
            )
            set_default_font(doc)
            doc.save(self.output_dir / f"report_{file_name_base}_empty_pre_scale.docx")
            return

        original_df_unscaled_for_lstm = (
            df_filled_for_scaling.copy()
        )  # This is the correctly processed, unscaled data
        df_scaled = original_df_unscaled_for_lstm.copy()

        features_all_scaled = [col for col in df_scaled.columns if col != "SO2"]

        if features_all_scaled:
            scaler_features_global = MinMaxScaler()
            df_scaled[features_all_scaled] = scaler_features_global.fit_transform(
                df_scaled[features_all_scaled]
            )
        else:
            doc.add_paragraph("No features (other than SO2) to scale.")

        scaler_target_global = MinMaxScaler()
        df_scaled[["SO2"]] = scaler_target_global.fit_transform(df_scaled[["SO2"]])
        doc.add_paragraph(f"Data scaled. Final shape for LSTM input: {df_scaled.shape}")

        # --- LSTM Training Round 1: All Features ---
        print(f"Starting LSTM Training Round 1 (All Features) for {file_name_base}...")
        model_all_features, scores_all_features, y_test_all_inv, y_pred_all_inv = (
            self._train_evaluate_report_lstm(
                doc,
                df_scaled.copy(),
                features_all_scaled,
                scaler_target_global,
                "all_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm,
            )
        )

        # --- SHAP Feature Importance ---
        doc.add_heading(f"Feature Importance (SHAP via XGBoost) - {file_name_base}", 1)
        selected_features_for_round2 = []
        feature_importance_df = pd.DataFrame()

        if features_all_scaled and not df_scaled[features_all_scaled].empty:
            X_flat_shap = df_scaled[
                features_all_scaled
            ].values  # Already scaled, NaNs handled by fillna(0) earlier
            y_flat_shap = df_scaled["SO2"].values

            if (
                X_flat_shap.shape[0] > 1
                and y_flat_shap.shape[0] > 1
                and X_flat_shap.shape[0] == y_flat_shap.shape[0]
            ):
                try:
                    print(
                        f"Running SHAP for {file_name_base} with X_flat_shap: {X_flat_shap.shape}"
                    )
                    xgb_model = xgboost.XGBRegressor(
                        objective="reg:squarederror",
                        random_state=self.CONST.SEED_VALUE,
                        n_estimators=100,
                    )  # Added n_estimators
                    xgb_model.fit(X_flat_shap, y_flat_shap)

                    sample_size_shap = min(
                        200, X_flat_shap.shape[0]
                    )  # Increase sample for SHAP
                    X_sample_shap = (
                        shap.sample(
                            X_flat_shap,
                            sample_size_shap,
                            random_state=self.CONST.SEED_VALUE,
                        )
                        if X_flat_shap.shape[0] > sample_size_shap
                        else X_flat_shap
                    )

                    if X_sample_shap.shape[0] > 0:
                        explainer = shap.TreeExplainer(xgb_model)
                        shap_values_sample = explainer.shap_values(
                            X_sample_shap
                        )  # Use X_sample_shap

                        plt.figure()  # Ensure new figure for each plot
                        shap.summary_plot(
                            shap_values_sample,
                            X_sample_shap,
                            feature_names=features_all_scaled,
                            plot_type="bar",
                            show=False,
                        )
                        plt.title(
                            f"SHAP Feature Importance (Bar) - {english_file_name}"
                        )
                        plt.tight_layout()
                        shap_bar_path = (
                            self.output_dir / f"shap_bar_{english_file_name}.png"
                        )
                        plt.savefig(shap_bar_path, dpi=300, bbox_inches="tight")
                        plt.close()
                        add_figure_to_doc(
                            doc,
                            shap_bar_path,
                            "SHAP Feature Importance (Bar Plot)",
                            "Mean absolute SHAP values.",
                        )

                        plt.figure()
                        shap.summary_plot(
                            shap_values_sample,
                            X_sample_shap,
                            feature_names=features_all_scaled,
                            show=False,
                        )
                        plt.title(
                            f"SHAP Feature Importance (Beeswarm) - {english_file_name}"
                        )
                        plt.tight_layout()
                        shap_beeswarm_path = (
                            self.output_dir / f"shap_beeswarm_{english_file_name}.png"
                        )
                        plt.savefig(shap_beeswarm_path, dpi=300, bbox_inches="tight")
                        plt.close()
                        add_figure_to_doc(
                            doc,
                            shap_beeswarm_path,
                            "SHAP Feature Importance (Beeswarm Plot)",
                            "Distribution of SHAP values.",
                        )

                        mean_abs_shap = np.abs(shap_values_sample).mean(axis=0)
                        feature_importance_df = (
                            pd.DataFrame(
                                {
                                    "Feature": features_all_scaled,
                                    "Mean_Abs_SHAP_Value": mean_abs_shap,
                                }
                            )
                            .sort_values(by="Mean_Abs_SHAP_Value", ascending=False)
                            .reset_index(drop=True)
                        )
                        add_table_to_doc(
                            doc,
                            feature_importance_df.head(10),
                            "Top 10 SHAP Feature Importance Ranking",
                            "Features ranked by mean absolute SHAP value.",
                        )

                        selected_features_for_round2 = feature_importance_df.head(
                            self.CONST.NUM_TOP_FEATURES_FOR_SELECTION
                        )["Feature"].tolist()
                        doc.add_paragraph(
                            f"Selected features for second training round (SHAP Top {self.CONST.NUM_TOP_FEATURES_FOR_SELECTION}): {', '.join(selected_features_for_round2) if selected_features_for_round2 else 'None'}"
                        )
                    else:
                        doc.add_paragraph("SHAP sample was empty. Skipping SHAP plots.")
                except Exception as e:
                    doc.add_paragraph(
                        f"Error during SHAP analysis: {type(e).__name__} - {e}"
                    )
                    print(
                        f"Error in SHAP for {file_name_base}: {type(e).__name__} - {e}"
                    )
            else:
                doc.add_paragraph(
                    "Not enough data or mismatched shapes for SHAP input."
                )
        else:
            doc.add_paragraph(
                "Skipping SHAP analysis: No features available or scaled feature data is empty."
            )

        # --- LSTM Training Round 2: Selected Features ---
        print(
            f"Starting LSTM Training Round 2 (Selected Features) for {file_name_base}..."
        )
        model_selected_features = None  # Initialize
        scores_selected_features = {}
        if selected_features_for_round2:
            (
                model_selected_features,
                scores_selected_features,
                y_test_sel_inv,
                y_pred_sel_inv,
            ) = self._train_evaluate_report_lstm(
                doc,
                df_scaled.copy(),
                selected_features_for_round2,
                scaler_target_global,
                "selected_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm,
            )
        else:
            doc.add_heading(
                f"LSTM Model Training & Evaluation (Selected Features) - {file_name_base}",
                1,
            )
            doc.add_paragraph(
                "No features selected by SHAP (or SHAP failed). Skipping second training round."
            )

        # --- Final 24-Hour Prediction Plot (using ALL Features model and SELECTED Features model) ---
        doc.add_heading(
            f"Final {self.CONST.PRED_LENGTH}-Hour Prediction Comparison - {file_name_base}",
            1,
        )
        if (
            not original_df_unscaled_for_lstm.empty
            and "SO2" in original_df_unscaled_for_lstm.columns
            and len(original_df_unscaled_for_lstm) >= self.CONST.PRED_LENGTH
        ):
            actual_last_24_inv = (
                original_df_unscaled_for_lstm["SO2"]
                .iloc[-self.CONST.PRED_LENGTH :]
                .values
            )
            time_index_actual = original_df_unscaled_for_lstm.index[
                -self.CONST.PRED_LENGTH :
            ]

            if isinstance(original_df_unscaled_for_lstm.index, pd.DatetimeIndex):
                last_known_time = original_df_unscaled_for_lstm.index[-1]
                future_timestamps_24h = pd.date_range(
                    start=last_known_time + pd.Timedelta(hours=1),
                    periods=self.CONST.PRED_LENGTH,
                    freq="h",
                )
            else:
                future_timestamps_24h = np.arange(len(actual_last_24_inv))  # Fallback

            plt.figure(figsize=(14, 7))
            plt.plot(
                time_index_actual,
                actual_last_24_inv,
                label="Actual SO2 (Last Known)",
                marker="o",
                linewidth=2,
            )

            # All Features Model Prediction
            if model_all_features and features_all_scaled:
                X_full_all_for_pred, _ = prepare_sequences(
                    df_scaled, features_all_scaled, self.CONST.SEQ_LENGTH, 0
                )
                if X_full_all_for_pred.size > 0:
                    pred_next_24_all_scaled = model_all_features.predict(
                        X_full_all_for_pred[-1:]
                    )
                    pred_next_24_all_inv = scaler_target_global.inverse_transform(
                        np.nan_to_num(pred_next_24_all_scaled)
                    )
                    plt.plot(
                        future_timestamps_24h,
                        pred_next_24_all_inv[0],
                        label="Predicted (All Features)",
                        linestyle="--",
                        marker="x",
                    )

            # Selected Features Model Prediction
            if model_selected_features and selected_features_for_round2:
                X_full_sel_for_pred, _ = prepare_sequences(
                    df_scaled, selected_features_for_round2, self.CONST.SEQ_LENGTH, 0
                )
                if X_full_sel_for_pred.size > 0:
                    pred_next_24_sel_scaled = model_selected_features.predict(
                        X_full_sel_for_pred[-1:]
                    )
                    pred_next_24_sel_inv = scaler_target_global.inverse_transform(
                        np.nan_to_num(pred_next_24_sel_scaled)
                    )
                    plt.plot(
                        future_timestamps_24h,
                        pred_next_24_sel_inv[0],
                        label="Predicted (Selected Features)",
                        linestyle=":",
                        marker="s",
                    )

            plt.title(
                f"Next {self.CONST.PRED_LENGTH}-Hour SO2 Prediction Comparison - {english_file_name}"
            )
            plt.xlabel("Time")
            plt.ylabel("SO2 Value")
            plt.legend()
            plt.grid(True)
            if isinstance(time_index_actual, pd.DatetimeIndex):
                plt.xticks(rotation=30, ha="right")
            final_pred_comp_path = (
                self.output_dir / f"final_pred_24h_comparison_{english_file_name}.png"
            )
            plt.savefig(final_pred_comp_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                final_pred_comp_path,
                f"Next {self.CONST.PRED_LENGTH}-Hour Prediction Comparison",
                "Comparison of 24-hour predictions from different models.",
            )
        else:
            doc.add_paragraph(
                f"Not enough data for final {self.CONST.PRED_LENGTH}-hour prediction plot."
            )

        # --- Iterative 1-Year Forecast ---
        if model_all_features:
            self._perform_iterative_forecast(
                doc,
                model_all_features,
                df_scaled.copy(),
                features_all_scaled,
                scaler_target_global,
                "all_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm.copy(),
                self.CONST.ITERATIVE_PRED_HORIZON_DAYS,
            )
        if model_selected_features:
            self._perform_iterative_forecast(
                doc,
                model_selected_features,
                df_scaled.copy(),
                selected_features_for_round2,
                scaler_target_global,
                "selected_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm.copy(),
                self.CONST.ITERATIVE_PRED_HORIZON_DAYS,
            )

        # --- Summary: All Features vs Selected Features ---
        doc.add_heading(
            "Model Performance Comparison (All Features vs. SHAP Selected Features)", 1
        )
        if scores_all_features and scores_selected_features:
            comp_data = {
                "Metric": ["MSE", "MAE", "R2 Score"],
                "All Features": [
                    scores_all_features.get("mse", np.nan),
                    scores_all_features.get("mae", np.nan),
                    scores_all_features.get("r2", np.nan),
                ],
                "Selected Features (SHAP)": [
                    scores_selected_features.get("mse", np.nan),
                    scores_selected_features.get("mae", np.nan),
                    scores_selected_features.get("r2", np.nan),
                ],
            }
            comp_df = pd.DataFrame(comp_data).round(4)
            add_table_to_doc(
                doc,
                comp_df,
                "Metrics Comparison Table",
                "Comparison of performance metrics between the model using all features and the model using SHAP-selected features.",
            )

            # Textual summary
            r2_all = scores_all_features.get("r2", -1)
            r2_selected = scores_selected_features.get("r2", -1)
            summary_text = "Comparing model performance:\n"
            if (
                r2_all > r2_selected and abs(r2_all - r2_selected) > 0.01
            ):  # Threshold for meaningful difference
                summary_text += f"- The model with ALL features performed better (R2: {r2_all:.4f}) than the model with SHAP-selected features (R2: {r2_selected:.4f}).\n"
            elif r2_selected > r2_all and abs(r2_selected - r2_all) > 0.01:
                summary_text += f"- The model with SHAP-SELECTED features performed better (R2: {r2_selected:.4f}) than the model with all features (R2: {r2_all:.4f}). This suggests SHAP successfully identified more impactful features.\n"
            else:  # Similar performance or one is NaN
                summary_text += f"- Both models (All Features R2: {r2_all:.4f}, Selected Features R2: {r2_selected:.4f}) showed similar performance, or one/both could not be reliably scored.\n"

            mae_all = scores_all_features.get("mae", np.inf)
            mae_selected = scores_selected_features.get("mae", np.inf)
            if (
                mae_selected < mae_all and abs(mae_all - mae_selected) > 0.01 * mae_all
            ):  # MAE improved by more than 1%
                summary_text += f"- The Selected Features model had a lower MAE ({mae_selected:.4f} vs {mae_all:.4f}), indicating smaller average errors.\n"
            elif (
                mae_all < mae_selected
                and abs(mae_selected - mae_all) > 0.01 * mae_selected
            ):
                summary_text += f"- The All Features model had a lower MAE ({mae_all:.4f} vs {mae_selected:.4f}).\n"

            doc.add_paragraph(summary_text)

        elif scores_all_features:
            doc.add_paragraph(
                "Only the 'All Features' model was successfully evaluated. No comparison possible."
            )
        else:
            doc.add_paragraph(
                "Neither model variant (All Features or Selected Features) could be fully evaluated for comparison."
            )

        # --- Save Report ---
        print(f"Finalizing and saving report for {file_name_base}...")
        set_default_font(doc)  # Ensure font is applied before saving
        report_filename = self.output_dir / f"report_{file_name_base}.docx"
        try:
            doc.save(report_filename)
            print(f"Report saved: {report_filename}")
        except Exception as e:
            print(f"Error saving report {report_filename}: {e}")
            try:
                fallback_report_filename = (
                    self.output_dir / f"report_{file_name_base}_fallback.docx"
                )
                doc.save(fallback_report_filename)
                print(f"Fallback report saved: {fallback_report_filename}")
            except Exception as e_fb:
                print(f"Error saving fallback report: {e_fb}")

        # Memory cleanup
        del (
            df,
            df_original,
            df_after_initial_cleaning,
            df_after_so2_trim,
            df_scaled,
            original_df_unscaled_for_lstm,
        )
        if "model_all_features" in locals():
            del model_all_features
        if "model_selected_features" in locals():
            del model_selected_features
        if "xgb_model" in locals():
            del xgb_model
        gc.collect()
        tf.keras.backend.clear_session()

    def run(self):
        for csv_file_path in self.csv_files:
            try:
                self.process_file(csv_file_path)
            except Exception as e:
                print(
                    f"--- CRITICAL ERROR processing file {csv_file_path.name}: {type(e).__name__} - {e} ---"
                )
                import traceback

                traceback.print_exc()
            finally:
                gc.collect()
                tf.keras.backend.clear_session()


if __name__ == "__main__":
    input_directory = (
        "/mnt/e/MayThesis2025/src/labs4/s02_v01_02_clean/output-680507-2319"
    )
    if not os.path.exists(input_directory):
        print(f"Error: Input directory does not exist: {input_directory}")
    else:
        processor = DynamicInputDataset(input_directory)
        processor.run()
