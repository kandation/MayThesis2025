# %% [markdown]
# # Predicting SO2 Levels Using LSTM with XGBoost-based SHAP Tree Importance
# This notebook builds an LSTM model with XGBoost-based SHAP Tree importance and detailed reporting.

# %% [code]
# Import Libraries
print("Step 1: Importing Libraries...")
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
import xgboost
import shap
from docx import Document
from docx.shared import Inches
import random
import os

# Set random seeds for reproducibility
print("Step 2: Setting Random Seeds...")
seed_value = 42
random.seed(seed_value)
np.random.seed(seed_value)
tf.random.set_seed(seed_value)

# Configure GPU memory growth for LSTM
print("Step 3: Configuring GPU for LSTM...")
gpus = tf.config.experimental.list_physical_devices('GPU')
if gpus:
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print("GPU is set up and ready for LSTM!")
    except RuntimeError as e:
        print(e)

print("Num GPUs Available: ", len(tf.config.list_physical_devices('GPU')))

# Helper functions for .docx reporting
def add_table(doc, df, title):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])

def add_figure(doc, filename, title):
    doc.add_heading(title, level=2)
    doc.add_picture(filename, width=Inches(6))

# %% [code]
# Load Data
print("Step 4: Loading Data...")
file_path = "/mnt/e/MayThesis2025/cleanned_datasets/(37t)สถานีอุตุนิยมวิทยาลำปาง.csv"
df = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
df = df.iloc[:1000]  # Limit to first 1000 rows for testing
doc = Document()
doc.add_heading("LSTM Model Report for SO2 Prediction with XGBoost SHAP FI", 0)
doc.add_paragraph(f"Dataset: {file_path} (Limited to 1000 rows)")

# %% [code]
# EDA
print("Step 5: Performing EDA...")
doc.add_heading("Exploratory Data Analysis", 1)
doc.add_paragraph(f"Shape: {df.shape}")
doc.add_paragraph(f"Missing Values Total: {df.isnull().sum().sum()}")

plt.figure(figsize=(12, 6))
sns.heatmap(df.isnull(), cbar=False, cmap='viridis')
plt.title("Missing Data Heatmap")
plt.savefig("eda_heatmap.png", dpi=300)
plt.close()
add_figure(doc, "eda_heatmap.png", "Missing Data Heatmap")

plt.figure(figsize=(10, 8))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', fmt='.2f')
plt.title("Correlation Matrix")
plt.savefig("eda_corr.png", dpi=300)
plt.close()
add_figure(doc, "eda_corr.png", "Correlation Matrix")

plt.figure(figsize=(15, 6))
df.boxplot(figsize=(15, 6), flierprops=dict(marker='o', markersize=5))
plt.xticks(rotation=45)
plt.title("Feature Outlier Distribution")
plt.savefig("eda_boxplot.png", dpi=300)
plt.close()
add_figure(doc, "eda_boxplot.png", "Feature Outlier Distribution")

n_cols = 4
n_rows = (len(df.columns) + n_cols - 1) // n_cols
plt.figure(figsize=(15, n_rows * 3))
for i, col in enumerate(df.columns):
    plt.subplot(n_rows, n_cols, i + 1)
    df[col].hist(bins=50)
    plt.title(col)
plt.tight_layout()
plt.savefig("eda_hist.png", dpi=300)
plt.close()
add_figure(doc, "eda_hist.png", "Feature Histograms")

# %% [code]
# Data Preprocessing
print("Step 6: Preparing Data...")
doc.add_heading("Data Preprocessing", 1)

# Fill missing dates
all_dates = pd.date_range(start=df.index.min(), end=df.index.max(), freq='H')
df = df.reindex(all_dates).interpolate(method='linear')

# Drop features with >50% missing values
missing_pct = df.isnull().mean()
features_to_drop = missing_pct[missing_pct > 0.5].index
df = df.drop(columns=features_to_drop)
doc.add_paragraph(f"Dropped features (>50% missing): {list(features_to_drop)}")

# Normalize data
features = df.drop(columns=['SO2']).columns
scaler_features = MinMaxScaler()
scaler_target = MinMaxScaler()
df[features] = scaler_features.fit_transform(df[features])
df['SO2'] = scaler_target.fit_transform(df[['SO2']])

# Drop any remaining NaN rows after interpolation
df = df.dropna()

# Prepare sequences for LSTM
print("Step 7: Preparing Sequences for LSTM...")
def prepare_sequences(data, seq_length, pred_length):
    X, y = [], []
    for i in range(len(data) - seq_length - pred_length + 1):
        X.append(data[features].iloc[i:i + seq_length].values)
        y.append(data['SO2'].iloc[i + seq_length:i + seq_length + pred_length].values)
    return np.array(X), np.array(y)

seq_length = 24
pred_length = 24
X_full, y_full = prepare_sequences(df, seq_length, pred_length)
print(f"X_full shape: {X_full.shape}, y_full shape: {y_full.shape}")

# Prepare flat data for XGBoost SHAP
X_full_flat = df[features].values
y_full_flat = df['SO2'].values
print(f"X_full_flat shape: {X_full_flat.shape}, y_full_flat shape: {y_full_flat.shape}")

# %% [code]
# Training with TimeSeriesSplit using All Features (LSTM)
print("Step 8: Training LSTM Model with All Features...")
doc.add_heading("Model Training with LSTM", 1)
doc.add_paragraph(f"Model: LSTM\nBatch Size: 1024\nEpochs: 50\nSplits: 5\nFeatures Used: {list(features)}")

tscv = TimeSeriesSplit(n_splits=5)
batch_size = 1024
scores = {"mse": [], "mae": [], "r2": []}

for fold, (train_idx, test_idx) in enumerate(tscv.split(X_full)):
    print(f"Step 8.{fold+1}: Training Fold {fold+1}...")
    X_train, X_test = X_full[train_idx], X_full[test_idx]
    y_train, y_test = y_full[train_idx], y_full[test_idx]

    # Create tf.data.Dataset for batching
    train_dataset = tf.data.Dataset.from_tensor_slices((X_train, y_train)).batch(batch_size)
    test_dataset = tf.data.Dataset.from_tensor_slices((X_test, y_test)).batch(batch_size)

    # Clear Keras backend
    tf.keras.backend.clear_session()

    model = Sequential([
        LSTM(128, input_shape=(seq_length, len(features)), return_sequences=True),
        Dropout(0.2),
        LSTM(64),
        Dropout(0.2),
        Dense(32, activation='relu'),
        Dense(pred_length)
    ])
    model.compile(optimizer='adam', loss='mse')

    # Fit model
    history = model.fit(train_dataset, epochs=50, validation_data=test_dataset, verbose=1)

    y_pred = model.predict(test_dataset)
    y_test_inv = scaler_target.inverse_transform(y_test)
    y_pred_inv = scaler_target.inverse_transform(y_pred)

    mse = mean_squared_error(y_test_inv, y_pred_inv)
    mae = mean_absolute_error(y_test_inv, y_pred_inv)
    r2 = r2_score(y_test_inv.flatten(), y_pred_inv.flatten())
    scores["mse"].append(mse)
    scores["mae"].append(mae)
    scores["r2"].append(r2)

    # Loss plot
    plt.figure(figsize=(10, 6))
    plt.plot(history.history['loss'], label='Training Loss')
    plt.plot(history.history['val_loss'], label='Validation Loss')
    plt.title(f"Fold {fold+1} Loss")
    plt.legend()
    plt.savefig(f"fold_{fold+1}_loss.png", dpi=300)
    plt.close()

    # Actual vs Predicted (first 30 days or less if data is small)
    plt.figure(figsize=(12, 6))
    plt.plot(y_test_inv.flatten()[:720], label='Actual')
    plt.plot(y_pred_inv.flatten()[:720], label='Predicted')
    plt.title(f"Fold {fold+1} Actual vs Predicted (First 30 Days)")
    plt.legend()
    plt.savefig(f"fold_{fold+1}_pred.png", dpi=300)
    plt.close()

    doc.add_heading(f"Fold {fold+1}", 2)
    doc.add_paragraph(f"MSE: {mse:.4f}\nMAE: {mae:.4f}\nR²: {r2:.4f}")
    add_figure(doc, f"fold_{fold+1}_loss.png", f"Fold {fold+1} Loss")
    add_figure(doc, f"fold_{fold+1}_pred.png", f"Fold {fold+1} Predictions")

# Average metrics
avg_metrics = pd.DataFrame([
    {"Metric": "MSE", "Value": np.mean(scores["mse"])},
    {"Metric": "MAE", "Value": np.mean(scores["mae"])},
    {"Metric": "R²", "Value": np.mean(scores["r2"])}
])
add_table(doc, avg_metrics, "Average Metrics Across Folds")

# %% [code]
# XGBoost-based SHAP Tree Importance (Post-Training Analysis)
print("Step 9: Computing SHAP Tree Importance with XGBoost...")
doc.add_heading("Feature Importance with SHAP Tree via XGBoost (Post-Training)", 1)

# Train XGBoost model for SHAP analysis
xgb_model = xgboost.XGBRegressor(
    objective='reg:squarederror',
    tree_method='hist',  # Updated per XGBoost 2.0+ recommendation
    device='cuda',       # Use CUDA for GPU acceleration
    random_state=seed_value
)
xgb_model.fit(X_full_flat, y_full_flat)

# Use Tree Explainer (fallback due to GPUTree requiring CUDA extension)
X_sample_flat = X_full_flat[-100:]  # Sample for SHAP
explainer = shap.TreeExplainer(xgb_model)  # Use TreeExplainer instead of GPUTreeExplainer
shap_values = explainer.shap_values(X_sample_flat)

# Calculate mean absolute SHAP values for feature importance
shap_importance = np.abs(shap_values).mean(axis=0)
feature_importance = pd.DataFrame({
    'Feature': features,
    'Importance': shap_importance
}).sort_values('Importance', ascending=False)

# Save SHAP summary plot
shap.summary_plot(shap_values, X_sample_flat, feature_names=features, show=False)
plt.savefig("shap_tree.png", dpi=300, bbox_inches='tight')
plt.close()
add_figure(doc, "shap_tree.png", "SHAP Tree Summary Plot (XGBoost)")

# Add feature importance table
add_table(doc, feature_importance, "Feature Importance Ranking (SHAP Tree via XGBoost)")

# %% [code]
# Final Prediction with LSTM
print("Step 10: Generating Final Prediction with LSTM...")
doc.add_heading("Final Prediction with LSTM", 1)
last_data = X_full[-1:]
pred_next_24 = model.predict(last_data)
pred_next_24_inv = scaler_target.inverse_transform(pred_next_24)
future_timestamps = pd.date_range(start=df.index[-1], periods=25, freq='H')[1:]

plt.figure(figsize=(12, 6))
plt.plot(df.index[-24:], scaler_target.inverse_transform(df['SO2'][-24:].values.reshape(-1, 1)), label='Actual')
plt.plot(future_timestamps, pred_next_24_inv[0], label='Predicted', linestyle='--')
plt.title("Next 24-Hour SO2 Prediction (LSTM)")
plt.legend()
plt.savefig("final_pred.png", dpi=300)
plt.close()
add_figure(doc, "final_pred.png", "Next 24-Hour Prediction (LSTM)")

# Save report
print("Step 11: Saving Report...")
doc.save("report_lampang.docx")
print("Step 12: Process Completed!")