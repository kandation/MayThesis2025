import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os
import tensorflow as tf
import random
from config import Config

def setup_environment():
    """ตั้งค่าสภาพแวดล้อมเริ่มต้น"""
    np.random.seed(Config.SEED_VALUE)
    tf.random.set_seed(Config.SEED_VALUE)
    random.seed(Config.SEED_VALUE)
    
    gpus = tf.config.experimental.list_physical_devices('GPU')
    if gpus:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print("GPU is set up and ready!")
    print(f"Num GPUs Available: {len(gpus)}")

def load_data(file_path):
    """โหลดข้อมูลจากไฟล์ CSV"""
    df = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
    return df.copy()

def save_plot(fig, filename, title):
    """บันทึกกราฟ"""
    fig.savefig(Config.OUTPUT_DIR / f"{filename}.png", dpi=300, bbox_inches='tight')
    plt.close()

def add_doc_table(doc, df, title, desc=""):
    """เพิ่มตารางในเอกสาร Word"""
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for i, row in df.iterrows():
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = str(val)
    if desc:
        doc.add_paragraph(desc)

def add_doc_figure(doc, filename, title, desc=""):
    """เพิ่มรูปภาพในเอกสาร Word"""
    full_path = Config.OUTPUT_DIR / f"{filename}.png"
    if full_path.exists():
        doc.add_heading(title, level=2)
        doc.add_picture(str(full_path), width=Inches(6))
        if desc:
            doc.add_paragraph(desc)
    else:
        doc.add_paragraph(f"Warning: Image file {full_path} not found.")