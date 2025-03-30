# %%
# Get file from directory and print it

import os
import seaborn as sns
import docx
from docx import Document
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.ensemble import RandomForestRegressor
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
import matplotlib.pyplot as plt
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.metrics import mean_squared_error, roc_curve, auc

from sklearn.preprocessing import MinMaxScaler

from keras.preprocessing.sequence import TimeseriesGenerator
import matplotlib.pyplot as plt

import random

import tensorflow as tf

root_dir = r"D:/Users/User/PycharmProjects/may-project/labs/outputs_tmd/merged"
root_dir = r"D:/Users/User/PycharmProjects/may-project/labs/outputs/base/merged"
root_dir = r"E:\MayThesis2025\cleanned_datasets"

# get all csv files in the directory

files_name = [f for f in os.listdir(root_dir) if f.endswith(".csv")]

uris = [os.path.join(root_dir, f) for f in files_name]

index_file = 4
# 3

FEATURE_SET = 0

START_DATASET_DATE = "2019-01-01"
END_DATASET_DATE = "2019-12-31"

target = "ValueLn_SO2"


uri = uris[index_file]
current_file_name = files_name[index_file]

uris
files_name
print(files_name)
current_file_name
len(uris)

# %%
# Load and preprocess the data
data = pd.read_csv(uri, parse_dates=["Datetime"])

# %%
columns = data.columns.str.strip().str.lower().tolist()
IS_2009_DATA = False


if "SO2".lower() in columns:
    print(data.columns)
    IS_2009_DATA = True

    print("yes")


elif "ValueLn_SO2_2009".lower() in columns:
    # change ValueLn_SO2_2009 to ValueLn_SO2

    data.rename(columns={"ValueLn_SO2_2009": "ValueLn_SO2"}, inplace=True)
    data.rename(columns={"Value_SO2_2009": "Value_SO2"}, inplace=True)

    print("yes 2009")


else:
    print(data.columns)

    raise Exception("The dataset does not contain the target column")

data

# %% [markdown]
# # Summerize this notebook
#
#
#

# %%
# Set the random seed for reproducibility
seed_value = 42
random.seed(seed_value)
np.random.seed(seed_value)
tf.random.set_seed(seed_value)


# create docs


# sns

plt.style.use("seaborn-paper")

doc = Document()


# Set GPU memory growth
physical_devices = tf.config.list_physical_devices("GPU")
if len(physical_devices) > 0:
    tf.config.experimental.set_memory_growth(physical_devices[0], True)


data.reset_index()
data.set_index("Datetime", inplace=True)
data.sort_index(inplace=True)


# doc info
doc.add_heading("LSTM", 0)
# filename
doc.add_heading("Filename", level=1)
doc.add_paragraph(f"Filename: {uri}")

doc.add_heading("Data", level=1)
doc.add_paragraph(f"Data: {uri}")
doc.add_paragraph(f"Shape: {data.shape}")
doc.add_paragraph(f"Columns: {data.columns}")
doc.add_paragraph(f"Head: {data.head()}")
doc.add_paragraph(f"Tail: {data.tail()}")
doc.add_paragraph(f"Info: {data.info()}")


doc.add_heading("Date", level=2)
doc.add_paragraph(f"Start: {data.index.min()}")
doc.add_paragraph(f"End: {data.index.max()}")

# %%
data.index

# %%
data.columns

# %%
# Get columns containing 'Value_'

value_columns = [col for col in data.columns if "Value_" in col]
value_columns

data_raw = data[value_columns]

data_raw.describe()

# %%
data_raw.plot(subplots=True, figsize=(20, 10), lw=0.5)

# %%
doc.add_heading("Data Preprocessing", level=1)
doc.add_heading("Data Info", level=2)
doc.add_paragraph(f"The data has {data.shape[0]} rows and {data.shape[1]} columns")
doc.add_paragraph(f"The data has {data.isna().sum().sum()} missing values")
doc.add_paragraph(f"The data has {data.duplicated().sum()} duplicated values")


# add table for showing the first 5 rows of the data
doc.add_heading("Data", level=2)

tb_rows_num = 5
tb_cols_num = data.shape[1]


table = doc.add_table(rows=tb_rows_num + 1, cols=tb_cols_num, style="Table Grid")
table.style.border_width = 1

for i in range(tb_cols_num):
    table.cell(0, i).text = data.columns[i]

for i in range(tb_rows_num):
    print(i)
    for j in range(tb_cols_num):
        table.cell(i + 1, j).text = str(data.iloc[i, j])


doc.save("output.docx")


print(data.head())


# %%
def set_plot():
    # Set the legend as the title for each subplot

    for ax in plt.gcf().axes:
        handles, labels = ax.get_legend_handles_labels()
        if labels:
            ax.set_title(labels[0], fontsize=10)
            ax.legend_.remove()
        ax.set_xlabel("Datetime", fontsize=10)
        ax.xaxis.set_tick_params(labelsize=10, rotation=0)

        # set

        # add border to each subplot
        for spine in ax.spines.values():
            spine.set_edgecolor("black")
            spine.set_linewidth(0.5)

    plt.tight_layout()
    plt.savefig(
        "temp_fig.png",
        dpi=300,
        bbox_inches="tight",
        facecolor="white",
        transparent=False,
    )
    plt.show()


# %%
print(data.describe().to_markdown())

# add table
doc.add_heading("Data Description", level=2)
doc.add_paragraph("Descriptive statistics of the data:")

# describe
doc.add_paragraph(f"Data Description: {data.describe()}")


table = doc.add_table(rows=1, cols=9)
table.style = "Table Grid"
table.style.border_width = 1
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Column"
hdr_cells[1].text = "Count"
hdr_cells[2].text = "Mean"
hdr_cells[3].text = "Std"
hdr_cells[4].text = "Min"
hdr_cells[5].text = "25%"
hdr_cells[6].text = "50%"
hdr_cells[7].text = "75%"
hdr_cells[8].text = "Max"


for col in data.columns:
    print(col)
    row_cells = table.add_row().cells
    row_cells[0].text = col
    row_cells[1].text = str(data[col].count())
    # show mean if it is number, otherwise show  the value
    row_cells[2].text = (
        f"{data[col].mean():.4f}"
        if isinstance(data[col].mean(), (int, float))
        else data[col].mean()
    )
    row_cells[3].text = (
        f"{data[col].std():.4f}"
        if isinstance(data[col].std(), (int, float))
        else data[col].std()
    )
    row_cells[4].text = (
        f"{data[col].min():.4f}"
        if isinstance(data[col].min(), (int, float))
        else data[col].min()
    )
    row_cells[5].text = (
        f"{data[col].quantile(0.25):.4f}"
        if isinstance(data[col].quantile(0.25), (int, float))
        else data[col].quantile(0.25)
    )
    row_cells[6].text = (
        f"{data[col].quantile(0.50):.4f}"
        if isinstance(data[col].quantile(0.50), (int, float))
        else data[col].quantile(0.50)
    )
    row_cells[7].text = (
        f"{data[col].quantile(0.75):.4f}"
        if isinstance(data[col].quantile(0.75), (int, float))
        else data[col].quantile(0.75)
    )
    row_cells[8].text = (
        f"{data[col].max():.4f}"
        if isinstance(data[col].max(), (int, float))
        else data[col].max()
    )

# row_cells[4].text = f"{data[col].min():.4f}"
# row_cells[5].text = f"{data[col].quantile(0.25):.4f}"
# row_cells[6].text = f"{data[col].quantile(0.50):.4f}"
# row_cells[7].text = f"{data[col].quantile(0.75):.4f}"
# row_cells[8].text = f"{data[col].max():.4f}"


doc.save("output.docx")

# %%
# Get column startwith ValueLn_
value_columns = [col for col in data.columns if col.startswith("ValueLn_")]
value_columns

# %%
print(data[value_columns].isna().sum())
# add doc as table
doc.add_heading("Missing Values", level=2)
# add table
table = doc.add_table(rows=1, cols=2, style="Table Grid")
table.style.border_width = 1
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Column"
hdr_cells[1].text = "Missing Values"

for col in data[value_columns].columns:
    row_cells = table.add_row().cells
    row_cells[0].text = col
    row_cells[1].text = str(data[col].isna().sum())

doc.add_heading("Missing Values", level=2)
doc.add_paragraph(f"Missing values in each column:\n{data[value_columns].isna().sum()}")
doc.add_paragraph(f"Total missing values: {data[value_columns].isna().sum().sum()}")
doc.add_paragraph(
    f"Percentage of missing values: {data[value_columns].isna().sum().sum() / data.size * 100:.2f}%"
)
doc.add_paragraph(
    f"Columns with missing values: {data[value_columns].columns[data[value_columns].isna().any()].tolist()}"
)
doc.add_paragraph(
    f"Number of rows with missing values: {data[value_columns].isna().any(axis=1).sum()}"
)
doc.add_paragraph(
    f"Percentage of rows with missing values: {data[value_columns].isna().any(axis=1).sum() / data.shape[0] * 100:.2f}%"
)
doc.add_paragraph("Filling missing values with the mean of each column")
doc.save("output.docx")

# %%
columns_before_rm_na = data.columns

# %%
# Define the pool of potential features and target variable
feature_pool_set = [
    [
        "ValueLn_BP",
        "ValueLn_DEV",
        "ValueLn_NO2",
        "ValueLn_PM10",
        "ValueLn_RAIN",
        "ValueLn_RH",
        "ValueLn_TEMP",
        "ValueLn_TSP",
        "ValueLn_WD",
        "ValueLn_WS",
    ],
    #  for filw relv
    [
        "ValueLn_CO",
        "ValueLn_NO",
        "ValueLn_O3",
        "ValueLn_H2S",
        "ValueLn_NO2",
        "ValueLn_NOX",
        "ValueLn_Glob_rad",
        "ValueLn_Net_rad",
        "ValueLn_PM10",
        "ValueLn_Pressure",
        "ValueLn_Rain",
        "ValueLn_Rel_hum",
        "ValueLn_Temp",
        "ValueLn_Wind_dir",
        "ValueLn_Wind_speed",
        "ValueLn_PM2.5",
    ],
]

feature_pool = feature_pool_set[FEATURE_SET]


def select_available_features(data, feature_pool):
    available_features = [
        feature for feature in feature_pool if feature in data.columns
    ]
    return available_features


# %%
# find skipping datetime


def find_missing_key(data, key="Datetime"):
    # generate datetime every 1 hor since 2010-01-01 to  2019-12-31

    date_start_sim = pd.date_range(
        start=data.index.min(), end=data.index.max(), freq="1H"
    )
    print(len(date_start_sim))
    print(len(data.index))
    # find missing datetime

    missing_date = date_start_sim.difference(data.index)
    print(len(missing_date))

    return missing_date


missing_date = find_missing_key(data)

# %%
# Merge missing date to data
missing_data = pd.DataFrame(index=missing_date)
data = pd.concat([data, missing_data], axis=0)

data.sort_index(inplace=True)

# %% [markdown]
# ## Check column after drop

# %%
columns_after_rm_na = data.columns
columns_compare_diff = columns_after_rm_na.difference(columns_before_rm_na)


print("columns_before_rm_na", len(columns_before_rm_na), columns_before_rm_na.tolist())


print("columns_after_rm_na", len(columns_after_rm_na), columns_after_rm_na.tolist())
print(
    "columns_after_rm_na.difference(columns_before_rm_na)",
    len(columns_compare_diff),
    columns_compare_diff.tolist(),
)
# add doc


doc.add_heading("Remove Missing Values", level=2)


text_with_join = "\n".join(columns_before_rm_na)


doc.add_paragraph(f"Columns before removing missing values: {text_with_join}")


doc.add_paragraph(f"Columns after removing missing values: {columns_after_rm_na}")


doc.add_paragraph(f"Columns removed: {columns_compare_diff}")


doc.save("output.docx")

# %%
feature_pool

[feature for feature in data.columns if feature in feature_pool]

# %% [markdown]
# ## Feature remove non avalible

# %%
# split Datetime (index) start from year 2010-01-01
if IS_2009_DATA:
    data = data.loc["2010-01-01":]


data_raw = data.copy()


available_features = [feature for feature in feature_pool if feature in data.columns]


print("available_features", available_features)


# Select available features in feature_pool


features_available = select_available_features(data, feature_pool)


print(f"Available features: {features_available}")


# make data for feature available


data = data[[target] + features_available]


doc.add_heading("Feature Selection", level=2)


doc.add_paragraph(f"Available features: {features_available}")

# %%
# Drop rows with missing values in the selected features
# data = data.dropna(subset=features_available + [target])

# %%
# Drop feature if it has missing values more than 50%
# drop Null values of target since begin of null to end of null (Select only exist data from target )
data = data.dropna(subset=[target])
data.plot(subplots=True, figsize=(20, 10), lw=0.5)
missing_values = data.isna().sum() / data.shape[0]
missing_values = missing_values[missing_values > 0.5]
print(f"Features with more than 50% missing values: {missing_values}")

# Drop features with more than 50% missing values
data = data.drop(columns=missing_values.index)

# Fill missing values with the mean of each column
data = data.fillna(data.mean())

# Check if there are still missing values
missing_values = data.isna().sum()
print(f"Missing values: {missing_values}")


data.plot(subplots=True, figsize=(20, 10), lw=0.5)


doc.add_heading("Missing Values", level=2)
doc.add_paragraph(f"Missing values in each column:\n{data.isna().sum()}")
doc.add_paragraph(f"Total missing values: {data.isna().sum().sum()}")
doc.add_paragraph(
    f"Percentage of missing values: {data.isna().sum().sum() / data.size * 100:.2f}%"
)
doc.add_paragraph(
    f"Columns with missing values: {data.columns[data.isna().any()].tolist()}"
)
doc.add_paragraph(
    f"Number of rows with missing values: {data.isna().any(axis=1).sum()}"
)
doc.add_paragraph(
    f"Percentage of rows with missing values: {data.isna().any(axis=1).sum() / data.shape[0] * 100:.2f}%"
)
doc.add_paragraph("Filling missing values with the mean of each column")


doc.save("output.docx")

# %% [markdown]
# ## Recalculate Avalible Feature

# %%
features_available = select_available_features(data, feature_pool)
print(f"Available features: {features_available}")

# %% [markdown]
# ## Plot Outliner

# %%
import docx


def plot_outliner(data, feature):
    # calculate best scale of width of figure
    best_fig_size = 2 + 0.7 * len(feature)

    # Plot data to bloxplot by subplots split by features'
    fig, axes = plt.subplots(1, len(feature), figsize=(best_fig_size, 15))

    # Reduce the size of the markers using 'flierprops'
    flierprops = dict(
        marker="o", markersize=5, linewidth=1, markeredgecolor="black", alpha=0.5
    )

    for i, col in enumerate(features_available):
        data.boxplot(col, ax=axes[i], fontsize=5, grid=False, flierprops=flierprops)

        # remove ValueLn_ from the title

        axes[i].set_title(col[8:], fontsize=5)

        # box plot circle line width = 0.5

        for line in axes[i].lines:
            line.set_linewidth(1)

        # asix y value size = 10

        axes[i].yaxis.set_tick_params(labelsize=6)

        # add space between subplots

        plt.subplots_adjust(wspace=1)

        axes[i].xaxis.set_tick_params(labelsize=6, rotation=0)

        # add border to each subplot

        for spine in axes[i].spines.values():
            spine.set_edgecolor("black")

            spine.set_linewidth(0.5)

    # save png no transparent and white background

    plt.savefig(
        "temp_fig.png",
        dpi=300,
        bbox_inches="tight",
        facecolor="white",
        transparent=False,
    )

    plt.show()


plot_outliner(data, features_available)


def save_doc_fig(header, text):
    doc.add_heading(header, level=2)
    doc.add_paragraph(text)

    doc.add_picture("temp_fig.png")


save_doc_fig("Outliers", "Outliers in the data (original):")

# %%
# length index
length_index = len(data.index)
length_index

# %%
# select colunm startwith ValueLn_
features_to_normalize = [col for col in data.columns if col.startswith("ValueLn_")]


print(features_to_normalize)

# remove target from features_to_normalize
if target in features_to_normalize:
    features_to_normalize.remove(target)

best_hight_fig_size = 2 + (1.4 * len(features_to_normalize))

data[[target] + features_to_normalize].plot(
    subplots=True,
    figsize=(13, best_hight_fig_size),
    lw=0.1,
    fontsize=10,
)

set_plot()


save_doc_fig("Data Distribution", "Data distribution before normalization:")


# %%
# remove outliers using IQD
def remove_outliers_iqr(data, threshold=1.5):
    Q1 = data.quantile(0.5)  # 0.25
    Q3 = data.quantile(0.95)  # 0.75
    IQR = Q3 - Q1
    return data[
        ~((data < (Q1 - threshold * IQR)) | (data > (Q3 + threshold * IQR))).any(axis=1)
    ]


# Normalize outliers using winsorization
def winsorize(data, lower_percentile=5, upper_percentile=95):
    lower_threshold = np.percentile(data, lower_percentile)
    upper_threshold = np.percentile(data, upper_percentile)
    return np.clip(data, lower_threshold, upper_threshold)


# %%
features_to_normalize

# %%
from enum import Enum


class OutlierDetech(Enum):
    IQR = 1
    Winsorization = 2


use_outlier_detection = OutlierDetech.IQR

if use_outlier_detection == OutlierDetech.IQR:
    data_cleaned = remove_outliers_iqr(data[features_to_normalize])
elif use_outlier_detection == OutlierDetech.Winsorization:
    data_cleaned = data[features_to_normalize].apply(winsorize)


doc.add_heading("Remove Outliers", level=2)
doc.add_paragraph("Remove outliers using {}:".format(use_outlier_detection.name))
doc.add_paragraph("Number of rows before removing outliers: {}".format(length_index))
doc.add_paragraph(
    "Number of rows after removing outliers: {}".format(len(data_cleaned))
)
doc.save("output.docx")

# %%
# # Step 3: Remove missing values if they span more than 1 day
# data_remove_missing = data.copy()


# max_missing_days = pd.Timedelta(days=1)

# data_remove_missing["missing_flag"] = data_remove_missing[[
#     target]].isna().any(axis=1)

# data_remove_missing["gap_duration"] = data_remove_missing["missing_flag"].groupby(
#     (
#         data_remove_missing["missing_flag"]
#         != data_remove_missing["missing_flag"].shift()
#     ).cumsum()
# ).transform("size") * pd.Timedelta(days=1)

# data_remove_missing = data_remove_missing[
#     data_remove_missing["gap_duration"] <= max_missing_days
# ].drop(columns=["missing_flag", "gap_duration"])


# data_remove_missing

# %%
# # plot data and move legend to upper top


# f_to_plot = features_to_normalize

# n_features = len(f_to_plot)
# best_hight_fig_size = 10 + (1.4 * n_features)

# fig, axs = plt.subplots(n_features, 1, sharex=True, figsize=(20, best_hight_fig_size))

# # Plot the data on each subplot
# for i, col in enumerate(f_to_plot):
#     data_cleaned[col].plot(ax=axs[i], lw=0.5, fontsize=10)
#     axs[i].set_title(col[8:], fontsize=10)
#     axs[i].yaxis.set_tick_params(labelsize=10)
#     # legend to the upper top

#     axs[i].legend(loc="upper right", ncol=3, fancybox=True, shadow=True)


# # Show the plot
# plt.show()

# # plot data and move legend to upper top

# f_to_plot = features_to_normalize

# n_features = len(f_to_plot)
# best_hight_fig_size = 10 + (1.4 * n_features)

# fig, axs = plt.subplots(n_features, 1, sharex=True, figsize=(20, best_hight_fig_size))

# # Plot the data on each subplot
# for i, col in enumerate(f_to_plot):
#     data[col].plot(ax=axs[i], lw=0.5, fontsize=10)
#     axs[i].set_title(col[8:], fontsize=10)
#     axs[i].yaxis.set_tick_params(labelsize=10)
#     # legend to the upper top

#     axs[i].legend(loc="upper right", ncol=3, fancybox=True, shadow=True)


# # Show the plot
# plt.show()

# %%
print(len(data_cleaned.index))


print(len(data.index))

# %%
# # mereg data_cleaned with target
# data_cleaned = pd.concat([data[target], data_cleaned], axis=1)
# data_cleaned.sort_index(inplace=True)
# # plot since jan 2014 to jan 2015 from datetime
# # KeyError: 'Value based partial slicing on non-monotonic DatetimeIndexes with non-existing keys is not allowed.'
# data_cleaned.loc["2014-01-01":"2015-01-01"]


# # find not unique index
# len(data_cleaned.index[data_cleaned.index.duplicated()])


# Example concat

df1 = pd.DataFrame(
    {
        "A": ["A0", "A1", "A2", "A3"],
        "B": ["B0", "B1", "B2", "B3"],
        "C": ["C0", "C1", "C2", "C3"],
        "D": ["D0", "D1", "D2", "D3"],
    },
    index=[0, 1, 2, 3],
)

df2 = pd.DataFrame(
    {
        "A": ["A4", "A5", "A6", "A7"],
        "B": ["B4", "B5", "B6", "B7"],
        "C": ["C4", "C5", "C6", "C7"],
        "D": ["D4", "D5", "D6", "D7"],
    },
    index=[3, 4, 5, 6],
)

pd.concat([df1, df2], axis=1)

# %%
print("Length of data_cleaned: ", len(data_cleaned.index))
print("Length of data: ", len(data.index))


# find duplicated index
print(len(data_cleaned.index[data_cleaned.index.duplicated()]))
display(data_cleaned.index[data_cleaned.index.duplicated()])
print(len(data.index[data.index.duplicated()]))
display(data.index[data.index.duplicated()])

# drop duplicated index
data_cleaned = data_cleaned[~data_cleaned.index.duplicated(keep="first")]
data = data[~data.index.duplicated(keep="first")]


rs_data_cleaned = data_cleaned.copy()
# rs_data_cleaned.reset_index(drop=True, inplace=True)

rs_data = data.copy()
# rs_data.reset_index(drop=True, inplace=True)

rs_concat = pd.concat([rs_data, rs_data_cleaned], axis=1)
print(rs_concat.columns)

# find duplicated index
print(len(rs_concat.index[rs_concat.index.duplicated()]))

display(rs_concat)

# %%
# find not duplicated index from data and data_cleaned
not_duplicated_index = data.index.difference(data_cleaned.index)
display(not_duplicated_index)


# find not  duplicated column from data and data_cleaned
not_duplicated_columns = data.columns.difference(data_cleaned.columns)
not_duplicated_columns

# %% [markdown]
# # Correlation Matrix

# %%
# add to docx
doc.add_heading("Correlation Matrix", level=2)
doc.add_paragraph("Correlation matrix before and after cleaning the data")
font_size = 7
# mereg data_cleaned with target
data_cleaned = pd.concat([data[target], data_cleaned], axis=1)
plt.figure(figsize=(10, 10))
res = sns.heatmap(
    data[[target] + features_to_normalize].corr(),
    annot=True,
    cmap="coolwarm",
    fmt=".2f",
    vmin=-1,
    vmax=1,
    annot_kws={"size": font_size},
)
# # Drawing the frame
# res.axhline(y=0, color="k", linewidth=5)
# res.axhline(y=data_cleaned.shape[1], color="k", linewidth=5)
# res.axvline(x=0, color="k", linewidth=5)
# res.axvline(x=data_cleaned.shape[0], color="k", linewidth=5)
plt.title("Correlation Matrix Before Cleaning", fontsize=font_size)
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
save_doc_fig("Correlation Matrix", "Correlation matrix before cleaning the data")
plt.show()
# plot correlation matrix data_cleaned
plt.figure(figsize=(10, 10))
res = sns.heatmap(
    data_cleaned.corr(),
    annot=True,
    cmap="coolwarm",
    fmt=".2f",
    vmin=-1,
    vmax=1,
    annot_kws={"size": font_size},
)
# # Drawing the frame
# res.axhline(y=0, color="k", linewidth=5)
# res.axhline(y=data_cleaned.shape[1], color="k", linewidth=5)
# res.axvline(x=0, color="k", linewidth=5)
# res.axvline(x=data_cleaned.shape[0], color="k", linewidth=5)
plt.title("Correlation Matrix After Cleaning", fontsize=font_size)
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
save_doc_fig("Correlation Matrix", "Correlation matrix after cleaning the data")
plt.show()

# %%
# Sort correlation matrix
correlation_matrix = data_cleaned.corr()

# Create a DataFrame with feature, correlation, and absolute correlation columns
correlation_df = pd.DataFrame(
    {
        "feature": correlation_matrix.index,
        "correlation": correlation_matrix[target],
        "abs_correlation": correlation_matrix[target].abs(),
    }
)

# Sort the DataFrame by absolute correlation in descending order
correlation_df.sort_values(by="abs_correlation", ascending=False, inplace=True)

# Reset the index
correlation_df.reset_index(drop=True, inplace=True)


display(correlation_df)


# make table

doc.add_heading("Correlation Matrix", level=2)
doc.add_paragraph("Correlation sorted by absolute correlation with the target variable")

# add table
table = doc.add_table(rows=1, cols=3, style="Table Grid")
table.style.border_width = 1
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Feature"
hdr_cells[1].text = "Correlation"
hdr_cells[2].text = "Abs Correlation"

for i in range(correlation_df.shape[0]):
    row_cells = table.add_row().cells
    row_cells[0].text = correlation_df.iloc[i, 0]
    row_cells[1].text = f"{correlation_df.iloc[i, 1]:.4f}"
    row_cells[2].text = f"{correlation_df.iloc[i, 2]:.4f}"

doc.save("output.docx")

# %%
data_cleaned[[target] + features_to_normalize].plot(
    subplots=True,
    figsize=(15, 15),
    lw=0.1,
    fontsize=10,
)


set_plot()


save_doc_fig("Data Distribution", "Data distribution after normalization:")

# %%
# # crate subplot
# features_to_normalize_raw = features_to_normalize.copy()

# %%
# Sculter plot for observer destribute between so2 and other features


# features_to_normalize = features_to_normalize_raw[:3]

print(features_to_normalize)

layout_col = 3
layout_row = len(features_to_normalize) // layout_col + 1

print(layout_row)
cal_fig_sz_width = 6 * 1.2 * layout_col
cal_fig_sz_height = 4 * 1.2 * layout_row


fig, axs = plt.subplots(
    layout_row,
    layout_col,
    figsize=(cal_fig_sz_width, cal_fig_sz_height),
    layout="tight",
)
# make 3 x 2 layout
# fig, axs = plt.subplots(3, 2, figsize=(20, 15), layout="tight")


# flatten axs
axs = axs.flatten()
print("layout_row", layout_row)

for i, col in enumerate(features_to_normalize):
    axs[i].scatter(data_cleaned[col], data_cleaned[target], s=5, alpha=1)
    axs[i].set_title(col[8:], fontsize=10)
    axs[i].set_xlabel(col[8:], fontsize=10)
    axs[i].set_ylabel("SO2", fontsize=10)

    # set font size
    axs[i].xaxis.set_tick_params(labelsize=10)
    axs[i].yaxis.set_tick_params(labelsize=10)


plt.tight_layout()

plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)

# save to docx
doc.add_heading("Scatter Plot", level=2)
doc.add_paragraph("Scatter plot between SO2 and other features")
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# ma
plt.show()

# %% [markdown]
#

# %%
plot_outliner(data_cleaned, features_available)

save_doc_fig("Outliers", "Outliers in the data (cleaned):")

# %% [markdown]
# # Feature Best
#

# %%
data = data_cleaned.copy()

# Handle missing values in features
data[features_available] = data[features_available].fillna(
    data[features_available].mean()
)
# Fill missing SO2 values with the mean
data[target] = data[target].fillna(data[target].mean())


# Select relevant features using feature importance scoring
X = data[features_available]
y = data[target]
print(f"X shape: {X.shape}")
print(f"y shape: {y.shape}")

# %% [markdown]
# ## Select Best Features

# %%
# selecte best features by correlation

# Select the top 10 features with the highest absolute correlation with the target variable
selected_features = correlation_df["feature"].head(10).tolist()
print(f"Top features: {selected_features}")

# print selected features and score
score_corr_table = correlation_df[correlation_df["feature"].isin(selected_features)]
display(score_corr_table)


doc.add_heading("Feature Selection (from correlation)", level=2)
doc.add_paragraph("Selected features based on correlation with the target variable")
doc.add_paragraph(f"Top features: {selected_features}")

# add table
table = doc.add_table(rows=1, cols=3, style="Table Grid")
table.style.border_width = 1
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Feature"
hdr_cells[1].text = "Correlation"
hdr_cells[2].text = "Abs Correlation"

for i in range(score_corr_table.shape[0]):
    row_cells = table.add_row().cells
    row_cells[0].text = score_corr_table.iloc[i, 0]
    row_cells[1].text = f"{score_corr_table.iloc[i, 1]:.4f}"
    row_cells[2].text = f"{score_corr_table.iloc[i, 2]:.4f}"


doc.save("output.docx")


selected_features

# %%
# # select best features using xgboost
# import xgboost as xgb
# from matplotlib import pyplot
# from sklearn.model_selection import train_test_split

# SELECTED_BEST_FEATURES = 5
# features = features_available
# model = xgb.XGBRegressor()
# model.fit(X, y)

# # Scale the input features
# scaler = MinMaxScaler()
# X_scaled = scaler.fit_transform(data[features])
# y = data[target].values

# # Split the data into training and testing sets
# X_train, X_test, y_train, y_test = train_test_split(
#     X_scaled, y, test_size=0.2, random_state=42
# )


# # Perform feature selection using XGBoost
# dtrain = xgb.DMatrix(X_train, label=y_train)
# params = {"objective": "reg:squarederror", "eval_metric": "rmse"}
# model_xgb = xgb.train(params, dtrain, num_boost_round=100)


# # Get feature importances from XGBoost
# importances = model_xgb.get_score(importance_type="gain")
# print(f"Feature importances: {importances}")
# selected_features = sorted(importances, key=importances.get, reverse=True)[
#     :SELECTED_BEST_FEATURES
# ]


# # Transfer the selected features to the original feature names
# selected_features = [features[int(f[1:])] for f in selected_features]
# print(f"Selected features: {selected_features}")


# # Filter the selected features
# X_train_selected = X_train[:, [features.index(f) for f in selected_features]]
# X_test_selected = X_test[:, [features.index(f) for f in selected_features]]
# print(f"Features: {features_available}")
# print(f"Selected features: {features}")
# print(f"X_train shape: {X_train_selected.shape}")
# print(f"X_test shape: {X_test_selected.shape}")
# sorted_importance = sorted(
#     importances.items(), key=lambda x: x[1], reverse=False)


# # Extract the feature names and importance values
# features = [features[int(f[0][1:])] for f in sorted_importance]
# importance_values = [f[1] for f in sorted_importance]
# print(importance_values)


# # Plot the feature importance
# plt.figure(figsize=(8, 3))
# plt.barh(range(len(importance_values)), importance_values, height=0.6)
# plt.yticks(range(len(features)), features)
# plt.ylabel("Features", fontsize=10)
# plt.xlabel("Importance (Gain)", fontdict={"size": 10})
# plt.title("Feature Importance (Gain)", fontdict={"size": 10})
# plt.tight_layout()
# # save figure
# plt.savefig(
#     "temp_fig.png", dpi=300, bbox_inches="tight", facecolor="white", transparent=False
# )
# plt.show()

# # add to doc
# doc.add_heading("Feature Selection", level=2)
# doc.add_paragraph(f"Features: {features_available}")
# doc.add_paragraph(f"Selected features: {features}")
# doc.add_paragraph(f"X_train shape: {X_train_selected.shape}")
# doc.add_paragraph(f"X_test shape: {X_test_selected.shape}")
# # save fdig to doc
# doc.add_picture("temp_fig.png")
# # dict
# info_dict = {
#     "features": features_available,
#     "selected_features": selected_features,
#     "X_train_shape": X_train_selected.shape,
#     "X_test_shape": X_test_selected.shape,
# }
# doc.save("output.docx")

# %%
# Save feature selection to docx

# %%
# plot preview all features data as line graph

# except for the target variable
dont_plot = ["ValueLn_WD"]

plt.figure(figsize=(25, 6))
for feature in feature_pool + [target]:
    if feature in data.columns and feature not in dont_plot:
        plt.plot(
            data_raw.index, data_raw[feature], label=feature, alpha=0.7, linewidth=0.3
        )
    # if feature in data.columns:
    #     plt.plot(data.index, data[feature], label=feature)
plt.xlabel("Datetime")
plt.ylabel("Value")
plt.title("Features")
plt.legend()
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()


doc.add_heading("Features", level=2)
doc.add_paragraph("Preview of all features data:")
doc.add_picture("temp_fig.png")

# %%
data[selected_features].plot(
    subplots=True,
    figsize=(20, 10),
    lw=0.4,
    fontsize=10,
)


set_plot()


save_doc_fig("Selected Features", "Selected features data:")

# %%
# # Drop rows with missing values in the selected features
# data_cleaned = data_cleaned.dropna(subset=selected_features + [target])


# plot preview all features data as line graph
data_cleaned[[target] + selected_features].plot(
    subplots=True,
    figsize=(20, 10),
    lw=0.4,
    fontsize=10,
)
set_plot()


save_doc_fig("Selected Features", "Selected features data after cleaning:")


data = data_cleaned.copy()

# %%
# data_cleaned[data_cleaned["ValueLn_WS"] < ].plot()

# %%
# head
print(data.head().to_markdown())


X_selected = data[selected_features]
# x_selected_scaled = scaler.fit_transform(X_selected)
y = data[target].values


# X_selected
print(f"X_selected shape: {X_selected.shape}")


print(X_selected.head().to_markdown())
# x_selected_scaled
# print(f'x_selected_scaled shape: {x_selected_scaled.shape}')
# y
print(f"y shape: {y.shape}")
print(y[:5])


# Scale the input features
scaler = MinMaxScaler()
X_scaled = scaler.fit_transform(X_selected)
y = data[target].values


X_scaled


print()
print(f"X_scaled shape: {X_scaled.shape}")


# head
print(X_scaled)


# add to doc
doc.add_heading("Data Preprocessing", level=1)
doc.add_heading("Data Info", level=2)
doc.add_paragraph(f"The data has {data.shape[0]} rows and {data.shape[1]} columns")
doc.add_paragraph(f"The data has {data.isna().sum().sum()} missing values")
doc.add_paragraph(f"The data has {data.duplicated().sum()} duplicated values")


# add table for showing the first 5 rows of the data
doc.add_heading("Data", level=2)


tb_rows_num = 5
tb_cols_num = data.shape[1]


table = doc.add_table(rows=tb_rows_num + 1, cols=tb_cols_num, style="Table Grid")
table.style.border_width = 1


for i in range(tb_cols_num):
    table.cell(0, i).text = data.columns[i]


for i in range(tb_rows_num):
    print(i)
    for j in range(tb_cols_num):
        table.cell(i + 1, j).text = str(data.iloc[i, j])


doc.save("output.docx")

# %% [markdown]
# # Trainning

# %%
# GLOBAL SETTING FOR TRAIN

N_SPLITS = 5

# %%
# Example for TimeSeriesSplit


x_example = np.arange(50).reshape(50, 1)
tscv_X = TimeSeriesSplit(n_splits=N_SPLITS)


# Keep size for plot graph
dataset_size = []


for t, v in tscv_X.split(x_example):
    print(f"Train: {t}, Validation: {v}")
    dataset_size.append({"train": len(t), "validation": len(v)})


def plot_tscv(dataset_size):
    # plot hbar sorted by index ascending
    df = pd.DataFrame(dataset_size)
    df = df.sort_index(ascending=False)
    df.plot(kind="barh", stacked=True)

    # show value on the bar and show total value
    for i in range(len(df)):
        total_text = df.iloc[i]["train"] + df.iloc[i]["validation"]
        train_text = df.iloc[i]["train"]
        train_percentage = train_text / total_text * 100
        validation_text = df.iloc[i]["validation"]
        validation_percentage = validation_text / total_text * 100
        plt.text(
            total_text,
            i,
            f"{total_text}",
            ha="left",
            va="center",
            color="black",
            fontsize=7,
        )
        plt.text(
            train_text / 2,
            i,
            f"{train_text} ({train_percentage:.2f}%)",
            ha="center",
            va="center",
            color="white",
            fontsize=7,
        )
        plt.text(
            train_text + validation_text / 2,
            i,
            f"{validation_text} ({validation_percentage:.2f}%)",
            ha="center",
            va="center",
            color="black",
            fontsize=7,
        )

    plt.xlabel("Dataset Size")
    plt.ylabel("Split (Fold)")
    plt.title("Time Series Cross Validation")
    plt.legend(["Train", "Validation"], loc="upper right")
    plt.savefig(
        "temp_fig.png",
        dpi=300,
        bbox_inches="tight",
        facecolor="white",
        transparent=False,
    )
    plt.show()


plot_tscv(dataset_size)

# %%
# null_indices = data["ValueLn_WS"].isnull()
# null_indices


# vvv = "ValueLn_WS"
# data[vvv].plot()


# df = data.copy()


# # Calculate the mean of the values
# mean_value = df[vvv].mean()


# # Define the threshold for outliers (5% of the mean)
# threshold = 0.05 * mean_value


# # Create a boolean mask for outliers based on the threshold and near-zero condition
# outlier_mask = (df[vvv] < threshold) & (df[vvv] > 0)


# # Create a boolean mask for continuous outliers
# continuous_mask = outlier_mask.rolling(window=3, min_periods=1).sum() >= 3
# final_mask = outlier_mask & continuous_mask


# #
# #  Find the start and end indices of each segment satisfying all conditions
# starts = np.where(final_mask & ~final_mask.shift(1))[0]
# ends = np.where(final_mask & ~final_mask.shift(-1))[0] + 1


# # Plot the original graph
# plt.figure(figsize=(10, 5))
# plt.plot(df.index, df["value"], color="green", label="Original")


# # Paint the segments that satisfy all conditions in red
# for start, end in zip(starts, ends):
#     plt.axvspan(start - 0.5, end - 0.5, color="red", alpha=0.3)


# plt.xlabel("Index")
# plt.ylabel("Value")
# plt.title("Graph with Segments Satisfying All Conditions in Red")
# plt.legend()
# plt.tight_layout()
# plt.show()


# # Combine the outlier and continuous masks
# final_mask = outlier_mask & continuous_mask


# # Find the minimum outlier value
# min_outlier = df.loc[final_mask, vvv].min()
# min_outlier_index = df.loc[final_mask, vvv].idxmin()


# # Create a new DataFrame with the minimum outlier value masked in red
# df_masked = df.copy()
# df_masked.loc[min_outlier_index, vvv] = np.nan


# # Plot the original graph
# plt.figure(figsize=(10, 5))
# plt.plot(df.index, df[vvv], color="green", label="Original")


# # Plot the masked graph with the minimum outlier in red
# plt.plot(df_masked.index, df_masked[vvv], color="green", label="Masked")
# plt.scatter(min_outlier_index, min_outlier,
#             color="red", label="Minimum Outlier")


# plt.xlabel("Index")
# plt.ylabel("Value")
# plt.title("Graph with Minimum Outlier Masked in Red")
# plt.legend()
# plt.tight_layout()
# plt.show()


# print("Minimum outlier value:", min_outlier)

# %% [markdown]
# ## Import all train

# %%
import pandas as pd
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from keras.models import Sequential
from keras.layers import LSTM, Dense
from keras.preprocessing.sequence import TimeseriesGenerator
import matplotlib.pyplot as plt
from keras.layers import CuDNNLSTM, Bidirectional

# %% [markdown]
# ## Train Only So2

# %%
tscv = TimeSeriesSplit(n_splits=N_SPLITS)
# Metrics for the testing set
scores = {"test_mse": [], "test_mae": [], "r2": [], "train_mse": [], "train_mae": []}

timesteps = 48
batch_size = 2048
fold_counter = 1
dataset_size = []

key_target_norm = f"{target}_normalize"

# Normalize the SO2 values
scaler = MinMaxScaler(feature_range=(0, 1))

data_for_so2 = data.copy()
data_for_so2[key_target_norm] = scaler.fit_transform(
    data_for_so2[target].values.reshape(-1, 1)
)


# Store loss values for each fold
train_loss_values = []
val_loss_values = []


doc.add_heading("Trainning", level=1)
doc.add_heading("Model", level=2)
doc.add_paragraph("Model: LSTM")
doc.add_paragraph("Optimizer: Adam")
doc.add_paragraph("Loss: Mean Squared Error")
doc.add_paragraph("Batch Size: {}".format(batch_size))
doc.add_paragraph("Timesteps: {}".format(timesteps))
doc.add_paragraph("Number of Splits: {}".format(N_SPLITS))
doc.add_paragraph("Feature Scaling: MinMaxScaler")
doc.add_paragraph("Feature Range: 0 to 1")
doc.add_paragraph("Selected Features: {}".format(selected_features))
doc.add_paragraph("Target Variable: {}".format(target))
doc.add_paragraph("Normalized Target Variable: {}".format(key_target_norm))
doc.add_paragraph("Model Summary:")
doc.save("output.docx")

# %%
fold_counter = 1

doc.add_heading("Model Training (for Use only So2 as feature)", level=2)
doc.add_paragraph("Model: LSTM")
doc.add_paragraph("Cross Validation by Time Series Split")


for train_index, test_index in tscv.split(data_for_so2):
    dataset_size.append({"train": len(train_index), "validation": len(test_index)})
    train_data = data_for_so2.iloc[train_index]
    val_data = data_for_so2.iloc[test_index]

    doc

    # Create TimeseriesGenerator for training and validation data
    train_generator = TimeseriesGenerator(
        train_data[key_target_norm].values,
        train_data[key_target_norm].values,
        length=timesteps,
        batch_size=batch_size,
    )

    val_generator = TimeseriesGenerator(
        val_data[key_target_norm].values,
        val_data[key_target_norm].values,
        length=timesteps,
        batch_size=batch_size,
    )

    # Build and train the LSTM model
    model_for_so2 = Sequential()
    model_for_so2.add(CuDNNLSTM(50, input_shape=(timesteps, 1)))
    model_for_so2.add(Dense(1))
    model_for_so2.compile(loss="mean_squared_error", optimizer="adam")
    history_for_so2 = model_for_so2.fit(
        train_generator,
        validation_data=val_generator,
        epochs=50,
        batch_size=batch_size,
        verbose=1,
    )

    # Store loss values for each fold
    train_loss_values.append(history_for_so2.history["loss"])
    val_loss_values.append(history_for_so2.history["val_loss"])

    # Predict the validation test set
    val_pred = model_for_so2.predict(val_generator)

    # Inverse transform the predicted values
    val_pred = scaler.inverse_transform(val_pred)
    val_actual = scaler.inverse_transform(
        val_data[key_target_norm].values.reshape(-1, 1)
    )

    # Calculate evaluation metrics
    val_actual_subset = val_actual[timesteps:]
    val_pred_subset = val_pred[: len(val_actual_subset)]

    # Check for NaN values
    nan_mask = np.isnan(val_actual_subset) | np.isnan(val_pred_subset)

    # Remove rows with NaN values
    val_actual_subset = val_actual_subset[~nan_mask]
    val_pred_subset = val_pred_subset[~nan_mask]

    mse = mean_squared_error(val_actual_subset, val_pred_subset)
    mae = mean_absolute_error(val_actual_subset, val_pred_subset)
    r2 = r2_score(val_actual_subset, val_pred_subset)

    scores["test_mse"].append(mse)
    scores["test_mae"].append(mae)
    scores["r2"].append(r2)

    fold_counter += 1

    # Plot the loss graph each fold
    plt.figure(figsize=(10, 5))
    plt.plot(history_for_so2.history["loss"], label="Train")
    plt.plot(history_for_so2.history["val_loss"], label="Validation")
    plt.legend(loc="upper right")
    plt.xlabel("Epoch")
    plt.ylabel("Loss")
    plt.title(f"Model Loss - Fold {fold_counter}")

    plt.savefig(
        "temp_fig.png",
        dpi=300,
        bbox_inches="tight",
        facecolor="white",
        transparent=False,
    )

    plt.show()

    doc.add_heading(f"Fold {fold_counter}", level=3)
    doc.add_paragraph(f"MSE: {mse:.4f}")
    doc.add_paragraph(f"MAE: {mae:.4f}")
    doc.add_paragraph(f"R-squared: {r2:.4f}")
    doc.add_picture("temp_fig.png")
    doc.save("output.docx")


# Calculate average metrics across all folds
avg_mse = np.mean(scores["test_mse"])
avg_mae = np.mean(scores["test_mae"])
avg_r2 = np.mean(scores["r2"])


print(f"Average MSE: {avg_mse:.4f}")
print(f"Average MAE: {avg_mae:.4f}")
print(f"Average R-squared: {avg_r2:.4f}")

doc.add_heading("Average Metrics", level=2)
doc.add_paragraph(f"Average MSE: {avg_mse:.4f}")
doc.add_paragraph(f"Average MAE: {avg_mae:.4f}")
doc.add_paragraph(f"Average R-squared: {avg_r2:.4f}")
doc.save("output.docx")


# Plot the loss graph
plt.figure(figsize=(10, 5))
for i in range(N_SPLITS):
    plt.plot(train_loss_values[i], label=f"Train - Fold {i + 1}")
    plt.plot(val_loss_values[i], label=f"Validation - Fold {i + 1}")


plt.legend(loc="upper right")
plt.xlabel("Epoch")
plt.ylabel("Loss")
plt.title("Model Loss")
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()

doc.add_heading("Model Loss", level=2)
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# %%
# Plot Cross Validation Size split

doc.add_heading("Time Series Cross Validation", level=2)
doc.add_paragraph("Time Series Cross Validation split size")

plot_tscv(dataset_size)


doc.add_picture("temp_fig.png")
doc.save("output.docx")

# %%
# Model Test Unseen Data


train_index, test_index = reversed(list(tscv.split(data))).__next__()

train_data = data_for_so2.iloc[train_index]
test_data = data_for_so2.iloc[test_index]

test_generator = TimeseriesGenerator(
    test_data[key_target_norm].values,
    test_data[key_target_norm].values,
    length=timesteps,
    batch_size=batch_size,
)

# Predict the test set
test_pred = model_for_so2.predict(test_generator)

# Inverse transform the predicted values
test_pred = scaler.inverse_transform(test_pred)


doc.add_heading("Model Testing (for only So2)", level=2)
doc.add_paragraph("Model Testing on Unseen Data")
# size of validation
doc.add_paragraph(f"Validation Size: {len(test_data)}")
doc.add_paragraph(f"Date Start: {test_data.index[0]}")
doc.add_paragraph(f"Date End: {test_data.index[-1]}")


# plot the predicted values
plt.figure(figsize=(10, 5))
plt.plot(
    test_data.index[timesteps:],
    test_data[target].values[timesteps:],
    label="Actual",
)
plt.plot(test_data.index[timesteps:], test_pred, label="Predicted")
plt.xlabel("Datetime")
plt.ylabel("Value")
plt.title("Actual vs. Predicted SO2 Values")
plt.xlim(test_data.index[-timesteps * 30], test_data.index[-1])
plt.legend()
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)

plt.show()

doc.add_paragraph("Actual vs. Predicted SO2 Values (zoom)")
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# plot the predicted values
plt.figure(figsize=(10, 5))
plt.plot(
    test_data.index[timesteps:],
    test_data[target].values[timesteps:],
    label="Actual",
)
plt.plot(test_data.index[timesteps:], test_pred, label="Predicted")
plt.xlabel("Datetime")
plt.ylabel("Value")
plt.title("Actual vs. Predicted SO2 Values")
plt.legend()
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()
doc.add_paragraph("Actual vs. Predicted SO2 Values (all)")
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# %%
# Plot Score of MSE Each fold


# Calculate the average scores

avg_scores = {k: np.mean(v) for k, v in scores.items()}

print(f"Average scores: {avg_scores}")

# Plot scores each index as bar graph

for k, v in scores.items():
    plt.figure(figsize=(10, 5))
    plt.bar(
        range(len(v)),
        v,
        label=k,
    )
    plt.xlabel("Index")
    plt.ylabel(k)
    plt.title(f"{k} for each index")
    # save fig\
    plt.savefig(
        "temp_fig.png",
        dpi=300,
        bbox_inches="tight",
        facecolor="white",
        transparent=False,
    )
    save_doc_fig(f"{k} for each index", f"{k} for each index")
    doc.add_paragraph(f"Average {k}: {avg_scores[k]:.4f}")
    plt.show()

# save file

doc.save("output.docx")

# %%
# plot prediction graph


# Predict the next 24 hours
last_input = data_for_so2[key_target_norm].values[-timesteps:]

print(key_target_norm)
print(last_input.shape)
print("Timesteps", timesteps)

print(last_input.shape)
print(last_input.reshape(1, timesteps, 1))
next_24_hours = []
for _ in range(timesteps):
    next_hour = model_for_so2.predict(last_input.reshape(1, timesteps, 1))
    next_24_hours.append(next_hour[0][0])
    last_input = np.append(last_input[1:], next_hour)


# Inverse transform the predicted values
next_24_hours = scaler.inverse_transform(np.array(next_24_hours).reshape(-1, 1))


# Plot the training and testing data
plt.figure(figsize=(10, 6))
plt.plot(data.index, data[target], label="Training")
plt.plot(data.index[-timesteps:], data[target][-timesteps:], label="Testing")
plt.plot(
    pd.date_range(start=data.index[-1], periods=timesteps + 1, freq="H")[1:],
    next_24_hours,
    label="Predicted",
    linestyle="--",
    marker="o",
)
plt.xlabel("Datetime")
plt.ylabel("ValueLn_SO2_2009")
plt.title("Application for predicting SO2 values for the next 24 hours")
plt.legend()
plt.xticks(rotation=45)
plt.grid(True)


plt.xlim(
    data.index[-24],
    pd.date_range(start=data.index[-1], periods=timesteps * 4, freq="H")[24],
)

plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()

doc.add_heading("Model Prediction (Example)", level=2)
doc.add_paragraph("Model Prediction for the next 24 hours")
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# %% [markdown]
# ## Multiple Features

# %%
# Parameters
timesteps = 24
batch_size = 512
EPOCHS = 100


CLEAN_MISSING_METHOD = "drop"


data_for_multi = data.copy()
print("Index", data_for_multi.index.__len__())

selected_features = data_for_multi.columns.tolist()
selected_features.remove(target)

# Select feature only 2^n -1
selected_features = selected_features
print(f"Selected features: {selected_features}")

if CLEAN_MISSING_METHOD == "drop":
    print("Drop missing values")
    # Drop missing
    data_for_multi = data_for_multi.dropna()
    data_for_multi.plot(subplots=True, figsize=(20, 10), lw=0.5)

elif CLEAN_MISSING_METHOD == "interpolate":
    # interplolate missing values
    data_for_multi = data_for_multi.interpolate(method="linear")
    data_for_multi.plot(subplots=True, figsize=(20, 10), lw=0.5)


print("Index", data_for_multi.index.__len__())


plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)


# Save to doc
doc.add_heading("Multi Feature Lstm train", level=2)
doc.add_heading("Data Cleaning ", level=3)
doc.add_paragraph(f"Clean missing values using {CLEAN_MISSING_METHOD}")
doc.add_picture("temp_fig.png")
doc.save("output.docx")


key_target_norm = f"{target}_normalize"

# Scale the input features
scaler = MinMaxScaler(feature_range=(0, 1))
data_scaled = pd.DataFrame(
    scaler.fit_transform(data_for_multi[selected_features]),
    columns=selected_features,
    index=data_for_multi.index,
)

# Normalize the target variable
target_scaler = MinMaxScaler(feature_range=(0, 1))
data_scaled[key_target_norm] = target_scaler.fit_transform(
    data_for_multi[target].values.reshape(-1, 1)
)

# Create TimeSeriesSplit object
tscv = TimeSeriesSplit(n_splits=N_SPLITS)

# Metrics for the testing set
scores = {"test_mse": [], "test_mae": [], "r2": [], "train_mse": [], "train_mae": []}
dataset_size = []

# Store loss values for each fold
train_loss_values = []
val_loss_values = []

fold_counter = 1

# split train and validation

train_index, test_index = tscv.split(data_scaled).__next__()

dataset_size.append({"train": train_index.shape, "validation": test_index.shape})
train_data = data_scaled.iloc[train_index]
test_data = data_scaled.iloc[test_index]

print(
    "train_data[key_target_norm].values.shape: ",
    train_data[key_target_norm].values.shape,
)

print(
    "train_data[selected_features].values.shape: ",
    train_data[selected_features].values.shape,
)

print(
    "train_data[key_target_norm].values.reshape(-1, 1)",
    train_data[key_target_norm].values.reshape(-1, 1).shape,
)

# Create TimeseriesGenerator for training and validation data
train_generator = TimeseriesGenerator(
    train_data[selected_features].values,
    train_data[key_target_norm].values,
    length=timesteps,
    batch_size=batch_size,
)


val_generator = TimeseriesGenerator(
    test_data[selected_features].values,
    test_data[key_target_norm].values.reshape(-1, 1),
    length=timesteps,
    batch_size=batch_size,
)


# Get the shape of the generated data
for i in range(len(train_generator)):
    x, y = train_generator[i]
    print(f"Batch {i + 1}:")
    print(f"X shape: {x.shape}")
    print(f"y shape: {y.shape}")
    print()


# Build and train the LSTM model
model_multi = Sequential()
model_multi.add(CuDNNLSTM(50, input_shape=(timesteps, len(selected_features))))
# model_multi.add(CuDNNLSTM(50, input_shape=(timesteps, 1)))
model_multi.add(Dense(1))


model_multi.compile(
    loss="mean_squared_error",
    optimizer="adam",
    metrics=["mae", "accuracy"],
)

history_multi = model_multi.fit(
    train_generator,
    epochs=EPOCHS,
    verbose=1,
    validation_data=val_generator,
)


doc.add_heading("Model Training (for all features)", level=2)
doc.add_paragraph("Model: LSTM")
doc.add_paragraph(f"Epochs: {EPOCHS}")
doc.add_paragraph(f"Batch Size: {batch_size}")
doc.add_paragraph(f"Timesteps: {timesteps}")
doc.add_paragraph(f"Number of Splits: {N_SPLITS}")
doc.add_paragraph(f"Feature Scaling: MinMaxScaler")
doc.add_paragraph(f"Feature Range: 0 to 1")
doc.add_paragraph(f"Selected Features: {selected_features}")
doc.add_paragraph(f"Target Variable: {target}")
doc.add_paragraph(f"Train size {train_index.shape}")
doc.add_paragraph(f"Validation size {test_index.shape}")
doc.add_paragraph("Model Summary:")
doc.save("output.docx")

# %%
print(history_multi.history)
# plot history (dict) , training and val_los
plt.figure()

plt.plot(history_multi.history["loss"], label="loss")
plt.plot(history_multi.history["val_loss"], label="val_loss")
plt.plot(history_multi.history["mae"], label="mae")
plt.plot(history_multi.history["accuracy"], label="accuracy")
plt.legend()

plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()

doc.add_heading("Model Loss", level=2)
doc.add_paragraph("Model Loss")
doc.add_picture("temp_fig.png")
doc.save("output.docx")

# Predict the validation set
val_pred = model_multi.predict(val_generator)

# Inverse transform the predicted values
val_pred = target_scaler.inverse_transform(val_pred)
val_actual = target_scaler.inverse_transform(
    test_data[key_target_norm].values.reshape(-1, 1)
)

# Calculate evaluation metrics
val_actual_subset = val_actual[timesteps:]
val_pred_subset = val_pred[: len(val_actual_subset)]

print("val_actual_subset", val_actual_subset.shape)
print("val_pred_subset", val_pred_subset.shape)


mse = mean_squared_error(val_actual_subset, val_pred_subset)
mae = mean_absolute_error(val_actual_subset, val_pred_subset)
r2 = r2_score(val_actual_subset, val_pred_subset)

scores["test_mse"].append(mse)
scores["test_mae"].append(mae)
scores["r2"].append(r2)

fold_counter += 1

# plot val_pred and val_actual
plt.figure()

plt.plot(val_actual_subset, label="actual")
plt.plot(val_pred_subset, label="predicted")
plt.legend()
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)
plt.show()

# -----------------------------------------

# Calculate average metrics across all folds
avg_mse = np.mean(scores["test_mse"])
avg_mae = np.mean(scores["test_mae"])
avg_r2 = np.mean(scores["r2"])

print(f"Average MSE: {avg_mse:.4f}")
print(f"Average MAE: {avg_mae:.4f}")
print(f"Average R-squared: {avg_r2:.4f}")

doc.add_heading("Average Metrics", level=2)
doc.add_paragraph(f"Average MSE: {avg_mse:.4f}")
doc.add_paragraph(f"Average MAE: {avg_mae:.4f}")
doc.add_paragraph(f"Average R-squared: {avg_r2:.4f}")

doc.add_picture("temp_fig.png")
doc.save("output.docx")

# %%
data.head().to_json(orient="records")

# %% [markdown]
# ## Plot Result of train only SO2

# %%
pd.date_range(start=data.index[5], periods=timesteps * 4, freq="H")[1]

# %% [markdown]
# ## Model Old
#

# %%
from keras.utils import plot_model


plot_model(
    model_for_so2,
    to_file="model_structure.png",
    show_shapes=True,
    show_layer_names=True,
)


# Add the model structure image to the document
doc.add_heading("Model Structure", level=2)
doc.add_paragraph("The structure of the LSTM model for SO2 as only feature:")
doc.add_picture("model_structure.png")
doc.save("output.docx")


plot_model(
    model_multi, to_file="model_structure.png", show_shapes=True, show_layer_names=True
)


# Add the model structure image to the document

doc.add_paragraph("The structure of the LSTM model: for multi features:")
doc.add_picture("model_structure.png")
doc.save("output.docx")

# %%
# from keras import activations


# # # Split the data into training and testing sets using TimeSeriesSplit
# # tscv = TimeSeriesSplit(n_splits=5, test_size=48)


# # doc heading
# doc.add_heading("Model Training", level=1)
# doc.add_heading("Model Info", level=2)
# doc.add_paragraph(f"Model: LSTM")
# doc.add_paragraph(f"Features: {selected_features}")
# doc.add_paragraph(f"Target: {target}")
# doc.add_paragraph(f"Number of splits: {N_SPLITS}")


# doc.add_heading(f"Split {i+1}", level=2)
# doc.add_paragraph(f"Train index: {train_index}")
# doc.add_paragraph(f"Test index: {test_index}")


# # revers tscv.split(X_scaled)


# for ix, (train_index, test_index) in list(enumerate(tscv.split(data))):
#     print("TRAIN:", train_index, "TEST:", test_index)

#     X_train, X_test = X_scaled[train_index], X_scaled[test_index]
#     y_train, y_test = y[train_index], y[test_index]

#     print(f"Train: {X_train.shape} Test: {X_test.shape}")

#     # Reshape the input data for LSTM
#     X_train = X_train.reshape((X_train.shape[0], 1, X_train.shape[1]))
#     X_test = X_test.reshape((X_test.shape[0], 1, X_test.shape[1]))

#     look_back = 1
#     forecast_horizon = 1

#     # Build the LSTM model
#     with tf.device("/gpu:0"):  # Use GPU for model training
#         model = Sequential()
#         model.add(
#             LSTM(
#                 128,
#                 input_shape=(look_back, len(selected_features)),
#                 return_sequences=True,
#             )
#         )
#         model.add(Dropout(0.2))
#         model.add(LSTM(64, return_sequences=True))
#         model.add(Dropout(0.2))
#         model.add(LSTM(32))
#         model.add(Dropout(0.2))
#         model.add(Dense(forecast_horizon))  # Add a dense layer for the output
#         model.compile(loss="mean_squared_error", optimizer="adam")

#     # Train the model and get the history
#     history = model.fit(
#         X_train, y_train, epochs=50, batch_size=1024, validation_data=(X_test, y_test)
#     )

#     # Plot the loss graph
#     plt.figure()  # figsize=(8, 6)
#     plt.plot(history.history["loss"], label="Training Loss", lw=0.5)
#     plt.plot(history.history["val_loss"], label="Validation Loss", lw=0.5)
#     plt.xlabel("Epoch", fontsize=10)
#     plt.ylabel("Loss", fontsize=10)
#     plt.title("Training and Validation Loss over Epochs", fontsize=10)
#     plt.legend()
#     plt.savefig(
#         "temp_fig.png",
#         dpi=300,
#         bbox_inches="tight",
#         facecolor="white",
#         transparent=False,
#     )
#     plt.show()

#     # Make predictions for the training and testing sets
#     y_train_pred = model.predict(X_train)
#     y_test_pred = model.predict(X_test)

#     # Make predictions for the testing set
#     y_pred = model.predict(X_test)

#     y_pred = y_pred.reshape(y_test.shape)

#     # Calculate evaluation metrics
#     mse = mean_squared_error(y_test, y_pred)
#     mae = mean_absolute_error(y_test, y_pred)
#     r2 = r2_score(y_test, y_pred)

#     train_mse = model.evaluate(X_train, y_train, verbose=0)
#     train_mae = mean_absolute_error(y_train, y_train_pred)

#     scores["test_mse"].append(mse)
#     scores["test_mae"].append(mae)
#     scores["r2"].append(r2)
#     scores["train_mse"].append(train_mse)
#     scores["train_mae"].append(train_mae)

#     print("Testing Set Metrics:")
#     print("Mean Squared Error (MSE):", mse)
#     print("Mean Absolute Error (MAE):", mae)
#     print("R-squared (R2):", r2)
#     print("Train MSE:", train_mse)
#     print("Train MAE:", train_mae)

#     # show shape
#     print(f"X_train shape: {X_train.shape}")
#     print(f"X_test shape: {X_test.shape}")
#     print(f"y_train shape: {y_train.shape}")
#     print(f"y_test shape: {y_test.shape}")

#     # add to doc
#     doc.add_heading("Model Training", level=2)
#     doc.add_paragraph(f"Testing Set Metrics:")
#     doc.add_paragraph(f"Mean Squared Error (MSE): {mse}")
#     doc.add_paragraph(f"Mean Absolute Error (MAE): {mae}")
#     doc.add_paragraph(f"R-squared (R2): {r2}")
#     doc.add_paragraph(f"Train MSE: {train_mse}")
#     doc.add_paragraph(f"Train MAE: {train_mae}")

#     # add training and testing size
#     doc.add_heading("Data Size", level=2)
#     doc.add_paragraph(f"indices: {train_index}")
#     doc.add_paragraph(f"X_train shape: {X_train.shape}")
#     doc.add_paragraph(f"X_test shape: {X_test.shape}")
#     doc.add_paragraph(f"y_train shape: {y_train.shape}")
#     doc.add_paragraph(f"y_test shape: {y_test.shape}")

#     # percent of training and testing
#     doc.add_paragraph(
#         f"Training size: {X_train.shape[0] / X_scaled.shape[0] * 100:.2f}%"
#     )
#     doc.add_paragraph(f"Testing size: {X_test.shape[0] / X_scaled.shape[0] * 100:.2f}%")

#     # add fig
#     doc.add_picture(
#         "temp_fig.png",
#     )

#     # add line
#     doc.add_paragraph("--------------------------------------------------")

#     doc.save("output.docx")

#     # save score to text
#     with open("output_score.txt", "a") as f:
#         f.write(str(scores))


# # Make predictions for the entire dataset
# X_scaled_full = X_scaled.reshape((X_scaled.shape[0], 1, X_scaled.shape[1]))
# predictions = model.predict(X_scaled_full)


# predictions = predictions.reshape(y.shape)  # Reshape predictions to match y


# # # Make predictions for the next 48 hours
# # predictions = model.predict(X_test)


# # Create a dataframe for the predicted values
# predicted_data = pd.DataFrame(predictions, columns=["Predicted"], index=data.index)


# # # Create a dataframe for the predicted values
# # predicted_data = pd.DataFrame(predictions, columns=['Predicted'], index=data.index[test_index])


# # Plot the timeseries graph for the entire 10-year period
# plt.figure(figsize=(20, 6))
# plt.plot(data.index, data[target], label="Real", lw=0.5, alpha=0.7)
# plt.plot(
#     predicted_data.index,
#     predicted_data["Predicted"],
#     label="Predicted",
#     lw=0.5,
#     color="red",
# )
# plt.xlabel("Datetime")
# plt.ylabel("SO2 Value")
# plt.title("Real vs Predicted SO2 Values (10 Years)")
# plt.legend()
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )


# plt.show()


# # add to doc
# doc.add_heading("Model Prediction", level=1)
# doc.add_heading("Model Info", level=2)
# doc.add_paragraph(f"Model: LSTM")
# doc.add_paragraph(f"Features: {selected_features}")
# doc.add_paragraph(f"Target: {target}")


# # add fig
# doc.add_picture("temp_fig.png")


# doc.save("output.docx")

# %%
# # Show Predicted graph only year 2015
# plt.figure(figsize=(20, 6))
# plt.plot(data.index, data[target], label="Real", lw=0.5, alpha=0.7, color="blue")
# plt.plot(
#     predicted_data.index,
#     predicted_data["Predicted"],
#     label="Predicted",
#     lw=1,
#     color="red",
# )
# plt.xlabel("Datetime", fontsize=10)
# plt.ylabel("SO2 Value", fontsize=10)
# plt.title("Real vs Predicted SO2 Values (2016)", fontsize=10)
# plt.legend(fontsize=10, loc="upper right")
# # xlimit from datetime
# xlimit = pd.to_datetime(["2016-01-01", "2016-12-31"])
# plt.xlim(xlimit)


# # ylime auto
# plt.ylim(-10, 120)


# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )


# # add to doc
# # title is show only year 2015
# doc.add_heading("Model Prediction (one year)", level=1)


# # add fig
# doc.add_picture("temp_fig.png")

# %%
# train_mse = model.evaluate(X_train, y_train, verbose=0)
# train_mse

# %%
# test_mse = model.evaluate(X_test, y_test, verbose=0)
# test_mse

# %%
# y_pred_full = predictions.copy()

# %%
# # Assuming you have the predicted and observed values in the following variables:
# # predictions: a 1D array of predicted values
# # target: a 1D array of observed values


# y_target = data[[target]]


# print(y_target.shape)
# print(y_pred_full.shape)


# # Create a scatter plot
# plt.scatter(y_target, y_pred_full, color="blue", alpha=0.5, s=0.5)


# # sort y_target
# y_target_sort = y_target.sort_values(by=target)
# y_pred_full_sort = pd.DataFrame(predictions, index=y_target.index).sort_values(by=0)


# plt.plot(y_target_sort, y_pred_full_sort, color="red", linestyle="-", linewidth=2)


# # Set the labels and title
# plt.xlabel("Observed SO2 concentrations (μg/m³)")
# plt.ylabel("Predicted SO2 concentrations (μg/m³)")
# plt.title("Predicted and observed values of the test set")
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )


# # Display the plot
# plt.show()


# # add to doc , scatter plot of predicted and observed values
# doc.add_heading("Model Prediction", level=1)
# doc.add_heading("Model Info", level=2)
# doc.add_paragraph(f"Model: LSTM")
# doc.add_paragraph(f"Features: {selected_features}")
# doc.add_paragraph(f"Target: {target}")


# save_doc_fig("Scatter Plot", "Scatter plot of predicted and observed values:")
# doc.save("output.docx")

# %%
# import numpy as np
# import seaborn as sns
# import matplotlib.pyplot as plt


# # Assuming you have Y_test and y_pred as numpy arrays
# Y_test = np.array(y_target)
# y_pred = np.array(y_pred_full)
# # # Plot scatter plot with fitting line
# # sns.regplot(x=Y_test, y=y_pred)


# print(Y_test.shape)
# print(y_pred.shape)


# plt.scatter(range(len(Y_test)), Y_test, color="blue", s=0.5)
# plt.scatter(range(len(y_pred)), y_pred, color="red", s=0.5)
# plt.title("regressorName")
# plt.legend(["Real", "Predicted"], loc="upper right", fontsize=10)


# # save fig\
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )


# save_doc_fig("Scatter Plot", "Scatter plot of predicted and observed values:")
# plt.show()


# # line plot y_test


# plt.plot(
#     Y_test,
#     label="Real",
#     color="blue",
#     lw=0.5,
#     marker="o",
#     markersize=1,
#     markeredgecolor="black",
# )
# plt.plot(y_pred, label="Predicted", color="red", lw=0.5)
# plt.legend(loc="upper right", fontsize=10)


# # save fig
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )


# plt.show()


# # plt.show()


# save_doc_fig("Scatter Plot", "Scatter plot of predicted and observed values:")

# %% [markdown]
# # Descript Info

# %%
# hist = data.plot.hist(subplots=True, figsize=(20, 20), bins=50, alpha=0.5).flatten()


# Determine the number of rows needed
num_features = len(data.columns)
num_rows = (num_features + 3) // 4  # Round up to the nearest integer


# Create a subplot grid
fig, axes = plt.subplots(nrows=num_rows, ncols=4, figsize=(15, 10))
axes = axes.flatten()


# Iterate over each feature and create a histogram subplot
for i, col in enumerate(data.columns):
    data[col].plot.hist(ax=axes[i], bins=50, alpha=0.5)
    # sns.histplot(data=data, x=col, ax=axes[row, col])
    axes[i].set_title(col, fontsize=10)
    axes[i].set_xlabel("")
    axes[i].set_ylabel("")


# Adjust the spacing between subplots
plt.tight_layout()


# save fig\
plt.savefig(
    "temp_fig.png",
    dpi=300,
    bbox_inches="tight",
    facecolor="white",
    transparent=False,
)


# Show the plot
plt.show()


# add header
doc.add_heading("Data Distribution", level=1)
doc.add_heading("Data Info", level=2)
doc.add_paragraph(f"The data has {data.shape[0]} rows and {data.shape[1]} columns")
doc.add_paragraph(f"The data has {data.isna().sum().sum()} missing values")
doc.add_paragraph(f"The data has {data.duplicated().sum()} duplicated values")


save_doc_fig("Data Distribution", "Data distribution of all features:")

# %%
# PREDICT_TIME = 34564
# PREDICT_TIME_END = PREDICT_TIME + 24


# # Predict the next 48 hours
# X_scaled_full_hour = X_scaled_full[PREDICT_TIME:PREDICT_TIME_END]
# predictions = model.predict(X_scaled_full_hour)


# # Create a dataframe for the predicted values
# predicted_data = pd.DataFrame(
#     predictions, columns=["Predicted"], index=data.index[PREDICT_TIME:PREDICT_TIME_END]
# )


# pd_plot = data[[target] + selected_features][PREDICT_TIME:PREDICT_TIME_END].plot(
#     subplots=True, lw=1, figsize=(15, 15), fontsize=10
# )


# set_plot()


# save_doc_fig(
#     "Real vs Predicted SO2 Values (Next 24 Hours)",
#     "Real vs Predicted SO2 Values (Next 24 Hours)",
# )


# # plot the timeseries graph for the next 48 hours period and add pd_plot to the subplot
# plt.figure(figsize=(15, 5))
# plt.plot(
#     data.index[PREDICT_TIME:PREDICT_TIME_END],
#     data[target][PREDICT_TIME:PREDICT_TIME_END],
#     label="Real",
#     lw=1,
# )
# plt.plot(predicted_data.index, predicted_data["Predicted"], label="Predicted", lw=1)
# plt.xlabel("Datetime")
# plt.ylabel("SO2 Value")
# plt.title("Real vs Predicted Ambient SO2 Concentration (Next 24 Hours)")
# plt.legend()
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )
# plt.show()


# save_doc_fig(
#     "Real vs Predicted SO2 Values (Next 24 Hours)",
#     "Real vs Predicted SO2 Values (Next 24 Hours)",
# )


# doc.save("output.docx")

# %%
# print(data[[target]][PREDICT_TIME:PREDICT_TIME_END].shape)
# print(predictions.shape)

# plt.figure(figsize=(10, 8))
# plt.scatter(data[[target]][PREDICT_TIME:PREDICT_TIME_END], predictions, alpha=0.5)
# # save fig\
# plt.savefig(
#     "temp_fig.png",
#     dpi=300,
#     bbox_inches="tight",
#     facecolor="white",
#     transparent=False,
# )

# save_doc_fig("Scatter Plot", "Scatter plot of predicted and observed values:")

# %%
# rename output.docx to [file_name].docx


import os


os.replace("output.docx", f"docx/corr_{current_file_name}.docx")

current_file_name
