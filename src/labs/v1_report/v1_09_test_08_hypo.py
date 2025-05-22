# Import Libraries
print("Step 1: Importing Libraries...")
import gc
import os
import random
import time  # For timing operations
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import shap
import tensorflow as tf
import xgboost
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import MinMaxScaler
from tensorflow.keras.callbacks import EarlyStopping
from tensorflow.keras.layers import LSTM, Dense, Dropout
from tensorflow.keras.models import Sequential


def winsorize_series(series, lower_percentile=0.01, upper_percentile=0.99):
    """
    Performs Winsorization on a pandas Series.
    Clips outliers at the specified percentiles.
    """
    if not pd.api.types.is_numeric_dtype(series):
        print(f"DEBUG: winsorize_series - Series '{series.name}' is not numeric. Skipping.")
        return series  # Return original if not numeric

    # Ensure series is not all NaN
    if series.isnull().all():
        print(f"DEBUG: winsorize_series - Series '{series.name}' is all NaN. Skipping.")
        return series

    q_low = series.quantile(lower_percentile)
    q_high = series.quantile(upper_percentile)

    # Handle cases where q_low or q_high might be NaN (e.g., if series has too many NaNs or is constant)
    if pd.isna(q_low) and pd.isna(q_high):  # Both NaN, likely constant or all NaN
        print(f"DEBUG: winsorize_series - Both percentiles are NaN for '{series.name}'. Skipping clipping.")
        return series
    if pd.isna(q_low):  # Only lower is NaN, clip only upper
        print(f"DEBUG: winsorize_series - Lower percentile is NaN for '{series.name}'. Clipping only upper bound.")
        return series.clip(upper=q_high)
    if pd.isna(q_high):  # Only upper is NaN, clip only lower
        print(f"DEBUG: winsorize_series - Upper percentile is NaN for '{series.name}'. Clipping only lower bound.")
        return series.clip(lower=q_low)

    # Ensure q_low is not greater than q_high (can happen with very skewed data or small unique values)
    if q_low > q_high:
        print(f"DEBUG: winsorize_series - Lower percentile ({q_low:.2f}) > Upper percentile ({q_high:.2f}) for '{series.name}'. This might indicate issues with data distribution. Skipping clipping.")
        return series

    return series.clip(lower=q_low, upper=q_high)


# --- Helper Functions for Report Generation (Mostly Unchanged from your last provided version) ---
def describe_dataframe_skewness_for_report(df, doc, section_title, english_file_name):
    doc.add_heading(f"{section_title} - Skewness Analysis - {english_file_name}", level=3)
    numeric_df = df.select_dtypes(include=np.number)
    if numeric_df.empty:
        doc.add_paragraph("No numeric features to calculate skewness.")
        return
    try:
        skew_series = numeric_df.skew()
        if skew_series.empty or skew_series.isnull().all():
            doc.add_paragraph("Cannot calculate meaningful skewness (e.g., all NaNs or insufficient data in numeric columns).")
            return
    except Exception as e:
        doc.add_paragraph(f"Error calculating skewness: {e}")
        return
    abs_skew_series = skew_series.abs().dropna()
    if abs_skew_series.empty:
        doc.add_paragraph("No valid (non-NaN) absolute skewness values found.")
        return
    skew_max_val = abs_skew_series.max()
    skew_max_col = abs_skew_series.idxmax()
    skew_min_val = abs_skew_series.min()
    skew_min_col = abs_skew_series.idxmin()
    desc_para = doc.add_paragraph()
    desc_para.add_run(f"Feature with highest absolute skewness: ").bold = True
    desc_para.add_run(f"'{skew_max_col}' ({skew_max_val:.2f}).\n")
    desc_para.add_run(f"Feature with lowest absolute skewness: ").bold = True
    desc_para.add_run(f"'{skew_min_col}' ({skew_min_val:.2f}).\n")
    if skew_max_val > 1:
        desc_para.add_run("The highest skewness indicates a highly skewed distribution. ")
    elif skew_max_val > 0.5:
        desc_para.add_run("The highest skewness indicates a moderately skewed distribution. ")
    else:
        desc_para.add_run("The highest skewness indicates an approximately symmetric distribution. ")
    try:
        if skew_max_col in df.columns:
            stats_series = df[skew_max_col]
            desc_para.add_run(f"\nStats for '{skew_max_col}': Min={stats_series.min():.2f}, Max={stats_series.max():.2f}, Mean={stats_series.mean():.2f}, Std={stats_series.std():.2f}.")
        else:
            desc_para.add_run(f"\nCould not retrieve stats for '{skew_max_col}' (column not found).")
    except Exception as e_stats:
        desc_para.add_run(f"\nError retrieving stats for '{skew_max_col}': {e_stats}.")


def set_default_font(doc, font_name="TH SarabunPSK", font_size=11):
    for section in doc.sections:
        for header in section.header.paragraphs:
            for run in header.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        for footer in section.footer.paragraphs:
            for run in footer.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
    styles = doc.styles
    for style_name in [
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Caption",
        "List Paragraph",
    ]:
        try:
            style = styles[style_name]
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            if "Heading" in style_name:
                if style_name == "Heading 1":
                    font.size = Pt(16)
                    font.bold = True
                elif style_name == "Heading 2":
                    font.size = Pt(14)
                    font.bold = True
                elif style_name == "Heading 3":
                    font.size = Pt(12)
                    font.bold = True
            if hasattr(font._element, "rPr") and font._element.rPr is not None:
                rfonts = font._element.rPr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = type(font._element.rPr)._new_rFonts()
                    font._element.rPr.append(rfonts)
                rfonts.set(qn("w:ascii"), font_name)
                rfonts.set(qn("w:hAnsi"), font_name)
                rfonts.set(qn("w:cs"), font_name)
                rfonts.set(qn("w:eastAsia"), font_name)
        except KeyError:
            print(f"Warning: Style '{style_name}' not found.")


def add_table_to_doc(doc, df, title, description=""):
    doc.add_heading(title, level=2)
    if not isinstance(df, pd.DataFrame) or df.empty:
        doc.add_paragraph(f"Cannot add table '{title}': DataFrame is not valid or empty.")
        if description:
            doc.add_paragraph(description)
        return
    try:
        df_str = df.astype(str)
        table = doc.add_table(rows=df_str.shape[0] + 1, cols=df_str.shape[1], style="Table Grid")
        for j, col_name in enumerate(df_str.columns):
            table.cell(0, j).text = str(col_name)
        for i in range(df_str.shape[0]):
            for j in range(df_str.shape[1]):
                table.cell(i + 1, j).text = df_str.iloc[i, j]
        if description:
            doc.add_paragraph(description)
    except Exception as e:
        doc.add_paragraph(f"Error adding table '{title}': {e}")
        print(f"Error adding table '{title}' to docx: {e}")


def add_figure_to_doc(doc, filename, title, description=""):
    if isinstance(filename, Path):
        filename = str(filename)
    doc.add_heading(title, level=2)
    if os.path.exists(filename):
        try:
            doc.add_picture(filename, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Warning: Could not add image {filename}. Error: {e}")
    else:
        doc.add_paragraph(f"Warning: Image file {filename} not found.")
    if description:
        doc.add_paragraph(description)


def prepare_sequences(data, features, seq_length, pred_length):
    print(f"DEBUG: prepare_sequences - Input data shape: {data.shape}, num_features: {len(features)}")
    start_time = time.time()
    X, y = [], []
    if not features or "SO2" not in data.columns:
        print("DEBUG: prepare_sequences - Features list empty or SO2 missing. Returning empty arrays.")
        return np.array(X), np.array(y)

    valid_features = [f for f in features if f in data.columns]
    if not valid_features:
        print(f"DEBUG: prepare_sequences - None of provided features {features} found. Returning empty arrays.")
        return np.array(X), np.array(y)

    if len(data) < seq_length + pred_length:
        print(f"DEBUG: prepare_sequences - Data length {len(data)} too short. Returning empty arrays.")
        return np.array(X), np.array(y)

    # Optimized sequence preparation (slightly)
    feature_values = data[valid_features].values
    target_values = data["SO2"].values

    num_sequences = len(data) - seq_length - pred_length + 1
    if num_sequences <= 0:
        print("DEBUG: prepare_sequences - num_sequences is not positive. Returning empty arrays.")
        return np.array(X), np.array(y)

    # Pre-allocate numpy arrays if memory allows (can be very large)
    # This might be an issue for "หลายแสนบรรทัด" if seq_length is large
    # For very large data, appending to lists and then converting might be more memory-flexible
    # but slower. Let's stick to lists for now for flexibility.

    for i in range(num_sequences):
        X.append(feature_values[i : i + seq_length])
        y.append(target_values[i + seq_length : i + seq_length + pred_length])

    if not X or not y:
        print("DEBUG: prepare_sequences - X or y list is empty after loop. Returning empty arrays.")
        return np.array(X), np.array(y)

    X_np = np.array(X).astype(np.float32)
    y_np = np.array(y).astype(np.float32)
    end_time = time.time()
    print(f"DEBUG: prepare_sequences - X_full shape: {X_np.shape}, y_full shape: {y_np.shape}. Time taken: {end_time - start_time:.2f}s")
    return X_np, y_np


# --- Main Class ---
class DynamicInputDataset:
    class CONST:
        SEED_VALUE = 42
        BATCH_SIZE = 1024  # Adjusted for potentially large datasets, was 64
        EPOCHS = 50  # Can be reduced for faster debug runs, e.g., 5-10
        TEST_SPLIT_RATIO = 0.2
        SEQ_LENGTH = 24
        PRED_LENGTH = 24
        OUTPUT_FOLDER = "outputs_detailed_eda_debug"
        NUM_TOP_FEATURES_FOR_SELECTION = 5
        SHAP_IMPORTANCE_THRESHOLD = 0.005
        ITERATIVE_PRED_HORIZON_DAYS = 365
        # MAX_ROWS_TO_PROCESS = 100000 # New: Limit rows for debugging if needed
        MAX_ROWS_TO_PROCESS = None  # Set to None to process all rows

        valid_ranges = {
            "BP": (500, 1100),  # mmHg
            "NO2": (0, 10),  # ppm
            "PM10": (0, 1000),  # µg/m³
            "RAIN": (0, 500),  # mm (practical upper limit)
            "RH": (0, 100),  # %RH
            "SO2": (0, 10),  # ppm
            "SW": (0, 1500),  # w/m²
            "TEMP": (-50, 60),  # °C
            "TSP": (0, 1000),  # µg/m³
            "WD": (0, 360),  # degrees
            "WS": (0, 100),  # m/s
        }

    file_name_mapping = {
        "(37t)ศาลหลักเมือง(ปิดสถานี)": "San_Lak_Mueang_Closed_Station",
        "(37t)สถานีอุตุนิยมวิทยาลำปาง": "Lampang_Meteorological_Station",
        "(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด": "Ban_Sopad_Hospital",
        "(39t)โรงพยาบาลส่งเสริมสุขภาพตำบลท่าสี": "Tha_Si_Hospital",
        "(40t)การประปาส่วนภูมิภาคแม่เมาะ": "Mae_Mo_Waterworks",
    }

    def __init__(self, input_dir):
        self.input_dir = Path(input_dir)
        self.csv_files = sorted(list(self.input_dir.glob("*.csv")))
        self.output_dir = Path.cwd() / self.CONST.OUTPUT_FOLDER
        self.output_dir.mkdir(exist_ok=True)
        random.seed(self.CONST.SEED_VALUE)
        np.random.seed(self.CONST.SEED_VALUE)
        tf.random.set_seed(self.CONST.SEED_VALUE)
        gpus = tf.config.experimental.list_physical_devices("GPU")
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
                print("GPU memory growth enabled.")
            except RuntimeError as e:
                print(f"Error setting GPU memory growth: {e}")
        print(f"Num GPUs Available: {len(gpus)}")

    def _apply_range_validation(self, df_to_validate):
        """
        Validates data based on physically realistic ranges defined in self.CONST.valid_ranges.
        Values outside these ranges are set to NaN. WD is wrapped to 0-360.
        Returns the validated DataFrame and a dictionary of cleaning statistics.
        """
        print(f"DEBUG: _apply_range_validation - Starting. Input DF shape: {df_to_validate.shape}")
        if not hasattr(self.CONST, "valid_ranges"):
            print("WARNING: self.CONST.valid_ranges not defined. Range validation will be skipped.")
            return df_to_validate.copy(), {}  # Return a copy and empty stats

        valid_ranges = self.CONST.valid_ranges
        df_validated = df_to_validate.copy()
        range_cleaning_stats = {}  # Stores count of new NaNs per column
        cols_changed_by_range = []  # Stores names of cols affected

        for col, (min_val, max_val) in valid_ranges.items():
            if col in df_validated.columns and pd.api.types.is_numeric_dtype(df_validated[col]):
                original_nan_count = df_validated[col].isnull().sum()
                condition_out_of_range = (df_validated[col] < min_val) | (df_validated[col] > max_val)

                if condition_out_of_range.any():
                    if col not in cols_changed_by_range:  # Avoid duplicate logging message parts
                        cols_changed_by_range.append(f"'{col}' (range: {min_val}-{max_val})")
                    df_validated.loc[condition_out_of_range, col] = np.nan

                # Specific handling for WD (Wind Direction) after NaN setting
                if col == "WD" and df_validated[col].notna().any():
                    # Apply modulo only to non-NaN values to ensure they wrap around 360 degrees
                    df_validated.loc[df_validated[col].notna(), col] = df_validated.loc[df_validated[col].notna(), col] % 360

                new_nan_count = df_validated[col].isnull().sum()
                if new_nan_count > original_nan_count:
                    range_cleaning_stats[col] = new_nan_count - original_nan_count
            elif col in df_validated.columns:
                print(f"DEBUG: _apply_range_validation - Column '{col}' found but is not numeric. Skipping range validation for it.")

        print(f"DEBUG: _apply_range_validation - Finished. Output DF shape: {df_validated.shape}")
        return df_validated, range_cleaning_stats, cols_changed_by_range

    def _apply_winsorization(self, df_to_winsorize):
        """
        Applies Winsorization to all numeric columns of the DataFrame.
        Returns the Winsorized DataFrame, list of considered columns, and list of changes.
        """
        print(f"DEBUG: _apply_winsorization - Starting. Input DF shape: {df_to_winsorize.shape}")
        df_winsorized = df_to_winsorize.copy()
        numeric_cols_considered = df_winsorized.select_dtypes(include=np.number).columns.tolist()
        winsorized_details_list = []

        if numeric_cols_considered:
            print(f"DEBUG: _apply_winsorization - Applying to: {', '.join(numeric_cols_considered)}")
            for col_w in numeric_cols_considered:
                if df_winsorized[col_w].isnull().all() or df_winsorized[col_w].nunique() < 2:
                    print(f"DEBUG: Winsorize skipping '{col_w}' in _apply_winsorization (all NaN or constant).")
                    continue

                original_min_w = df_winsorized[col_w].min()
                original_max_w = df_winsorized[col_w].max()

                df_winsorized[col_w] = winsorize_series(df_winsorized[col_w])  # Call the global/static winsorize_series

                min_after = df_winsorized[col_w].min()
                max_after = df_winsorized[col_w].max()

                original_min_w_str = f"{original_min_w:.2f}" if pd.notna(original_min_w) else "NaN"
                original_max_w_str = f"{original_max_w:.2f}" if pd.notna(original_max_w) else "NaN"
                min_after_str = f"{min_after:.2f}" if pd.notna(min_after) else "NaN"
                max_after_str = f"{max_after:.2f}" if pd.notna(max_after) else "NaN"

                changed = False
                if pd.notna(original_min_w) and pd.notna(min_after) and not np.isclose(original_min_w, min_after):
                    changed = True
                if pd.notna(original_max_w) and pd.notna(max_after) and not np.isclose(original_max_w, max_after):
                    changed = True
                if pd.isna(original_min_w) != pd.isna(min_after) or pd.isna(original_max_w) != pd.isna(max_after):
                    changed = True

                if changed:
                    winsorized_details_list.append(f"'{col_w}' (orig min/max: {original_min_w_str}/{original_max_w_str}, new min/max: {min_after_str}/{max_after_str})")
        else:
            print("DEBUG: _apply_winsorization - No numeric columns found.")

        print(f"DEBUG: _apply_winsorization - Finished. Output DF shape: {df_winsorized.shape}")
        return df_winsorized, numeric_cols_considered, winsorized_details_list

    def _perform_full_eda(self, doc, df_input, eda_stage_name, english_file_name, file_name_base):
        """Helper function to perform and report a full EDA stage."""
        print(f"DEBUG: Starting EDA Stage: {eda_stage_name} for {file_name_base}. DF shape: {df_input.shape}")
        doc.add_heading(f"Exploratory Data Analysis ({eda_stage_name}) - {file_name_base}", level=1)
        doc.add_paragraph(f"Current Data Shape: {df_input.shape}")
        doc.add_paragraph(f"Total Missing Values: {df_input.isnull().sum().sum()}")

        if df_input.empty:
            doc.add_paragraph(f"DataFrame is empty at '{eda_stage_name}' stage. EDA cannot proceed.")
            print(f"DEBUG: DataFrame empty at EDA stage {eda_stage_name}.")
            return

        doc.add_heading(f"Data Preview ({eda_stage_name})", level=2)
        doc.add_heading(f"Columns List ({eda_stage_name})", level=3)
        for col in df_input.columns:
            doc.add_paragraph(f"- {col}")
        add_table_to_doc(
            doc,
            df_input.head(10),
            f"Sample of First 10 Rows ({eda_stage_name})",
            f"Showing the first 10 rows at '{eda_stage_name}' stage.",
        )
        doc.add_heading(f"Data Information (dtypes, non-null count) ({eda_stage_name})", level=3)
        from io import StringIO

        buffer = StringIO()
        df_input.info(buf=buffer)
        info_str = buffer.getvalue()
        doc.add_paragraph(info_str)

        plot_suffix = f"{eda_stage_name.lower().replace(' ', '_')}_{english_file_name}"
        plt.figure(figsize=(12, 6))
        sns.heatmap(df_input.isnull(), cbar=True, cmap="viridis", annot=False)
        plt.title(f"Missing Data Heatmap ({eda_stage_name}) - {english_file_name}")
        heatmap_path = self.output_dir / f"heatmap_{plot_suffix}.png"
        plt.savefig(heatmap_path, dpi=300, bbox_inches="tight")
        plt.close()
        add_figure_to_doc(
            doc,
            heatmap_path,
            f"Missing Data Heatmap ({eda_stage_name})",
            f"Heatmap of missing values at '{eda_stage_name}'.",
        )

        if "SO2" in df_input.columns and pd.api.types.is_numeric_dtype(df_input["SO2"]):
            numeric_df_for_corr = df_input.select_dtypes(include=np.number)
            if not numeric_df_for_corr.empty and "SO2" in numeric_df_for_corr.columns:
                corr_matrix = numeric_df_for_corr.corr()
                if "SO2" in corr_matrix.columns:
                    so2_corr = corr_matrix[["SO2"]].sort_values(by="SO2", ascending=False)
                    add_table_to_doc(
                        doc,
                        so2_corr,
                        f"Sorted Correlation with SO2 ({eda_stage_name})",
                        f"Correlation with SO2 at '{eda_stage_name}'.",
                    )
                plt.figure(
                    figsize=(
                        max(12, len(corr_matrix.columns) * 0.5),
                        max(10, len(corr_matrix.columns) * 0.4),
                    )
                )  # Dynamic size for corr matrix
                sns.heatmap(corr_matrix, annot=True, cmap="coolwarm", fmt=".2f", linewidths=0.5)
                plt.title(f"Correlation Matrix ({eda_stage_name}) - {english_file_name}")
                plt.xticks(rotation=45, ha="right")
                plt.yticks(rotation=0)
                corr_heatmap_path = self.output_dir / f"corr_heatmap_{plot_suffix}.png"
                plt.savefig(corr_heatmap_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    corr_heatmap_path,
                    f"Correlation Matrix Heatmap ({eda_stage_name})",
                    f"Correlations at '{eda_stage_name}'.",
                )
            else:
                doc.add_paragraph(f"Not enough numeric data or SO2 for correlation at '{eda_stage_name}'.")
        else:
            doc.add_paragraph(f"SO2 not found/numeric for correlation at '{eda_stage_name}'.")

        numeric_cols_for_boxplot = df_input.select_dtypes(include=np.number).columns
        if not numeric_cols_for_boxplot.empty:
            plt.figure(figsize=(15, max(6, len(numeric_cols_for_boxplot) * 0.6)))
            df_input[numeric_cols_for_boxplot].boxplot(rot=45, grid=False)
            plt.title(f"Feature Outlier Distribution ({eda_stage_name}) - {english_file_name}")
            plt.tight_layout()
            boxplot_path = self.output_dir / f"boxplot_{plot_suffix}.png"
            plt.savefig(boxplot_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                boxplot_path,
                f"Feature Outlier Distribution ({eda_stage_name})",
                f"Boxplots at '{eda_stage_name}'.",
            )
        else:
            doc.add_paragraph(f"No numeric columns for boxplot at '{eda_stage_name}'.")

        numeric_cols_for_hist = df_input.select_dtypes(include=np.number).columns
        if not numeric_cols_for_hist.empty:
            n_cols_hist = min(4, len(numeric_cols_for_hist))
            n_rows_hist = (len(numeric_cols_for_hist) + n_cols_hist - 1) // n_cols_hist
            plt.figure(figsize=(15, n_rows_hist * 3.5))
            for i, col in enumerate(numeric_cols_for_hist):
                plt.subplot(n_rows_hist, n_cols_hist, i + 1)
                sns.histplot(df_input[col].dropna(), kde=True)
                plt.title(col)
            plt.tight_layout()
            hist_path = self.output_dir / f"histograms_{plot_suffix}.png"
            plt.savefig(hist_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                hist_path,
                f"Feature Histograms ({eda_stage_name})",
                f"Histograms at '{eda_stage_name}'.",
            )
            describe_dataframe_skewness_for_report(df_input, doc, f"Skewness ({eda_stage_name})", english_file_name)
        else:
            doc.add_paragraph(f"No numeric columns for histogram/skewness at '{eda_stage_name}'.")

        if isinstance(df_input.index, pd.DatetimeIndex) and not df_input.empty:
            numeric_cols_for_ts = df_input.select_dtypes(include=np.number).columns
            if not numeric_cols_for_ts.empty:
                n_features_ts = len(numeric_cols_for_ts)
                plt.figure(figsize=(15, n_features_ts * 2.5 if n_features_ts > 0 else 5))
                for i, col in enumerate(numeric_cols_for_ts):
                    plt.subplot(n_features_ts, 1, i + 1)
                    plt.plot(df_input.index, df_input[col], label=col)
                    plt.title(f"Timeseries of {col} - {english_file_name}")
                    plt.legend(loc="upper right")
                plt.tight_layout()
                ts_path = self.output_dir / f"timeseries_all_{plot_suffix}.png"
                plt.savefig(ts_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    ts_path,
                    f"Timeseries Plot for All Numeric Features ({eda_stage_name})",
                    f"Timeseries plots at '{eda_stage_name}'.",
                )
            else:
                doc.add_paragraph(f"No numeric columns for timeseries plot at '{eda_stage_name}'.")
        else:
            doc.add_paragraph(f"Index not DatetimeIndex or df empty. Skipping timeseries for '{eda_stage_name}'.")
        print(f"DEBUG: Finished EDA Stage: {eda_stage_name}")

    def _plot_train_test_split_bar_chart(self, train_size, test_size, model_suffix, english_file_name, doc):
        dataset_info = pd.DataFrame([{"Train": train_size, "Validation": test_size}])
        plt.figure(figsize=(6, 4))
        dataset_info.plot(kind="barh", stacked=True, ax=plt.gca())
        total_text = train_size + test_size
        train_percentage = (train_size / total_text * 100) if total_text > 0 else 0
        validation_percentage = (test_size / total_text * 100) if total_text > 0 else 0
        plt.text(
            total_text,
            0,
            f"{total_text}",
            ha="left",
            va="center",
            color="black",
            fontsize=9,
        )
        plt.text(
            train_size / 2,
            0,
            f"{train_size}\n({train_percentage:.1f}%)",
            ha="center",
            va="center",
            color="white",
            fontsize=9,
        )
        plt.text(
            train_size + test_size / 2,
            0,
            f"{test_size}\n({validation_percentage:.1f}%)",
            ha="center",
            va="center",
            color="black",
            fontsize=9,
        )
        plt.xlabel("Dataset Size (Sequences)")
        plt.ylabel("Split")
        plt.title(f"Train-Test Split ({model_suffix.replace('_', ' ').title()}) - {english_file_name}")
        plt.yticks([])
        plt.legend(["Train", "Validation"], loc="lower right")
        chart_path = self.output_dir / f"train_test_split_{model_suffix}_{english_file_name}.png"
        plt.savefig(
            chart_path,
            dpi=300,
            bbox_inches="tight",
            facecolor="white",
            transparent=False,
        )
        plt.close()
        add_figure_to_doc(
            doc,
            chart_path,
            f"Train-Test Split ({model_suffix.replace('_', ' ').title()})",
            f"Bar chart of sequences for {model_suffix.replace('_', ' ')} model.",
        )

    def _train_evaluate_report_lstm(
        self,
        doc,
        df_scaled_input,
        features_to_use,
        scaler_target,
        model_suffix,
        english_file_name,
        file_name_base,
        original_df_for_time_ref,
    ):
        print(f"DEBUG: _train_evaluate_report_lstm - Starting for model_suffix: {model_suffix}")
        report_title_suffix = model_suffix.replace("_", " ").title()
        doc.add_heading(
            f"LSTM Model Training & Evaluation ({report_title_suffix}) - {file_name_base}",
            1,
        )

        if not features_to_use:
            doc.add_paragraph(f"No features for '{report_title_suffix}'. Skipping training.")
            print(f"DEBUG: _train_evaluate_report_lstm - No features for {model_suffix}. Skipping.")
            return None, {}, None, None

        doc.add_paragraph(f"Features used: {', '.join(features_to_use)}")
        doc.add_paragraph(f"Seq Length: {self.CONST.SEQ_LENGTH}, Pred Length: {self.CONST.PRED_LENGTH}")

        print(f"DEBUG: _train_evaluate_report_lstm - Calling prepare_sequences for {model_suffix} with {len(features_to_use)} features.")
        X_full, y_full = prepare_sequences(
            df_scaled_input,
            features_to_use,
            self.CONST.SEQ_LENGTH,
            self.CONST.PRED_LENGTH,
        )

        if X_full.size == 0 or y_full.size == 0:
            doc.add_paragraph(f"Not enough data for sequences ({report_title_suffix}). X_shape: {X_full.shape}, y_shape: {y_full.shape}. Skipping.")
            print(f"DEBUG: _train_evaluate_report_lstm - Not enough data for sequences for {model_suffix}. X_full: {X_full.shape}, y_full: {y_full.shape}")
            return None, {}, None, None

        print(f"DEBUG: _train_evaluate_report_lstm - prepare_sequences successful. X_full shape: {X_full.shape}")

        print(f"DEBUG: _train_evaluate_report_lstm - Starting train_test_split for {model_suffix}.")
        split_start_time = time.time()
        X_train, X_test, y_train, y_test = train_test_split(
            X_full,
            y_full,
            test_size=self.CONST.TEST_SPLIT_RATIO,
            shuffle=False,
            random_state=self.CONST.SEED_VALUE,
        )
        split_end_time = time.time()
        print(f"DEBUG: _train_evaluate_report_lstm - train_test_split done. X_train: {X_train.shape}, X_test: {X_test.shape}. Time: {split_end_time - split_start_time:.2f}s")

        self._plot_train_test_split_bar_chart(len(X_train), len(X_test), model_suffix, english_file_name, doc)

        model_filename_h5 = self.output_dir / f"{file_name_base}-lstm-{model_suffix}.h5"
        history_data = {"loss": [], "val_loss": []}
        is_model_trained_this_run = False
        model = None

        if model_filename_h5.exists():
            print(f"DEBUG: _train_evaluate_report_lstm - Loading pre-trained model: {model_filename_h5}")
            try:
                model = tf.keras.models.load_model(model_filename_h5, compile=True)
            except Exception as e:
                print(f"DEBUG: Failed to load compiled {model_filename_h5}, trying compile=False. Error: {e}")
                try:
                    model = tf.keras.models.load_model(model_filename_h5, compile=False)
                    model.compile(optimizer="adam", loss="mse")
                except Exception as e_recompile:
                    print(f"DEBUG: Failed to load/recompile {model_filename_h5}: {e_recompile}")
                    model = None

        if model is None:
            print(f"DEBUG: _train_evaluate_report_lstm - Training new model for {model_suffix}. X_train shape: {X_train.shape}")
            tf.keras.backend.clear_session()
            model = Sequential(
                [
                    LSTM(
                        128,
                        input_shape=(self.CONST.SEQ_LENGTH, X_train.shape[-1]),
                        return_sequences=True,
                    ),
                    Dropout(0.2),
                    LSTM(64, return_sequences=False),
                    Dropout(0.2),
                    Dense(32, activation="relu"),
                    Dense(self.CONST.PRED_LENGTH),
                ]
            )
            model.compile(optimizer=tf.keras.optimizers.Adam(learning_rate=0.001), loss="mse")
            early_stopping = EarlyStopping(monitor="val_loss", patience=10, restore_best_weights=True, verbose=1)

            print(f"DEBUG: _train_evaluate_report_lstm - Starting model.fit for {model_suffix}.")
            fit_start_time = time.time()
            history = model.fit(
                X_train,
                y_train,
                epochs=self.CONST.EPOCHS,
                batch_size=self.CONST.BATCH_SIZE,
                validation_data=(X_test, y_test),
                verbose=1,
                callbacks=[early_stopping],
            )
            fit_end_time = time.time()
            print(f"DEBUG: _train_evaluate_report_lstm - model.fit completed for {model_suffix}. Time: {fit_end_time - fit_start_time:.2f}s")
            history_data = history.history
            model.save(model_filename_h5)
            is_model_trained_this_run = True

        model.summary(print_fn=lambda x: doc.add_paragraph(x))
        print(f"DEBUG: _train_evaluate_report_lstm - Starting model.predict for {model_suffix}.")
        predict_start_time = time.time()
        y_pred_scaled = model.predict(X_test)
        predict_end_time = time.time()
        print(f"DEBUG: _train_evaluate_report_lstm - model.predict done. Time: {predict_end_time - predict_start_time:.2f}s")

        y_test_inv = scaler_target.inverse_transform(np.nan_to_num(y_test))
        y_pred_inv = scaler_target.inverse_transform(np.nan_to_num(y_pred_scaled))
        valid_mask = ~np.isnan(y_test_inv).any(axis=1) & ~np.isnan(y_pred_inv).any(axis=1) if y_test_inv.ndim > 1 else ~np.isnan(y_test_inv) & ~np.isnan(y_pred_inv)
        y_test_inv_clean = y_test_inv[valid_mask]
        y_pred_inv_clean = y_pred_inv[valid_mask]
        scores = {"mse": np.nan, "mae": np.nan, "r2": np.nan}
        if y_test_inv_clean.size > 0 and y_pred_inv_clean.size > 0:
            scores["mse"] = mean_squared_error(y_test_inv_clean, y_pred_inv_clean)
            scores["mae"] = mean_absolute_error(y_test_inv_clean, y_pred_inv_clean)
            scores["r2"] = r2_score(y_test_inv_clean, y_pred_inv_clean)
        else:
            doc.add_paragraph(f"Not enough valid data for metrics ({report_title_suffix}).")

        if is_model_trained_this_run and history_data.get("loss") and history_data.get("val_loss"):
            if history_data["loss"] and history_data["val_loss"]:
                plt.figure(figsize=(10, 6))
                plt.plot(history_data["loss"], label="Training Loss")
                plt.plot(history_data["val_loss"], label="Validation Loss")
                plt.title(f"Model Loss ({report_title_suffix}) - {english_file_name}")
                plt.xlabel("Epoch")
                plt.ylabel("Loss (MSE)")
                plt.legend()
                loss_plot_path = self.output_dir / f"loss_{model_suffix}_{english_file_name}.png"
                plt.savefig(loss_plot_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    loss_plot_path,
                    f"Model Loss ({report_title_suffix})",
                    "Training and validation loss.",
                )
        else:
            doc.add_paragraph(f"Loss plot not generated for {report_title_suffix} (loaded/no history).")

        y_test_plot_main = y_test_inv_clean[:, 0] if y_test_inv_clean.ndim > 1 and y_test_inv_clean.shape[1] > 0 else y_test_inv_clean.flatten()
        y_pred_plot_main = y_pred_inv_clean[:, 0] if y_pred_inv_clean.ndim > 1 and y_pred_inv_clean.shape[1] > 0 else y_pred_inv_clean.flatten()
        num_points_24h = min(self.CONST.PRED_LENGTH, len(y_test_plot_main))
        if num_points_24h > 0:
            plt.figure(figsize=(12, 6))
            plt.plot(y_test_plot_main[:num_points_24h], label="Actual SO2 (Test)", marker=".")
            plt.plot(
                y_pred_plot_main[:num_points_24h],
                label="Predicted SO2 (Test)",
                linestyle="--",
                marker="x",
            )
            plt.title(f"Actual vs. Predicted SO2 ({report_title_suffix}, First {num_points_24h}h Test) - {english_file_name}")
            plt.xlabel("Time Step (Hours)")
            plt.ylabel("SO2 Value")
            plt.legend()
            ts_plot_path_24h = self.output_dir / f"timeseries_pred_24h_{model_suffix}_{english_file_name}.png"
            plt.savefig(ts_plot_path_24h, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                ts_plot_path_24h,
                f"Actual vs. Predicted (First {num_points_24h}h Test, {report_title_suffix})",
                f"Comparison for first {num_points_24h}h of test set.",
            )

        time_frames_scatter = {"Test Set (All Cleaned Data)": (y_test_plot_main, y_pred_plot_main)}
        if num_points_24h > 1:
            time_frames_scatter[f"Test Set (First {num_points_24h} Hours)"] = (
                y_test_plot_main[:num_points_24h],
                y_pred_plot_main[:num_points_24h],
            )
        for frame_name, (
            y_actual_scatter,
            y_pred_scatter,
        ) in time_frames_scatter.items():
            if len(y_actual_scatter) > 1:
                plt.figure(figsize=(8, 8))
                plt.scatter(
                    y_actual_scatter,
                    y_pred_scatter,
                    alpha=0.6,
                    label="Actual vs. Predicted",
                )
                min_val = min(y_actual_scatter.min(), y_pred_scatter.min()) if len(y_actual_scatter) > 0 else 0
                max_val = max(y_actual_scatter.max(), y_pred_scatter.max()) if len(y_actual_scatter) > 0 else 1
                plt.plot(
                    [min_val, max_val],
                    [min_val, max_val],
                    "k--",
                    lw=2,
                    label="Ideal y=x Line",
                )
                if len(np.unique(y_actual_scatter)) > 1:
                    reg = LinearRegression().fit(y_actual_scatter.reshape(-1, 1), y_pred_scatter)
                    r2_scat = reg.score(y_actual_scatter.reshape(-1, 1), y_pred_scatter)
                    plt.plot(
                        y_actual_scatter,
                        reg.predict(y_actual_scatter.reshape(-1, 1)),
                        color="red",
                        label=f"Fit (R²={r2_scat:.2f})",
                    )
                plt.xlabel("Actual SO2")
                plt.ylabel("Predicted SO2")
                plt.title(f"Actual vs. Predicted Scatter ({frame_name}, {report_title_suffix}) - {english_file_name}")
                plt.legend()
                plt.grid(True)
                scatter_fname_suffix = frame_name.lower().replace(" ", "_").replace("(", "").replace(")", "")
                scatter_plot_path = self.output_dir / f"scatter_{scatter_fname_suffix}_{model_suffix}_{english_file_name}.png"
                plt.savefig(scatter_plot_path, dpi=300, bbox_inches="tight")
                plt.close()
                add_figure_to_doc(
                    doc,
                    scatter_plot_path,
                    f"Scatter Plot ({frame_name}, {report_title_suffix})",
                    f"Scatter plot for {frame_name}.",
                )
            else:
                doc.add_paragraph(f"Not enough data for scatter plot ({frame_name}, {report_title_suffix}).")

        doc.add_heading(f"Performance Metrics ({report_title_suffix})", 2)
        if not all(np.isnan(v) for v in scores.values()):
            metrics_df = pd.DataFrame([scores]).round(4)
            add_table_to_doc(
                doc,
                metrics_df,
                f"Metrics Summary ({report_title_suffix})",
                "Performance metrics on test set.",
            )
        else:
            doc.add_paragraph(f"Metrics could not be calculated for {report_title_suffix}.")
        print(f"DEBUG: _train_evaluate_report_lstm - Finished for model_suffix: {model_suffix}")
        return model, scores, y_test_inv_clean, y_pred_inv_clean

    def _perform_iterative_forecast(
        self,
        doc,
        model,
        df_scaled_input,
        features_to_use,
        scaler_target,
        model_suffix,
        english_file_name,
        file_name_base,
        original_df_for_time_ref,
        horizon_days,
    ):
        print(f"DEBUG: _perform_iterative_forecast - Starting for {model_suffix}, horizon: {horizon_days} days.")
        doc.add_heading(
            f"Iterative Forecast ({horizon_days} Days - {model_suffix}) - {file_name_base}",
            2,
        )
        if model is None or df_scaled_input.empty or not features_to_use or scaler_target is None:
            doc.add_paragraph("Cannot perform iterative forecast: missing model, data, features, or scaler.")
            print("DEBUG: _perform_iterative_forecast - Aborting due to missing components.")
            return None, None

        last_known_sequence_scaled_X_only, _ = prepare_sequences(df_scaled_input, features_to_use, self.CONST.SEQ_LENGTH, 0)
        if last_known_sequence_scaled_X_only.size == 0:
            doc.add_paragraph("Could not get last known sequence for iterative forecast.")
            print("DEBUG: _perform_iterative_forecast - last_known_sequence_scaled_X_only is empty.")
            return None, None

        current_sequence_X = last_known_sequence_scaled_X_only[-1].reshape(1, self.CONST.SEQ_LENGTH, len(features_to_use))
        all_predictions_scaled_target_only = []  # Stores only the PRED_LENGTH SO2 predictions
        num_prediction_steps = horizon_days * 24 // self.CONST.PRED_LENGTH

        print(f"DEBUG: _perform_iterative_forecast - num_prediction_steps: {num_prediction_steps}")
        for step_idx in range(num_prediction_steps):
            if step_idx % 50 == 0:
                print(f"DEBUG: Iterative forecast step {step_idx}/{num_prediction_steps} for {model_suffix}")  # Progress for long forecasts

            pred_block_scaled_target = model.predict(current_sequence_X)  # Shape (1, PRED_LENGTH) - this is predicted SO2
            all_predictions_scaled_target_only.append(pred_block_scaled_target[0])

            # Construct the next input sequence current_sequence_X
            # This is the most complex part for multi-feature inputs where features are not just lagged target.
            # For now, we assume features_to_use might include lagged SO2.
            # The simplest way to roll the window if features are only lags of target and exogenous variables (which we assume constant or perfectly known - a big assumption)

            # Create PRED_LENGTH new feature vectors for the next sequence input
            new_feature_vectors_for_input = []
            # Start with the last (SEQ_LENGTH - PRED_LENGTH) feature vectors from current_sequence_X
            if self.CONST.SEQ_LENGTH > self.CONST.PRED_LENGTH:
                base_vectors = current_sequence_X[0, self.CONST.PRED_LENGTH :, :].tolist()  # List of arrays
            else:  # if PRED_LENGTH >= SEQ_LENGTH, we start fresh for each block, less ideal
                base_vectors = []

            # The `pred_block_scaled_target` (shape PRED_LENGTH) needs to be incorporated into the
            # feature vectors for the *next* `PRED_LENGTH` time steps.
            # This assumes that one of the `features_to_use` is the (scaled) SO2 value itself,
            # or a lag of it that we are now "filling in" with predictions.

            # Identify which column in `features_to_use` corresponds to 'SO2' or its relevant lag
            # This is crucial and depends on your feature engineering.
            # If 'SO2' (scaled target) is NOT directly one of the input features, this becomes harder.
            # Let's assume for simplicity that 'SO2' (scaled) IS the feature we are updating or that the model handles this internally by design.
            # A common strategy if you have exogenous features: keep them as they were in the last step of `current_sequence_X`,
            # and only update the feature(s) that correspond to the target variable.

            # For each of the PRED_LENGTH steps we just predicted:
            last_known_full_feature_vector = current_sequence_X[0, -1, :].copy()  # Last vector from previous input

            for i in range(self.CONST.PRED_LENGTH):
                # Create a new feature vector for time t+i
                new_vec = last_known_full_feature_vector.copy()  # Start with a copy of the last known features

                # Find which feature in `features_to_use` is 'SO2' (or the primary target lag)
                # This is a placeholder. You MUST define how `pred_block_scaled_target[0, i]` maps to your input features.
                # If your model expects, e.g., 'SO2_lag_1' as a feature, and you just predicted SO2 for t, then this becomes SO2_lag_1 for t+1.
                # This logic is highly dependent on `features_to_use`.
                # For a simple univariate (SO2 predicting SO2) or where SO2 is the main varying input:
                if "SO2" in features_to_use:  # Assuming 'SO2' is a feature name in your scaled data fed to `prepare_sequences`
                    so2_feature_idx = features_to_use.index("SO2")
                    new_vec[so2_feature_idx] = pred_block_scaled_target[0, i]
                elif len(features_to_use) == 1:  # If only one feature, assume it's SO2 or its derivative
                    new_vec[0] = pred_block_scaled_target[0, i]
                else:
                    # This is where it gets tricky for multi-variate.
                    # A robust solution requires knowing which feature is the target's lag.
                    # Fallback: if 'SO2' is not in features_to_use, but it's the only thing predicted,
                    # this iterative forecast might be conceptually flawed for multi-feature inputs
                    # unless the other features are static or have their own forecast.
                    # For now, let's assume the *first* feature is the one to update if 'SO2' isn't named. This is a guess.
                    # print(f"Warning: Iterative forecast updating first feature as SO2 proxy. Feature names: {features_to_use}")
                    # new_vec[0] = pred_block_scaled_target[0, i] # Risky assumption
                    pass  # If SO2 is not an explicit input feature, we just roll existing features. Model must be designed for this.

                base_vectors.append(new_vec)

            if len(base_vectors) != self.CONST.SEQ_LENGTH:
                print(f"DEBUG: Iterative forecast - Error in reconstructing sequence. Expected length {self.CONST.SEQ_LENGTH}, got {len(base_vectors)}. Stopping.")
                break  # Stop if sequence reconstruction fails

            current_sequence_X = np.array(base_vectors).reshape(1, self.CONST.SEQ_LENGTH, len(features_to_use))

        if not all_predictions_scaled_target_only:
            doc.add_paragraph("No predictions made in iterative forecast loop.")
            print("DEBUG: _perform_iterative_forecast - all_predictions_scaled_target_only is empty.")
            return None, None

        all_predictions_scaled_np = np.concatenate(all_predictions_scaled_target_only, axis=0)
        all_predictions_inv = scaler_target.inverse_transform(all_predictions_scaled_np.reshape(-1, 1)).flatten()
        forecast_time_index = None
        if isinstance(original_df_for_time_ref.index, pd.DatetimeIndex) and not original_df_for_time_ref.empty:
            last_known_time = original_df_for_time_ref.index[-1]
            forecast_time_index = pd.date_range(
                start=last_known_time + pd.Timedelta(hours=1),
                periods=len(all_predictions_inv),
                freq="h",
            )  # Changed to 'h'
        else:
            forecast_time_index = np.arange(len(all_predictions_inv))

        plt.figure(figsize=(15, 7))
        if isinstance(original_df_for_time_ref.index, pd.DatetimeIndex) and "SO2" in original_df_for_time_ref:
            recent_actual_len = min(len(original_df_for_time_ref), 24 * 30)  # Last 30 days actual
            plt.plot(
                original_df_for_time_ref.index[-recent_actual_len:],
                original_df_for_time_ref["SO2"].iloc[-recent_actual_len:],
                label="Recent Actual SO2",
            )
        plt.plot(
            forecast_time_index,
            all_predictions_inv,
            label=f"Iterative Forecast ({horizon_days} Days)",
            linestyle="--",
            color="red",
        )
        plt.title(f"{horizon_days}-Day Iterative SO2 Forecast ({model_suffix}) - {english_file_name}")
        plt.xlabel("Time")
        plt.ylabel("SO2 Value")
        plt.legend()
        plt.grid(True)
        if isinstance(forecast_time_index, pd.DatetimeIndex):
            plt.xticks(rotation=30, ha="right")
        iter_forecast_plot_path = self.output_dir / f"iter_forecast_{horizon_days}d_{model_suffix}_{english_file_name}.png"
        plt.savefig(iter_forecast_plot_path, dpi=300, bbox_inches="tight")
        plt.close()
        add_figure_to_doc(
            doc,
            iter_forecast_plot_path,
            f"{horizon_days}-Day Iterative Forecast ({model_suffix})",
            f"Iterative SO2 forecast for {horizon_days} days ({model_suffix}). High uncertainty for long forecasts.",
        )
        print(f"DEBUG: _perform_iterative_forecast - Finished for {model_suffix}.")
        return forecast_time_index, all_predictions_inv

    def process_file(self, file_path):
        print(f"\n--- Processing file: {file_path.name} ---")
        file_name_base = file_path.stem
        english_file_name = self.file_name_mapping.get(file_name_base, file_name_base)
        doc = Document()
        set_default_font(doc)
        doc.add_heading(f"LSTM Model Report for SO2 Prediction - {file_name_base}", 0)
        doc.add_paragraph(f"Dataset: {file_path.name}")

        try:
            print(f"DEBUG: Loading CSV: {file_path.name}")
            load_start_time = time.time()
            df_original = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
            load_end_time = time.time()
            print(f"DEBUG: CSV loaded. Shape: {df_original.shape}. Time: {load_end_time - load_start_time:.2f}s")
        except Exception as e:
            doc.add_paragraph(f"Error loading CSV {file_path.name}: {e}")
            report_filename = self.output_dir / f"report_{file_name_base}_load_error.docx"
            doc.save(report_filename)
            print(f"Failed to load {file_path.name}, report saved: {report_filename}")
            return

        # Convert full name column to short name for xxT file
        """Datetime,CO,NOX,NO,NO2,SO2,O3,H2S,PM10,Wind_speed,Wind_dir,Temp,Rel_hum,Glob_rad,Net_rad,Pressure,Rain"""
        set_of_columns = {
            "CO": "CO",
            "NOX": "NOx",
            "NO": "NO",
            "NO2": "NO2",
            "SO2": "SO2",
            "O3": "O3",
            "H2S": "H2S",
            "PM10": "PM10",
            "Wind_speed": "WS",
            "Wind_dir": "WD",
            "Temp": "Temp",
            "Rel_hum": "RH",
            "Glob_rad": "GlobalRad",
            "Net_rad": "NetRad",
            "Pressure": "Pressure",
            "Rain": "Rain",
        }
        df_original.rename(columns=set_of_columns, inplace=True)
        # Process all rows initially, or use MAX_ROWS_TO_PROCESS

        df = df_original.copy()

        if self.CONST.MAX_ROWS_TO_PROCESS is not None and len(df) > self.CONST.MAX_ROWS_TO_PROCESS:
            print(f"DEBUG: Truncating DataFrame to MAX_ROWS_TO_PROCESS: {self.CONST.MAX_ROWS_TO_PROCESS}")
            df = df.head(self.CONST.MAX_ROWS_TO_PROCESS)

        if "SO2" not in df.columns:
            doc.add_paragraph(f"Critical: SO2 column not found in {file_path.name}. Processing cannot continue.")
            report_filename = self.output_dir / f"report_{file_name_base}_no_SO2.docx"
            set_default_font(doc)
            doc.save(report_filename)
            return
        df["SO2"] = df["SO2"].clip(lower=0)
        if "dev" in df.columns:
            df = df.drop(columns=["dev"])

        self._perform_full_eda(doc, df.copy(), "Before Any Processing", english_file_name, file_name_base)
        df_after_initial_cleaning = df.copy()
        doc.add_heading("Data Cleaning Steps", level=1)
        missing_pct = df_after_initial_cleaning.isnull().mean()

        # Valid range check
        # Add data cleaning step
        print("Step 4.1: Cleaning Data...")
        # Step 2: Apply general range validation using _apply_range_validation
        print(f"DEBUG: process_file - Calling _apply_range_validation.")
        df_after_range_val, range_stats, cols_changed_by_range_val = self._apply_range_validation(df.copy())  # Pass a copy

        if cols_changed_by_range_val:
            doc.add_paragraph(f"Range validation (values outside physical limits set to NaN) affected: {', '.join(cols_changed_by_range_val)}.")
            if range_stats:
                range_log_details = " New NaNs due to range check: "
                for col_stat, count_stat in range_stats.items():
                    if count_stat > 0:
                        range_log_details += f"{col_stat}: {count_stat}; "
                doc.add_paragraph(range_log_details)
        else:
            doc.add_paragraph("Range validation performed. No columns were outside their defined valid ranges, or no ranges defined.")

        # Step 3: Apply Winsorization using _apply_winsorization
        print(f"DEBUG: process_file - Calling _apply_winsorization.")
        df_after_winsorize, considered_cols, winsorize_details = self._apply_winsorization(df_after_range_val.copy())  # Pass a copy

        if considered_cols:
            doc.add_paragraph(f"Winsorization (1st-99th percentile) considered for: {', '.join(considered_cols)}.")
            if winsorize_details:
                doc.add_paragraph(f"Winsorization affected (showing changes): {'; '.join(winsorize_details)}")
            else:
                doc.add_paragraph("Winsorization applied, no significant changes detected or columns skipped.")
        else:
            doc.add_paragraph("No numeric columns considered for Winsorization.")

        # df is now the result of these initial adjustments
        df_after_initial_cleaning = df_after_range_val  # Final df after these steps

        # Document cleaning results
        doc.add_heading(f"Data Cleaning Results - {english_file_name}", 1)
        doc.add_paragraph("Applied cleaning rules:")
        valid_range = DynamicInputDataset.CONST.valid_ranges
        for param, (min_val, max_val) in valid_range.items():
            doc.add_paragraph(f"- {param}: {min_val} to {max_val}")

        features_to_drop_missing = [col for col in missing_pct[missing_pct > 0.5].index if col != "SO2"]
        if features_to_drop_missing:
            df_after_initial_cleaning = df_after_initial_cleaning.drop(columns=features_to_drop_missing)
            doc.add_paragraph(f"Dropped features >50% missing: {', '.join(features_to_drop_missing)}. New shape: {df_after_initial_cleaning.shape}")
        else:
            doc.add_paragraph("No features >50% missing to drop.")
        self._perform_full_eda(
            doc,
            df_after_initial_cleaning.copy(),
            "After Removing High Missing Features",
            english_file_name,
            file_name_base,
        )

        df_after_so2_trim = df_after_initial_cleaning.copy()
        if "SO2" in df_after_so2_trim.columns and df_after_so2_trim["SO2"].notna().any():
            first_valid_idx = df_after_so2_trim["SO2"].first_valid_index()
            last_valid_idx = df_after_so2_trim["SO2"].last_valid_index()
            if first_valid_idx is not None and last_valid_idx is not None:
                df_after_so2_trim = df_after_so2_trim.loc[first_valid_idx:last_valid_idx]
                doc.add_paragraph(f"Trimmed by SO2 valid range. New shape: {df_after_so2_trim.shape}")
        self._perform_full_eda(
            doc,
            df_after_so2_trim.copy(),
            "After Trimming by SO2",
            english_file_name,
            file_name_base,
        )

        doc.add_heading("Data Summaries After Cleaning & Trimming", level=1)
        if not df_after_so2_trim.empty:
            doc.add_heading(f"Data Column Information (After Cleaning/Trimming)", level=2)
            summary_list = []
            for col in df_after_so2_trim.columns:
                col_data = df_after_so2_trim[col]
                is_numeric = pd.api.types.is_numeric_dtype(col_data)
                summary_list.append(
                    {
                        "Column": col,
                        "DataType": str(col_data.dtype),
                        "Count": col_data.count(),
                        "Unique": col_data.nunique(),
                        "Nulls": col_data.isnull().sum(),
                        "Mean": f"{col_data.mean():.2f}" if is_numeric else "N/A",
                        "Std": f"{col_data.std():.2f}" if is_numeric else "N/A",
                        "Min": f"{col_data.min():.2f}" if is_numeric else col_data.min() if not col_data.empty and not is_numeric else "N/A",
                        "Max": f"{col_data.max():.2f}" if is_numeric else col_data.max() if not col_data.empty and not is_numeric else "N/A",
                    }
                )
            summary_df = pd.DataFrame(summary_list)
            add_table_to_doc(
                doc,
                summary_df,
                "Detailed Column Information Table",
                "Summary stats post cleaning.",
            )
            numeric_df_final = df_after_so2_trim.select_dtypes(include=np.number)
            if not numeric_df_final.empty:
                stats_summary_table = numeric_df_final.describe().transpose().round(2)
                add_table_to_doc(
                    doc,
                    stats_summary_table,
                    "Statistical Summary of Numeric Features (Post Cleaning)",
                    "Descriptive stats post cleaning.",
                )
        else:
            doc.add_paragraph("DataFrame empty post cleaning. No summary tables.")

        print(f"DEBUG: Starting Data Preprocessing for LSTM for {file_name_base}. DF shape: {df_after_so2_trim.shape}")
        doc.add_heading(f"Data Preprocessing for LSTM - {file_name_base}", 1)
        df_to_process = df_after_so2_trim.copy()
        df_scaled = pd.DataFrame()
        scaler_features_global = None
        scaler_target_global = None
        features_all_scaled = []
        original_df_unscaled_for_lstm = pd.DataFrame()

        if df_to_process.empty:
            doc.add_paragraph("Data empty pre-interpolation. Skipping LSTM.")
            set_default_font(doc)
            doc.save(self.output_dir / f"report_{file_name_base}_empty_pre_interp.docx")
            return
        if not isinstance(df_to_process.index, pd.DatetimeIndex) or not df_to_process.index.is_monotonic_increasing:
            df_to_process = df_to_process.sort_index()
            doc.add_paragraph("Ensured index sorted pre-interpolation.")

        interp_start_time = time.time()
        try:
            if isinstance(df_to_process.index, pd.DatetimeIndex) and df_to_process.index.nunique() > 1:
                all_dates = pd.date_range(
                    start=df_to_process.index.min(),
                    end=df_to_process.index.max(),
                    freq="h",
                )  # Changed to 'h'
                df_reindexed = df_to_process.reindex(all_dates)
                df_interpolated = df_reindexed.interpolate(method="linear", limit_direction="both")
                doc.add_paragraph(f"Interpolated to hourly. Shape post interpolate: {df_interpolated.shape}")
            else:
                df_interpolated = df_to_process.interpolate(method="linear", limit_direction="both")
                doc.add_paragraph(f"Interpolated (no reindex). Shape post interpolate: {df_interpolated.shape}")
        except Exception as e_interp:
            doc.add_paragraph(f"Error during interpolation: {e_interp}. Using data pre-interpolation.")
            df_interpolated = df_to_process.copy()
        interp_end_time = time.time()
        print(f"DEBUG: Interpolation done. Shape: {df_interpolated.shape}. Time: {interp_end_time - interp_start_time:.2f}s")

        df_processed_for_scaling = df_interpolated.dropna(how="all")
        doc.add_paragraph(f"Shape post drop all NaNs (post-interp): {df_processed_for_scaling.shape}")
        if "SO2" not in df_processed_for_scaling.columns or df_processed_for_scaling["SO2"].isnull().all():
            doc.add_paragraph(f"Critical: SO2 missing/all NaN post-interp for {file_name_base}.")
            set_default_font(doc)
            doc.save(self.output_dir / f"report_{file_name_base}_no_SO2_post_interp.docx")
            return
        df_filled_for_scaling = df_processed_for_scaling.ffill().bfill()  # Changed from fillna(method=...)
        if df_filled_for_scaling.isnull().any().any():
            df_filled_for_scaling = df_filled_for_scaling.fillna(0)
            doc.add_paragraph("Warning: NaNs remained post ffill/bfill; filled with 0.")
        if df_filled_for_scaling.empty:
            doc.add_paragraph(f"Data empty pre-scaling for {file_name_base}. Skipping LSTM.")
            set_default_font(doc)
            doc.save(self.output_dir / f"report_{file_name_base}_empty_pre_scale.docx")
            return

        original_df_unscaled_for_lstm = df_filled_for_scaling.copy()
        df_scaled = original_df_unscaled_for_lstm.copy()
        features_all_scaled = [col for col in df_scaled.columns if col != "SO2"]
        if features_all_scaled:
            scaler_features_global = MinMaxScaler()
            df_scaled[features_all_scaled] = scaler_features_global.fit_transform(df_scaled[features_all_scaled])
        scaler_target_global = MinMaxScaler()
        df_scaled[["SO2"]] = scaler_target_global.fit_transform(df_scaled[["SO2"]])
        doc.add_paragraph(f"Data scaled. Final shape for LSTM: {df_scaled.shape}")
        print(f"DEBUG: Data scaling complete. DF_scaled shape: {df_scaled.shape}")

        model_all_features, scores_all_features, _, _ = self._train_evaluate_report_lstm(
            doc,
            df_scaled.copy(),
            features_all_scaled,
            scaler_target_global,
            "all_features",
            english_file_name,
            file_name_base,
            original_df_unscaled_for_lstm.copy(),
        )
        doc.add_heading(f"Feature Importance (SHAP via XGBoost) - {file_name_base}", 1)
        selected_features_for_round2 = []
        feature_importance_df = pd.DataFrame()
        if features_all_scaled and not df_scaled[features_all_scaled].empty:
            X_flat_shap = df_scaled[features_all_scaled].values
            y_flat_shap = df_scaled["SO2"].values
            if X_flat_shap.shape[0] > 1 and y_flat_shap.shape[0] > 1 and X_flat_shap.shape[0] == y_flat_shap.shape[0]:
                try:
                    print(f"DEBUG: Running SHAP for {file_name_base}. X_flat_shap shape: {X_flat_shap.shape}")
                    shap_start_time = time.time()
                    xgb_model = xgboost.XGBRegressor(
                        objective="reg:squarederror",
                        random_state=self.CONST.SEED_VALUE,
                        n_estimators=100,
                    )
                    xgb_model.fit(X_flat_shap, y_flat_shap)
                    sample_size_shap = min(200, X_flat_shap.shape[0])
                    X_sample_shap = (
                        shap.sample(
                            X_flat_shap,
                            sample_size_shap,
                            random_state=self.CONST.SEED_VALUE,
                        )
                        if X_flat_shap.shape[0] > sample_size_shap
                        else X_flat_shap
                    )
                    if X_sample_shap.shape[0] > 0:
                        explainer = shap.TreeExplainer(xgb_model)
                        shap_values_sample = explainer.shap_values(X_sample_shap)
                        plt.figure()
                        shap.summary_plot(
                            shap_values_sample,
                            X_sample_shap,
                            feature_names=features_all_scaled,
                            plot_type="bar",
                            show=False,
                        )
                        plt.title(f"SHAP Feature Importance (Bar) - {english_file_name}")
                        plt.tight_layout()
                        shap_bar_path = self.output_dir / f"shap_bar_{english_file_name}.png"
                        plt.savefig(shap_bar_path, dpi=300, bbox_inches="tight")
                        plt.close()
                        add_figure_to_doc(
                            doc,
                            shap_bar_path,
                            "SHAP Feature Importance (Bar Plot)",
                            "Mean absolute SHAP values.",
                        )
                        plt.figure()
                        shap.summary_plot(
                            shap_values_sample,
                            X_sample_shap,
                            feature_names=features_all_scaled,
                            show=False,
                        )
                        plt.title(f"SHAP Feature Importance (Beeswarm) - {english_file_name}")
                        plt.tight_layout()
                        shap_beeswarm_path = self.output_dir / f"shap_beeswarm_{english_file_name}.png"
                        plt.savefig(shap_beeswarm_path, dpi=300, bbox_inches="tight")
                        plt.close()
                        add_figure_to_doc(
                            doc,
                            shap_beeswarm_path,
                            "SHAP Feature Importance (Beeswarm Plot)",
                            "Distribution of SHAP values.",
                        )
                        mean_abs_shap = np.abs(shap_values_sample).mean(axis=0)
                        feature_importance_df = (
                            pd.DataFrame(
                                {
                                    "Feature": features_all_scaled,
                                    "Mean_Abs_SHAP_Value": mean_abs_shap,
                                }
                            )
                            .sort_values(by="Mean_Abs_SHAP_Value", ascending=False)
                            .reset_index(drop=True)
                        )
                        add_table_to_doc(
                            doc,
                            feature_importance_df.head(10),
                            "Top 10 SHAP Feature Importance Ranking",
                            "Features ranked by mean abs SHAP.",
                        )
                        selected_features_for_round2 = feature_importance_df.head(self.CONST.NUM_TOP_FEATURES_FOR_SELECTION)["Feature"].tolist()
                        doc.add_paragraph(f"Selected features for round 2 (SHAP Top {self.CONST.NUM_TOP_FEATURES_FOR_SELECTION}): {', '.join(selected_features_for_round2) if selected_features_for_round2 else 'None'}")
                    else:
                        doc.add_paragraph("SHAP sample empty. Skipping SHAP plots.")
                    shap_end_time = time.time()
                    print(f"DEBUG: SHAP analysis done. Time: {shap_end_time - shap_start_time:.2f}s")
                except Exception as e:
                    doc.add_paragraph(f"Error during SHAP: {type(e).__name__} - {e}")
                    print(f"Error in SHAP for {file_name_base}: {type(e).__name__} - {e}")
            else:
                doc.add_paragraph("Not enough data/mismatched shapes for SHAP.")
        else:
            doc.add_paragraph("Skipping SHAP: No features or scaled feature data empty.")

        model_selected_features = None
        scores_selected_features = {}
        if selected_features_for_round2:
            model_selected_features, scores_selected_features, _, _ = self._train_evaluate_report_lstm(
                doc,
                df_scaled.copy(),
                selected_features_for_round2,
                scaler_target_global,
                "selected_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm.copy(),
            )
        else:
            doc.add_heading(f"LSTM (Selected Features) - {file_name_base}", 1)
            doc.add_paragraph("No features selected by SHAP. Skipping round 2.")

        doc.add_heading(
            f"Final {self.CONST.PRED_LENGTH}-Hour Prediction Comparison - {file_name_base}",
            1,
        )
        if not original_df_unscaled_for_lstm.empty and "SO2" in original_df_unscaled_for_lstm.columns and len(original_df_unscaled_for_lstm) >= self.CONST.PRED_LENGTH:
            actual_last_24_inv = original_df_unscaled_for_lstm["SO2"].iloc[-self.CONST.PRED_LENGTH :].values
            time_index_actual = original_df_unscaled_for_lstm.index[-self.CONST.PRED_LENGTH :]
            if isinstance(original_df_unscaled_for_lstm.index, pd.DatetimeIndex):
                future_timestamps_24h = pd.date_range(
                    start=original_df_unscaled_for_lstm.index[-1] + pd.Timedelta(hours=1),
                    periods=self.CONST.PRED_LENGTH,
                    freq="h",
                )
            else:
                future_timestamps_24h = np.arange(len(actual_last_24_inv))
            plt.figure(figsize=(14, 7))
            plt.plot(
                time_index_actual,
                actual_last_24_inv,
                label="Actual SO2 (Last Known)",
                marker="o",
                linewidth=2,
            )
            if model_all_features and features_all_scaled:
                X_full_all_for_pred, _ = prepare_sequences(df_scaled, features_all_scaled, self.CONST.SEQ_LENGTH, 0)
                if X_full_all_for_pred.size > 0:
                    pred_next_24_all_inv = scaler_target_global.inverse_transform(np.nan_to_num(model_all_features.predict(X_full_all_for_pred[-1:])))
                    plt.plot(
                        future_timestamps_24h,
                        pred_next_24_all_inv[0],
                        label="Predicted (All Features)",
                        linestyle="--",
                        marker="x",
                    )
            if model_selected_features and selected_features_for_round2:
                X_full_sel_for_pred, _ = prepare_sequences(df_scaled, selected_features_for_round2, self.CONST.SEQ_LENGTH, 0)
                if X_full_sel_for_pred.size > 0:
                    pred_next_24_sel_inv = scaler_target_global.inverse_transform(np.nan_to_num(model_selected_features.predict(X_full_sel_for_pred[-1:])))
                    plt.plot(
                        future_timestamps_24h,
                        pred_next_24_sel_inv[0],
                        label="Predicted (Selected Features)",
                        linestyle=":",
                        marker="s",
                    )
            plt.title(f"Next {self.CONST.PRED_LENGTH}-Hour SO2 Prediction Comparison - {english_file_name}")
            plt.xlabel("Time")
            plt.ylabel("SO2 Value")
            plt.legend()
            plt.grid(True)
            if isinstance(time_index_actual, pd.DatetimeIndex):
                plt.xticks(rotation=30, ha="right")
            final_pred_comp_path = self.output_dir / f"final_pred_24h_comparison_{english_file_name}.png"
            plt.savefig(final_pred_comp_path, dpi=300, bbox_inches="tight")
            plt.close()
            add_figure_to_doc(
                doc,
                final_pred_comp_path,
                f"Next {self.CONST.PRED_LENGTH}-Hour Prediction Comparison",
                "Comparison of 24h predictions.",
            )
        else:
            doc.add_paragraph(f"Not enough data for final {self.CONST.PRED_LENGTH}h prediction plot.")

        if model_all_features:
            self._perform_iterative_forecast(
                doc,
                model_all_features,
                df_scaled.copy(),
                features_all_scaled,
                scaler_target_global,
                "all_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm.copy(),
                self.CONST.ITERATIVE_PRED_HORIZON_DAYS,
            )
        if model_selected_features:
            self._perform_iterative_forecast(
                doc,
                model_selected_features,
                df_scaled.copy(),
                selected_features_for_round2,
                scaler_target_global,
                "selected_features",
                english_file_name,
                file_name_base,
                original_df_unscaled_for_lstm.copy(),
                self.CONST.ITERATIVE_PRED_HORIZON_DAYS,
            )

        doc.add_heading("Model Performance Comparison (All Features vs. SHAP Selected Features)", 1)
        if scores_all_features and scores_selected_features and scores_all_features.get("r2") is not np.nan and scores_selected_features.get("r2") is not np.nan:  # Check if scores are valid
            comp_data = {
                "Metric": ["MSE", "MAE", "R2 Score"],
                "All Features": [
                    scores_all_features.get("mse", np.nan),
                    scores_all_features.get("mae", np.nan),
                    scores_all_features.get("r2", np.nan),
                ],
                "Selected Features (SHAP)": [
                    scores_selected_features.get("mse", np.nan),
                    scores_selected_features.get("mae", np.nan),
                    scores_selected_features.get("r2", np.nan),
                ],
            }
            comp_df = pd.DataFrame(comp_data).round(4)
            add_table_to_doc(doc, comp_df, "Metrics Comparison Table", "Comparison of model metrics.")
            r2_all = scores_all_features.get("r2", -1)
            r2_selected = scores_selected_features.get("r2", -1)
            summary_text = "Comparing model performance:\n"
            if r2_all > r2_selected and abs(r2_all - r2_selected) > 0.01:
                summary_text += f"- ALL features model better (R2: {r2_all:.4f}) than SHAP-selected (R2: {r2_selected:.4f}).\n"
            elif r2_selected > r2_all and abs(r2_selected - r2_all) > 0.01:
                summary_text += f"- SHAP-SELECTED features model better (R2: {r2_selected:.4f}) than all features (R2: {r2_all:.4f}). SHAP identified impactful features.\n"
            else:
                summary_text += f"- Both models (All R2: {r2_all:.4f}, Selected R2: {r2_selected:.4f}) similar performance or one/both not reliably scored.\n"
            mae_all = scores_all_features.get("mae", np.inf)
            mae_selected = scores_selected_features.get("mae", np.inf)
            if mae_selected < mae_all and abs(mae_all - mae_selected) > 0.01 * mae_all:
                summary_text += f"- Selected Features model lower MAE ({mae_selected:.4f} vs {mae_all:.4f}), smaller avg errors.\n"
            elif mae_all < mae_selected and abs(mae_selected - mae_all) > 0.01 * mae_selected:
                summary_text += f"- All Features model lower MAE ({mae_all:.4f} vs {mae_selected:.4f}).\n"
            doc.add_paragraph(summary_text)
        elif scores_all_features:
            doc.add_paragraph("Only 'All Features' model evaluated. No comparison.")
        else:
            doc.add_paragraph("Neither model fully evaluated for comparison.")

        print(f"Finalizing and saving report for {file_name_base}...")
        set_default_font(doc)
        report_filename = self.output_dir / f"report_{file_name_base}.docx"
        try:
            doc.save(report_filename)
            print(f"Report saved: {report_filename}")
        except Exception as e:
            print(f"Error saving report {report_filename}: {e}")
            try:
                fallback_report_filename = self.output_dir / f"report_{file_name_base}_fallback.docx"
                doc.save(fallback_report_filename)
                print(f"Fallback report saved: {fallback_report_filename}")
            except Exception as e_fb:
                print(f"Error saving fallback report: {e_fb}")

        del (
            df,
            df_original,
            df_after_initial_cleaning,
            df_after_so2_trim,
            df_scaled,
            original_df_unscaled_for_lstm,
        )
        if "model_all_features" in locals():
            del model_all_features
        if "model_selected_features" in locals():
            del model_selected_features
        if "xgb_model" in locals():
            del xgb_model
        gc.collect()
        tf.keras.backend.clear_session()
        print(f"--- Finished processing file: {file_path.name} ---")

    def run(self):
        for csv_file_path in self.csv_files:
            try:
                self.process_file(csv_file_path)
            except Exception as e:
                print(f"--- CRITICAL ERROR processing file {csv_file_path.name}: {type(e).__name__} - {e} ---")
                import traceback

                traceback.print_exc()
            finally:
                gc.collect()
                tf.keras.backend.clear_session()


if __name__ == "__main__":
    input_directory = "/mnt/e/MayThesis2025/src/labs4/s02_v01_02_clean/output-680507-2319"
    input_directory = "/mnt/e/MayThesis2025/cleanned_datasets"
    if not os.path.exists(input_directory):
        print(f"Error: Input directory does not exist: {input_directory}")
    else:
        processor = DynamicInputDataset(input_directory)
        processor.run()
