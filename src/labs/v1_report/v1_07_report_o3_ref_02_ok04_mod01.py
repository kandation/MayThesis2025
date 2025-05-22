# Import Libraries
print("Step 1: Importing Libraries...")
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
import xgboost
import shap
from docx import Document
from docx.shared import Inches
import random
import os
from pathlib import Path
import gc

from docx.oxml.ns import qn


# Do not change these functions:
def set_thsarabupsk_font(doc, font_name="TH SarabunPSK"):
    """Sets the specified font for all text in a Word document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font._element.rPr.rFonts.set(qn("w:ascii"), font_name)
            font._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
            font._element.rPr.rFonts.set(qn("w:cs"), font_name)
            font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            font._element.rPr.rFonts.set(qn("w:complexScript"), font_name)
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    font._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    font._element.rPr.rFonts.set(qn("w:cs"), font_name)
    font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    font._element.rPr.rFonts.set(qn("w:complexScript"), font_name)


def add_table(doc, df, title, description=""):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
    table.style.border_width = 1
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
    if description:
        doc.add_paragraph(description)


def add_figure(doc, filename, title, description=""):
    if isinstance(filename, Path):
        filename = str(filename)
    if os.path.exists(filename):
        doc.add_heading(title, level=2)
        doc.add_picture(filename, width=Inches(6))
        if description:
            doc.add_paragraph(description)
    else:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Warning: Image file {filename} not found.")


def prepare_sequences(data, features, seq_length, pred_length):
    X, y = [], []
    for i in range(len(data) - seq_length - pred_length + 1):
        X.append(data[features].iloc[i : i + seq_length].values)
        y.append(data["SO2"].iloc[i + seq_length : i + seq_length + pred_length].values)
    return np.array(X), np.array(y)


# Wrap everything into a class with a constant configuration class.
class DynamicInputDataset:
    # Constant configuration class for easy to config
    class CONST:
        SEED_VALUE = 42
        BATCH_SIZE = 1024
        EPOCHS = 50
        N_SPLITS = 5
        SEQ_LENGTH = 24
        PRED_LENGTH = 24
        OUTPUT_FOLDER = "outputs"

    # File name mapping (unchanged)
    file_name_mapping = {
        "(37t)ศาลหลักเมือง(ปิดสถานี)": "San_Lak_Mueang_Closed_Station",
        "(37t)สถานีอุตุนิยมวิทยาลำปาง": "Lampang_Meteorological_Station",
        "(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด": "Ban_Sopad_Hospital",
        "(39t)โรงพยาบาลส่งเสริมสุขภาพตำบลท่าสี": "Tha_Si_Hospital",
        "(40t)การประปาส่วนภูมิภาคแม่เมาะ": "Mae_Mo_Waterworks",
    }

    def __init__(self, input_dir):
        self.input_dir = Path(input_dir)
        self.csv_files = list(self.input_dir.glob("*.csv"))
        # Create output directory
        self.output_dir = Path.cwd() / self.CONST.OUTPUT_FOLDER
        self.output_dir.mkdir(exist_ok=True)
        # Set random seeds for reproducibility
        random.seed(self.CONST.SEED_VALUE)
        np.random.seed(self.CONST.SEED_VALUE)
        tf.random.set_seed(self.CONST.SEED_VALUE)
        # Configure GPU memory growth for LSTM
        gpus = tf.config.experimental.list_physical_devices("GPU")
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
                print("GPU is set up and ready for LSTM!")
            except RuntimeError as e:
                print(e)
        print("Num GPUs Available: ", len(tf.config.list_physical_devices("GPU")))

    def plot_tscv(self, dataset_size):
        # plot hbar sorted by index ascending
        df = pd.DataFrame(dataset_size)
        df = df.sort_index(ascending=False)
        df.plot(kind="barh", stacked=True)

        # show value on the bar and show total value
        for i in range(len(df)):
            total_text = df.iloc[i]["train"] + df.iloc[i]["validation"]
            train_text = df.iloc[i]["train"]
            train_percentage = train_text / total_text * 100
            validation_text = df.iloc[i]["validation"]
            validation_percentage = validation_text / total_text * 100
            plt.text(
                total_text,
                i,
                f"{total_text}",
                ha="left",
                va="center",
                color="black",
                fontsize=7,
            )
            plt.text(
                train_text / 2,
                i,
                f"{train_text} ({train_percentage:.2f}%)",
                ha="center",
                va="center",
                color="white",
                fontsize=7,
            )
            plt.text(
                train_text + validation_text / 2,
                i,
                f"{validation_text} ({validation_percentage:.2f}%)",
                ha="center",
                va="center",
                color="black",
                fontsize=7,
            )

        plt.xlabel("Dataset Size")
        plt.ylabel("Split (Fold)")
        plt.title("Time Series Cross Validation")
        plt.legend(["Train", "Validation"], loc="upper right")
        plt.savefig(
            "temp_fig.png",
            dpi=300,
            bbox_inches="tight",
            facecolor="white",
            transparent=False,
        )
        plt.show()
        plt.close()

    def process_file(self, file_path):
        print(f"Processing file: {file_path}")
        # Step 4: Loading Data using dynamic input
        print("Step 4: Loading Data...")
        df = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
        # remove dev if exist
        if "dev" in df.columns:
            df = df.drop(columns=["dev"])
        df_raw = df.copy()  # Keep raw data for analysis
        file_name = os.path.basename(file_path).replace(".csv", "")
        english_file_name = self.file_name_mapping.get(file_name, file_name)
        if all([char >= "ก" and char <= "๙" for char in file_name]):
            raise ValueError("Invalid filename: Must be in Thai characters [ก-๙]")

        doc = Document()
        doc.add_heading(
            f"LSTM Model Report for SO2 Prediction with XGBoost SHAP FI - {file_name}",
            0,
        )
        doc.add_paragraph(f"Dataset: {file_path} (Limited to 1000 rows)")

        print("Step 5: Showing Data Before Any Processing...")
        doc.add_heading(f"Data Before Any Processing ({file_name})", 1)
        doc.add_paragraph(f"Shape Before: {df.shape}")
        doc.add_paragraph(f"Missing Values Total Before: {df.isnull().sum().sum()}")

        print("Step 5.1: Listing Columns (Before)...")
        doc.add_heading(f"Columns List (Before) - {file_name}", 2)
        doc.add_paragraph(
            f"List of columns in the dataset before processing ({file_name}):"
        )
        for col in df.columns:
            doc.add_paragraph(f"- {col}")

        print("Step 5.2: Showing Data Sample (Before)...")
        df_head_before = df.head(10)
        add_table(
            doc,
            df_head_before,
            f"Sample of First 10 Rows (Before) - {file_name}",
            f"ตารางนี้แสดงตัวอย่าง 10 แถวแรกของข้อมูลดิบก่อนดำเนินการ ({file_name})",
        )

        print("Step 5.3: Computing Data Information (Before)...")
        doc.add_heading(f"Data Information (Before) - {file_name}", 2)
        info_output_before = df.info()
        doc.add_paragraph(str(info_output_before))

        print("Step 6: Performing EDA (Before Any Processing)...")
        doc.add_heading(
            f"Exploratory Data Analysis (Before Any Processing) - {file_name}", 1
        )
        doc.add_paragraph(f"Shape: {df.shape}")
        doc.add_paragraph(f"Missing Values Total: {df.isnull().sum().sum()}")

        plt.figure(figsize=(12, 6))
        heatmap = sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(f"Missing Data Heatmap (Before Any Processing) - {english_file_name}")
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_before_all = df.isnull().sum().sum()
        nan_max_col_before_all = (
            df.isnull().sum().idxmax() if nan_count_before_all > 0 else "None"
        )
        nan_max_val_before_all = (
            df.isnull().sum().max() if nan_count_before_all > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) ก่อนดำเนินการใด ๆ รวมทั้งหมด {nan_count_before_all} ค่า ({file_name}) "
        if nan_count_before_all > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_before_all}' ({nan_max_val_before_all} ค่า) "
            if nan_max_val_before_all > df.shape[0] * 0.5:
                desc += (
                    " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่ "
                )
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น "
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วง 1000 แถวนี้ "
        desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}"
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_before_all_{english_file_name}.png",
            f"Missing Data Heatmap (Before Any Processing) - {file_name}",
            desc,
        )

        corr_matrix_before = df.corr()
        correlation_df_before = pd.DataFrame(
            {
                "feature": corr_matrix_before.index,
                "correlation": corr_matrix_before["SO2"],
                "abs_correlation": corr_matrix_before["SO2"].abs(),
            }
        )
        correlation_df_before = correlation_df_before.sort_values(
            by="abs_correlation", ascending=False
        )
        add_table(
            doc,
            correlation_df_before,
            f"Sorted Correlation Matrix with SO2 (Before) - {file_name}",
            f"ตารางนี้แสดงค่าสหสัมพันธ์กับ SO2 เรียงลำดับตามค่าสัมประสิทธิ์สัมพันธ์สัมบูรณ์ก่อนดำเนินการ ({file_name})",
        )

        plt.figure(figsize=(10, 8))
        heatmap = sns.heatmap(
            corr_matrix_before,
            annot=True,
            cmap="RdBu",
            center=0,
            vmin=-1,
            vmax=1,
            fmt=".2f",
        )
        plt.title(f"Correlation Matrix (Before Any Processing) - {english_file_name}")
        plt.xticks(rotation=45, ha="right")
        plt.yticks(rotation=0)
        plt.savefig(
            self.output_dir / f"eda_corr_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        corr_max_before = df.corr()["SO2"].drop("SO2").abs().max()
        corr_max_col_before = df.corr()["SO2"].drop("SO2").abs().idxmax()
        corr_min_before = df.corr()["SO2"].drop("SO2").abs().min()
        corr_min_col_before = df.corr()["SO2"].drop("SO2").abs().idxmin()
        desc = f"กราฟนี้แสดงความสัมพันธ์ระหว่างฟีเจอร์ก่อนดำเนินการ ค่าสหสัมพันธ์สูงสุดกับ SO2 คือ {corr_max_before:.2f} จาก '{corr_max_col_before}' ({file_name}) "
        desc += (
            f"ค่าสหสัมพันธ์ต่ำสุดกับ SO2 คือ {corr_min_before:.2f} จาก '{corr_min_col_before}' "
        )
        if corr_max_before > 0.7:
            desc += "แสดงถึงความสัมพันธ์ที่สูงมาก "
        elif corr_max_before > 0.3:
            desc += "แสดงถึงความสัมพันธ์ปานกลาง "
        else:
            desc += "แสดงถึงความสัมพันธ์ที่ต่ำ "
        desc += "การกระจายของค่าสหสัมพันธ์มีแนวโน้มสมมาตรตามแนวทแยง"
        add_figure(
            doc,
            self.output_dir / f"eda_corr_before_all_{english_file_name}.png",
            f"Correlation Matrix (Before Any Processing) - {file_name}",
            desc,
        )

        plt.figure(figsize=(15, 6))
        boxplot = df.boxplot(figsize=(15, 6))
        plt.xticks(rotation=45, ha="right")
        plt.title(
            f"Feature Outlier Distribution (Before Any Processing) - {english_file_name}"
        )
        plt.savefig(
            self.output_dir / f"eda_boxplot_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        outliers_before = {}
        for col in df.columns:
            q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
            iqr = q3 - q1
            outlier_count = (
                (df[col] < (q1 - 1.5 * iqr)) | (df[col] > (q3 + 1.5 * iqr))
            ).sum()
            outliers_before[col] = outlier_count
        max_outlier_col_before = max(outliers_before, key=outliers_before.get)
        min_outlier_col_before = min(outliers_before, key=outliers_before.get)
        desc = f"กราฟนี้แสดงการกระจายและ outlier ของฟีเจอร์ก่อนดำเนินการ คอลัมน์ที่มี outlier มากที่สุดคือ '{max_outlier_col_before}' ({outliers_before[max_outlier_col_before]} ค่า) ({file_name}) "
        desc += f"คอลัมน์ที่มี outlier น้อยที่สุดคือ '{min_outlier_col_before}' ({outliers_before[min_outlier_col_before]} ค่า) "
        if outliers_before[max_outlier_col_before] > df.shape[0] * 0.1:
            desc += "ซึ่งเกิน 10% ของข้อมูล แสดงถึงการกระจายที่มีความผันผวนสูง "
        else:
            desc += "ซึ่งอยู่ในระดับต่ำ ไม่มีผลกระทบมาก "
        desc += f"ค่าสถิติของ '{max_outlier_col_before}': Mean={df[max_outlier_col_before].mean():.2f}, Std={df[max_outlier_col_before].std():.2f}, "
        desc += f"Min={df[max_outlier_col_before].min():.2f}, Max={df[max_outlier_col_before].max():.2f}"
        add_figure(
            doc,
            self.output_dir / f"eda_boxplot_before_all_{english_file_name}.png",
            f"Feature Outlier Distribution (Before Any Processing) - {file_name}",
            desc,
        )

        n_cols = 4
        n_rows = (len(df.columns) + n_cols - 1) // n_cols
        plt.figure(figsize=(15, n_rows * 3))
        for i, col in enumerate(df.columns):
            plt.subplot(n_rows, n_cols, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.hist(data, bins=50, edgecolor="black")
            else:
                plt.text(0.5, 0.5, "No valid data (all NaN)", ha="center", va="center")
            plt.title(col)
        plt.tight_layout()
        plt.savefig(
            self.output_dir / f"eda_histogram_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = (
            f"กราฟนี้แสดงการกระจายของแต่ละฟีเจอร์ก่อนดำเนินการในรูปแบบฮิสโตแกรม ({file_name}) "
        )
        skew_max_before = df.skew().abs().max()
        skew_max_col_before = df.skew().abs().idxmax()
        skew_min_before = df.skew().abs().min()
        skew_min_col_before = df.skew().abs().idxmin()
        desc += (
            f"ฟีเจอร์ที่มี skewness สูงสุดคือ '{skew_max_col_before}' ({skew_max_before:.2f}) "
        )
        desc += (
            f"ฟีเจอร์ที่มี skewness ต่ำสุดคือ '{skew_min_col_before}' ({skew_min_before:.2f}) "
        )
        if skew_max_before > 1:
            desc += "แสดงถึงการกระจายที่เบ้มาก (skewed distribution) "
        elif skew_max_before > 0.5:
            desc += "แสดงถึงการกระจายที่เบ้ปานกลาง "
        else:
            desc += "แสดงถึงการกระจายที่ใกล้เคียงปกติ "
        desc += f"ค่าสถิติของ '{skew_max_col_before}': Min={df[skew_max_col_before].min():.2f}, Max={df[skew_max_col_before].max():.2f}, "
        desc += f"Mean={df[skew_max_col_before].mean():.2f}, Std={df[skew_max_col_before].std():.2f}"
        add_figure(
            doc,
            self.output_dir / f"eda_histogram_before_all_{english_file_name}.png",
            f"Feature Histograms (Before Any Processing) - {file_name}",
            desc,
        )

        print(
            "Step 6.3: Creating Timeseries Plot for All Features (Before Any Processing)..."
        )
        n_features = len(df.columns)
        plt.figure(figsize=(15, n_features * 4))
        for i, col in enumerate(df.columns):
            plt.subplot(n_features, 1, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.plot(df.index[df[col].notna()], data, label=col)
            else:
                plt.text(0.5, 0.5, "No valid data (all NaN)", ha="center", va="center")
            plt.title(f"Timeseries of {col} (Yearly Labels) - {english_file_name}")
            plt.legend()
        plt.tight_layout()
        plt.savefig(
            self.output_dir / f"timeseries_plot_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = f"กราฟนี้แสดงข้อมูล Timeseries ของทุกฟีเจอร์ก่อนดำเนินการ โดยใช้ Subplot และแสดง Label แบบ Yearly (เช่น 2010, 2011, ...) ({file_name}) "
        desc += "ความสูงของภาพปรับตามจำนวนฟีเจอร์ (4 นิ้วต่อฟีเจอร์)"
        add_figure(
            doc,
            self.output_dir / f"timeseries_plot_before_all_{english_file_name}.png",
            f"Timeseries Plot for All Features (Before Any Processing) - {file_name}",
            desc,
        )

        print("Step 6.4: Removing Features with Missing > 50% (Excluding SO2)...")
        missing_pct = df.isnull().mean()
        features_to_drop = missing_pct[missing_pct > 0.5].index
        features_to_drop = [col for col in features_to_drop if col != "SO2"]
        if features_to_drop:
            df = df.drop(columns=features_to_drop)
            doc.add_paragraph(
                f"Dropped features (>50% missing, excluding SO2) - {file_name}: {features_to_drop}"
            )
        else:
            doc.add_paragraph(
                f"No features with missing > 50% (excluding SO2) to drop - {file_name}."
            )

        doc.add_heading(
            f"Exploratory Data Analysis (After Removing Missing > 50%, Before Trim and Trim SO2) - {file_name}",
            1,
        )
        doc.add_paragraph(
            f"Shape after removing features with missing > 50% (excluding SO2) - {file_name}: {df.shape}"
        )
        doc.add_paragraph(
            f"Missing Values Total after removing features - {file_name}: {df.isnull().sum().sum()}"
        )

        plt.figure(figsize=(12, 6))
        heatmap = sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(
            f"Missing Data Heatmap (After Removing Missing > 50%, Excluding SO2, Before Trim and Trim SO2) - {english_file_name}"
        )
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_after_remove_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_after_remove = df.isnull().sum().sum()
        nan_max_col_after_remove = (
            df.isnull().sum().idxmax() if nan_count_after_remove > 0 else "None"
        )
        nan_max_val_after_remove = (
            df.isnull().sum().max() if nan_count_after_remove > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) หลังลบฟีเจอร์ที่มี Missing > 50% (ไม่รวม SO2) แต่ก่อนตัดทอนและตัดตาม SO2 รวมทั้งหมด {nan_count_after_remove} ค่า ({file_name}) "
        if nan_count_after_remove > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_after_remove}' ({nan_max_val_after_remove} ค่า) "
            if nan_max_val_after_remove > df.shape[0] * 0.5:
                desc += (
                    " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่ "
                )
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น "
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วงนี้หลังลบฟีเจอร์ "
        desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}"
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_after_remove_{english_file_name}.png",
            f"Missing Data Heatmap (After Removing Missing > 50%, Excluding SO2, Before Trim and Trim SO2) - {file_name}",
            desc,
        )

        corr_matrix_after_remove = df.corr()
        correlation_df_after_remove = pd.DataFrame(
            {
                "feature": corr_matrix_after_remove.index,
                "correlation": corr_matrix_after_remove["SO2"],
                "abs_correlation": corr_matrix_after_remove["SO2"].abs(),
            }
        )
        correlation_df_after_remove = correlation_df_after_remove.sort_values(
            by="abs_correlation", ascending=False
        )
        add_table(
            doc,
            correlation_df_after_remove,
            f"Sorted Correlation Matrix with SO2 (After Removing Missing > 50%) - {file_name}",
            f"ตารางนี้แสดงค่าสหสัมพันธ์กับ SO2 เรียงลำดับตามค่าสัมประสิทธิ์สัมพันธ์สัมบูรณ์หลังลบฟีเจอร์ที่มี Missing > 50% (ไม่รวม SO2) ({file_name})",
        )

        plt.figure(figsize=(10, 8))
        heatmap = sns.heatmap(
            corr_matrix_after_remove,
            annot=True,
            cmap="RdBu",
            center=0,
            vmin=-1,
            vmax=1,
            fmt=".2f",
        )
        plt.title(
            f"Correlation Matrix (After Removing Missing > 50%, Before Trim and Trim SO2) - {english_file_name}"
        )
        plt.xticks(rotation=45, ha="right")
        plt.yticks(rotation=0)
        plt.savefig(
            self.output_dir / f"eda_corr_after_remove_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        corr_max_after = df.corr()["SO2"].drop("SO2").abs().max()
        corr_max_col_after = df.corr()["SO2"].drop("SO2").abs().idxmax()
        corr_min_after = df.corr()["SO2"].drop("SO2").abs().min()
        corr_min_col_after = df.corr()["SO2"].drop("SO2").abs().idxmin()
        desc = f"กราฟนี้แสดงความสัมพันธ์ระหว่างฟีเจอร์หลังลบ Missing > 50% (ไม่รวม SO2) แต่ก่อนตัดทอนและตัดตาม SO2 ค่าสหสัมพันธ์สูงสุดกับ SO2 คือ {corr_max_after:.2f} จาก '{corr_max_col_after}' ({file_name}) "
        desc += (
            f"ค่าสหสัมพันธ์ต่ำสุดกับ SO2 คือ {corr_min_after:.2f} จาก '{corr_min_col_after}' "
        )
        if corr_max_after > 0.7:
            desc += "แสดงถึงความสัมพันธ์ที่สูงมาก "
        elif corr_max_after > 0.3:
            desc += "แสดงถึงความสัมพันธ์ปานกลาง "
        else:
            desc += "แสดงถึงความสัมพันธ์ที่ต่ำ "
        desc += "การกระจายของค่าสหสัมพันธ์มีแนวโน้มสมมาตรตามแนวทแยง"
        add_figure(
            doc,
            self.output_dir / f"eda_corr_after_remove_{english_file_name}.png",
            f"Correlation Matrix (After Removing Missing > 50%, Before Trim and Trim SO2) - {file_name}",
            desc,
        )

        plt.figure(figsize=(15, 6))
        boxplot = df.boxplot(figsize=(15, 6))
        plt.xticks(rotation=45, ha="right")
        plt.title(
            f"Feature Outlier Distribution (After Removing Missing > 50%, Before Trim and Trim SO2) - {english_file_name}"
        )
        plt.savefig(
            self.output_dir / f"eda_boxplot_after_remove_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        outliers_after_remove = {}
        for col in df.columns:
            q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
            iqr = q3 - q1
            outlier_count = (
                (df[col] < (q1 - 1.5 * iqr)) | (df[col] > (q3 + 1.5 * iqr))
            ).sum()
            outliers_after_remove[col] = outlier_count
        max_outlier_col_after = max(
            outliers_after_remove, key=outliers_after_remove.get
        )
        min_outlier_col_after = min(
            outliers_after_remove, key=outliers_after_remove.get
        )
        desc = f"กราฟนี้แสดงการกระจายและ outlier ของฟีเจอร์หลังลบ Missing > 50% (ไม่รวม SO2) แต่ก่อนตัดทอนและตัดตาม SO2 คอลัมน์ที่มี outlier มากที่สุดคือ '{max_outlier_col_after}' ({outliers_after_remove[max_outlier_col_after]} ค่า) ({file_name}) "
        desc += f"คอลัมน์ที่มี outlier น้อยที่สุดคือ '{min_outlier_col_after}' ({outliers_after_remove[min_outlier_col_after]} ค่า) "
        if outliers_after_remove[max_outlier_col_after] > df.shape[0] * 0.1:
            desc += "ซึ่งเกิน 10% ของข้อมูล แสดงถึงการกระจายที่มีความผันผวนสูง "
        else:
            desc += "ซึ่งอยู่ในระดับต่ำ ไม่มีผลกระทบมาก "
        desc += f"ค่าสถิติของ '{max_outlier_col_after}': Mean={df[max_outlier_col_after].mean():.2f}, Std={df[max_outlier_col_after].std():.2f}, "
        desc += f"Min={df[max_outlier_col_after].min():.2f}, Max={df[max_outlier_col_after].max():.2f}"
        add_figure(
            doc,
            self.output_dir / f"eda_boxplot_after_remove_{english_file_name}.png",
            f"Feature Outlier Distribution (After Removing Missing > 50%, Before Trim and Trim SO2) - {file_name}",
            desc,
        )

        n_cols = 4
        n_rows = (len(df.columns) + n_cols - 1) // n_cols
        plt.figure(figsize=(15, n_rows * 3))
        for i, col in enumerate(df.columns):
            plt.subplot(n_rows, n_cols, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.hist(data, bins=50, edgecolor="black")
            else:
                plt.text(0.5, 0.5, "No valid data (all NaN)", ha="center", va="center")
            plt.title(col)
        plt.tight_layout()
        plt.savefig(
            self.output_dir / f"eda_histogram_after_remove_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = f"กราฟนี้แสดงการกระจายของแต่ละฟีเจอร์หลังลบ Missing > 50% (ไม่รวม SO2) แต่ก่อนตัดทอนและตัดตาม SO2 ในรูปแบบฮิสโตแกรม ({file_name}) "
        skew_max_after = df.skew().abs().max()
        skew_max_col_after = df.skew().abs().idxmax()
        skew_min_after = df.skew().abs().min()
        skew_min_col_after = df.skew().abs().idxmin()
        desc += (
            f"ฟีเจอร์ที่มี skewness สูงสุดคือ '{skew_max_col_after}' ({skew_max_after:.2f}) "
        )
        desc += (
            f"ฟีเจอร์ที่มี skewness ต่ำสุดคือ '{skew_min_col_after}' ({skew_min_after:.2f}) "
        )
        if skew_max_after > 1:
            desc += "แสดงถึงการกระจายที่เบ้มาก (skewed distribution) "
        elif skew_max_after > 0.5:
            desc += "แสดงถึงการกระจายที่เบ้ปานกลาง "
        else:
            desc += "แสดงถึงการกระจายที่ใกล้เคียงปกติ "
        desc += f"ค่าสถิติของ '{skew_max_col_after}': Min={df[skew_max_col_after].min():.2f}, Max={df[skew_max_col_after].max():.2f}, "
        desc += f"Mean={df[skew_max_col_after].mean():.2f}, Std={df[skew_max_col_after].std():.2f}"
        add_figure(
            doc,
            self.output_dir / f"eda_histogram_after_remove_{english_file_name}.png",
            f"Feature Histograms (After Removing Missing > 50%, Before Trim and Trim SO2) - {file_name}",
            desc,
        )

        print("Step 6.5: Trimming Head/Tails Based on SO2...")
        if "SO2" in df.columns and df["SO2"].notna().any():
            valid_so2 = df["SO2"].notna()
            if valid_so2.sum() > 0:
                df = df[valid_so2]
                doc.add_paragraph(
                    f"Trimmed head/tails based on non-NaN SO2 values - {file_name}. New shape: {df.shape}"
                )
            else:
                doc.add_paragraph(
                    f"No valid SO2 values found after removing missing > 50% - {file_name}. Using all available data."
                )
        else:
            doc.add_paragraph(
                f"SO2 column not found or all NaN after removing missing > 50% - {file_name}. Using all available data."
            )

        doc.add_heading(
            f"Exploratory Data Analysis (After Trimming Based on SO2) - {file_name}", 1
        )
        doc.add_paragraph(
            f"Shape after trimming based on SO2 - {file_name}: {df.shape}"
        )
        doc.add_paragraph(
            f"Missing Values Total after trimming - {file_name}: {df.isnull().sum().sum()}"
        )

        plt.figure(figsize=(12, 6))
        heatmap = sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(
            f"Missing Data Heatmap (After Trimming Based on SO2) - {english_file_name}"
        )
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_after_trim_so2_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_after_trim_so2 = df.isnull().sum().sum()
        nan_max_col_after_trim_so2 = (
            df.isnull().sum().idxmax() if nan_count_after_trim_so2 > 0 else "None"
        )
        nan_max_val_after_trim_so2 = (
            df.isnull().sum().max() if nan_count_after_trim_so2 > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) หลังตัดทอนตาม SO2 รวมทั้งหมด {nan_count_after_trim_so2} ค่า ({file_name}) "
        if nan_count_after_trim_so2 > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_after_trim_so2}' ({nan_max_val_after_trim_so2} ค่า) "
            if nan_max_val_after_trim_so2 > df.shape[0] * 0.5:
                desc += (
                    " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่ "
                )
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น "
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วงนี้หลังตัดทอน "
        desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}"
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_after_trim_so2_{english_file_name}.png",
            f"Missing Data Heatmap (After Trimming Based on SO2) - {file_name}",
            desc,
        )

        print(
            "Step 6.6: Showing Simple Data Sample (10 Rows) After Trimming/Remove Missing..."
        )
        df_head_after = df.head(10)
        add_table(
            doc,
            df_head_after,
            f"Sample of First 10 Rows (After Trimming/Remove Missing) - {file_name}",
            f"ตารางนี้แสดงตัวอย่าง 10 แถวแรกของข้อมูลหลังลบฟีเจอร์ที่มี Missing > 50% และตัดทอนตาม SO2 ({file_name})",
        )

        print("Step 6.7: Computing Data Information (After Trimming/Remove Missing)...")
        doc.add_heading(
            f"Data Information (After Trimming/Remove Missing) - {file_name}", 1
        )
        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        table.style.border_width = 1
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Column"
        hdr_cells[1].text = "Count"
        hdr_cells[2].text = "Mean"
        hdr_cells[3].text = "Std"
        hdr_cells[4].text = "Min"
        hdr_cells[5].text = "25%"
        hdr_cells[6].text = "50%"
        hdr_cells[7].text = "75%"
        hdr_cells[8].text = "Max"

        for col in df.columns:
            print(col)
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = str(df[col].count())
            row_cells[2].text = (
                str(df[col].mean()) if pd.api.types.is_numeric_dtype(df[col]) else "N/A"
            )
            row_cells[3].text = (
                str(df[col].std()) if pd.api.types.is_numeric_dtype(df[col]) else "N/A"
            )
            row_cells[4].text = (
                str(df[col].min()) if pd.api.types.is_numeric_dtype(df[col]) else "N/A"
            )
            row_cells[5].text = (
                str(df[col].quantile(0.25))
                if pd.api.types.is_numeric_dtype(df[col])
                else "N/A"
            )
            row_cells[6].text = (
                str(df[col].median())
                if pd.api.types.is_numeric_dtype(df[col])
                else "N/A"
            )
            row_cells[7].text = (
                str(df[col].quantile(0.75))
                if pd.api.types.is_numeric_dtype(df[col])
                else "N/A"
            )
            row_cells[8].text = (
                str(df[col].max()) if pd.api.types.is_numeric_dtype(df[col]) else "N/A"
            )
        desc = f"ตารางนี้แสดงข้อมูลพื้นฐานของข้อมูลหลังลบฟีเจอร์ที่มี Missing > 50% และตัดทอนตาม SO2 ค่าที่ไม่ใช่ตัวเลขจะแสดง N/A ({file_name})"
        doc.add_paragraph(desc)

        print(
            "Step 6.8: Creating Statistical Summary Table (After Trimming/Remove Missing)..."
        )
        doc.add_heading(
            f"Statistical Summary Table (After Trimming/Remove Missing) - {file_name}",
            1,
        )
        table_stats = doc.add_table(rows=1, cols=9)
        table_stats.style = "Table Grid"
        table_stats.style.border_width = 1
        hdr_cells_stats = table_stats.rows[0].cells
        hdr_cells_stats[0].text = "Column"
        hdr_cells_stats[1].text = "Count"
        hdr_cells_stats[2].text = "Mean"
        hdr_cells_stats[3].text = "Std"
        hdr_cells_stats[4].text = "Min"
        hdr_cells_stats[5].text = "25%"
        hdr_cells_stats[6].text = "50%"
        hdr_cells_stats[7].text = "75%"
        hdr_cells_stats[8].text = "Max"

        for col in df.columns:
            if col != "Datetime":
                row_cells_stats = table_stats.add_row().cells
                row_cells_stats[0].text = col
                row_cells_stats[1].text = str(df[col].count())
                row_cells_stats[2].text = (
                    str(df[col].mean())
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[3].text = (
                    str(df[col].std())
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[4].text = (
                    str(df[col].min())
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[5].text = (
                    str(df[col].quantile(0.25))
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[6].text = (
                    str(df[col].median())
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[7].text = (
                    str(df[col].quantile(0.75))
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
                row_cells_stats[8].text = (
                    str(df[col].max())
                    if pd.api.types.is_numeric_dtype(df[col])
                    else "N/A"
                )
        desc_stats = f"ตารางนี้แสดงสถิติพื้นฐานของข้อมูลหลังลบฟีเจอร์ที่มี Missing > 50% และตัดทอนตาม SO2 ค่าที่ไม่ใช่ตัวเลขจะแสดง N/A ({file_name})"
        doc.add_paragraph(desc_stats)

        print("Step 7: Preparing Data for LSTM...")
        doc.add_heading(f"Data Preprocessing for LSTM - {file_name}", 1)

        features = None
        scaler_features = None
        scaler_target = None

        if len(df) > 0:
            all_dates = pd.date_range(
                start=df.index.min(), end=df.index.max(), freq="H"
            )
            df = df.reindex(all_dates).interpolate(method="linear")
            df_trimmed = df.dropna()
            if len(df_trimmed) > 0:
                if "SO2" in df_trimmed.columns:
                    features = [col for col in df_trimmed.columns if col != "SO2"]
                    scaler_features = MinMaxScaler()
                    scaler_target = MinMaxScaler()
                    df_trimmed[features] = scaler_features.fit_transform(
                        df_trimmed[features]
                    )
                    df_trimmed["SO2"] = scaler_target.fit_transform(df_trimmed[["SO2"]])
                    trim_start = df_trimmed.index.min().strftime("%Y-%m-%d %H:%M")
                    trim_end = df_trimmed.index.max().strftime("%Y-%m-%d %H:%M")
                    rows_removed = len(df) - len(df_trimmed)
                    desc = f"ข้อมูลถูกตัดทอนจาก {len(df)} เหลือ {len(df_trimmed)} แถว ลบออก {rows_removed} แถว หลังการเติมข้อมูล ลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 ({file_name}) "
                    if rows_removed > 0:
                        desc += f"วันที่เริ่มต้นหลัง trim: {trim_start}, วันที่สิ้นสุด: {trim_end} "
                        if rows_removed > len(df) * 0.1:
                            desc += "จำนวนแถวที่ลบมากกว่า 10% อาจส่งผลต่อช่วงข้อมูลต้นหรือปลาย "
                        else:
                            desc += "จำนวนแถวที่ลบอยู่ในระดับต่ำ ไม่น่ามีผลกระทบมาก "
                    else:
                        desc += "ไม่มีแถวที่ถูกตัดทอนเพิ่มเติม"
                    doc.add_paragraph(desc)
                    df = df_trimmed
                else:
                    desc = f"คอลัมน์ 'SO2' ไม่พบในข้อมูลหลังตัดทอน ใช้ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 แทน ({file_name})"
                    doc.add_paragraph(desc)
                    features = [col for col in df.columns if col != "SO2"]
                    scaler_features = MinMaxScaler()
                    scaler_target = MinMaxScaler()
                    if df[features].isnull().any().any() or df["SO2"].isnull().any():
                        df[features] = df[features].fillna(0)
                        df["SO2"] = df["SO2"].fillna(0)
                        doc.add_paragraph(
                            f"คำเตือน: ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 มีค่า NaN ได้รับการแทนที่ด้วย 0 เพื่อดำเนินการต่อ ผลลัพธ์อาจไม่แม่นยำ ({file_name})"
                        )
                    df[features] = scaler_features.fit_transform(df[features])
                    df["SO2"] = scaler_target.fit_transform(df[["SO2"]])
                    doc.add_paragraph(
                        f"คำเตือน: ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 อาจมีค่า NaN หรือไม่สมบูรณ์ ผลลัพธ์อาจไม่แม่นยำ ({file_name})"
                    )
            else:
                desc = f"ไม่มีข้อมูลเหลือหลังการตัดทอน ใช้ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 แทนเพื่อดำเนินการต่อ ({file_name})"
                doc.add_paragraph(desc)
                features = [col for col in df.columns if col != "SO2"]
                scaler_features = MinMaxScaler()
                scaler_target = MinMaxScaler()
                if df[features].isnull().any().any() or df["SO2"].isnull().any():
                    df[features] = df[features].fillna(0)
                    df["SO2"] = df["SO2"].fillna(0)
                    doc.add_paragraph(
                        f"คำเตือน: ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 มีค่า NaN ได้รับการแทนที่ด้วย 0 เพื่อดำเนินการต่อ ผลลัพธ์อาจไม่แม่นยำ ({file_name})"
                    )
                df[features] = scaler_features.fit_transform(df[features])
                df["SO2"] = scaler_target.fit_transform(df[["SO2"]])
                doc.add_paragraph(
                    f"คำเตือน: ข้อมูลดิบหลังลบ Missing > 50% (ไม่รวม SO2) และตัดตาม SO2 อาจมีค่า NaN หรือไม่สมบูรณ์ ผลลัพธ์อาจไม่แม่นยำ ({file_name})"
                )
        else:
            desc = f"ไม่สามารถดำเนินการต่อไปได้ เนื่องจากไม่มีข้อมูลหลังลบฟีเจอร์ที่มี Missing > 50% (ไม่รวม SO2) และตัดทอนตาม SO2 ({file_name})"
            doc.add_paragraph(desc)
            doc.save(self.output_dir / f"report_{file_name}.docx")
            return

        print("Step 7.1: Preparing Sequences for LSTM...")
        seq_length = self.CONST.SEQ_LENGTH
        pred_length = self.CONST.PRED_LENGTH
        if features is not None:
            X_full, y_full = prepare_sequences(df, features, seq_length, pred_length)
            print(f"X_full shape: {X_full.shape}, y_full shape: {y_full.shape}")
        else:
            print(f"Features not defined. Cannot prepare sequences - {file_name}.")
            doc.add_paragraph(
                f"Features not defined due to empty or invalid data. Skipping sequence preparation - {file_name}."
            )
            doc.save(self.output_dir / f"report_{file_name}_error.docx")
            return

        if features is not None:
            if df[features].isnull().any().any() or df["SO2"].isnull().any():
                df[features] = df[features].fillna(0)
                df["SO2"] = df["SO2"].fillna(0)
                doc.add_paragraph(
                    f"Warning: Flat data for XGBoost contains NaN values, replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )
            X_full_flat = df[features].values
            y_full_flat = df["SO2"].values
            print(
                f"X_full_flat shape: {X_full_flat.shape}, y_full_flat shape: {y_full_flat.shape}"
            )
        else:
            print(
                f"Features not defined. Cannot prepare flat data for XGBoost SHAP - {file_name}."
            )
            doc.add_paragraph(
                f"Features not defined due to empty or invalid data. Skipping XGBoost SHAP analysis - {file_name}."
            )
            doc.save(self.output_dir / f"report_{file_name}_error.docx")
            return

        print("Step 8: Training LSTM Model with All Features...")
        doc.add_heading(f"Model Training with LSTM - {file_name}", 1)
        doc.add_paragraph(
            f"Model: LSTM\nBatch Size: {self.CONST.BATCH_SIZE}\nEpochs: {self.CONST.EPOCHS}\nSplits: {self.CONST.N_SPLITS}\nFeatures Used: {list(features)} ({file_name})"
        )

        tscv = TimeSeriesSplit(n_splits=self.CONST.N_SPLITS)
        batch_size = self.CONST.BATCH_SIZE
        scores = {"mse": [], "mae": [], "r2": []}

        # Keep size for plot graph
        dataset_size = []

        for fold, (train_idx, test_idx) in enumerate(tscv.split(X_full)):
            print(f"Step 8.{fold + 1}: Training Fold {fold + 1} - {file_name}...")
            # skip fold to last fold for saving time
            if fold < self.CONST.N_SPLITS - 1:
                continue
            dataset_size.append({"train": len(train_idx), "validation": len(test_idx)})
            X_train, X_test = X_full[train_idx], X_full[test_idx]
            y_train, y_test = y_full[train_idx], y_full[test_idx]

            train_dataset = tf.data.Dataset.from_tensor_slices(
                (X_train, y_train)
            ).batch(batch_size)
            test_dataset = tf.data.Dataset.from_tensor_slices((X_test, y_test)).batch(
                batch_size
            )

            tf.keras.backend.clear_session()

            model = Sequential(
                [
                    LSTM(
                        128,
                        input_shape=(seq_length, len(features)),
                        return_sequences=True,
                    ),
                    Dropout(0.2),
                    LSTM(64),
                    Dropout(0.2),
                    Dense(32, activation="relu"),
                    Dense(pred_length),
                ]
            )
            model.compile(optimizer="adam", loss="mse")

            history = model.fit(
                train_dataset,
                epochs=self.CONST.EPOCHS,
                validation_data=test_dataset,
                verbose=1,
            )

            y_pred = model.predict(test_dataset)
            if np.any(np.isnan(y_test)) or np.any(np.isnan(y_pred)):
                doc.add_paragraph(
                    f"Warning: y_test or y_pred in Fold {fold + 1} contains NaN values found before inverse transform ({file_name})"
                )
                y_test = np.nan_to_num(y_test, nan=0.0)
                y_pred = np.nan_to_num(y_pred, nan=0.0)
                doc.add_paragraph(
                    f"NaN values in y_test and y_pred replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )

            y_test_inv = scaler_target.inverse_transform(y_test)
            y_pred_inv = scaler_target.inverse_transform(y_pred)

            if np.any(np.isnan(y_test_inv)) or np.any(np.isnan(y_pred_inv)):
                doc.add_paragraph(
                    f"Warning: y_test_inv or y_pred_inv in Fold {fold + 1} contains NaN values after inverse transform ({file_name})"
                )
                mask = ~np.isnan(y_test_inv) & ~np.isnan(y_pred_inv)
                if np.any(mask):
                    y_test_inv_clean = y_test_inv[mask]
                    y_pred_inv_clean = y_pred_inv[mask]
                else:
                    doc.add_paragraph(
                        f"No valid data in Fold {fold + 1}. Skipping metrics calculation ({file_name})"
                    )
                    continue
            else:
                y_test_inv_clean = y_test_inv
                y_pred_inv_clean = y_pred_inv

            mse = mean_squared_error(y_test_inv_clean, y_pred_inv_clean)
            mae = mean_absolute_error(y_test_inv_clean, y_pred_inv_clean)
            r2 = r2_score(y_test_inv_clean.flatten(), y_pred_inv_clean.flatten())
            scores["mse"].append(mse)
            scores["mae"].append(mae)
            scores["r2"].append(r2)

            plt.figure(figsize=(10, 6))
            plt.plot(history.history["loss"], label="Training Loss")
            plt.plot(history.history["val_loss"], label="Validation Loss")
            plt.title(f"Fold {fold + 1} Loss - {english_file_name}")
            plt.legend()
            plt.savefig(
                self.output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
            )
            plt.close()

            train_loss_mean = np.mean(history.history["loss"])
            train_loss_std = np.std(history.history["loss"])
            val_loss_mean = np.mean(history.history["val_loss"])
            val_loss_std = np.std(history.history["val_loss"])
            desc = f"กราฟนี้แสดง Loss จากการฝึก Fold {fold + 1} ({file_name}) "
            desc += f"ค่าเฉลี่ย Loss การฝึก: {train_loss_mean:.4f} (±{train_loss_std:.4f}), การทดสอบ: {val_loss_mean:.4f} (±{val_loss_std:.4f}) "
            if val_loss_mean > train_loss_mean * 1.2:
                desc += "แสดงถึงภาวะ overfitting ที่อาจเกิดขึ้นในช่วงท้ายการฝึก เนื่องจาก Loss ทดสอบสูงกว่าการฝึกอย่างมีนัยสำคัญ "
            elif val_loss_mean < train_loss_mean * 0.8:
                desc += "แสดงถึงภาวะ underfitting หรือการฝึกที่ยังไม่สมบูรณ์ เนื่องจาก Loss ทดสอบต่ำกว่าการฝึกอย่างมาก "
            else:
                desc += "แสดงถึงการฝึกที่สมดุลดี ระหว่างการฝึกและทดสอบ "
            desc += f"จำนวน epochs ที่ใช้: {self.CONST.EPOCHS}, Batch Size: {batch_size}"
            add_figure(
                doc,
                self.output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png",
                f"Fold {fold + 1} Loss - {file_name}",
                desc,
            )

            plt.figure(figsize=(12, 6))
            plt.plot(y_test_inv_clean.flatten()[:720], label="Actual")
            plt.plot(
                y_pred_inv_clean.flatten()[:720], label="Predicted", linestyle="--"
            )
            plt.title(
                f"Fold {fold + 1} Actual vs Predicted (First 30 Days) - {english_file_name}"
            )
            plt.legend()
            plt.savefig(
                self.output_dir / f"fold_{fold + 1}_pred_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
            )
            plt.close()

            pred_diff = np.abs(
                y_test_inv_clean.flatten()[:720] - y_pred_inv_clean.flatten()[:720]
            )
            max_diff_idx = np.argmax(pred_diff)
            max_diff_date = None
            test_start_idx = test_idx[0]
            date_idx = test_start_idx + seq_length + max_diff_idx
            if date_idx < len(df.index):
                max_diff_date = df.index[date_idx].strftime("%Y-%m-%d %H:%M")
            actual_mean = y_test_inv_clean.mean()
            pred_mean = y_pred_inv_clean.mean()
            actual_std = y_test_inv_clean.std()
            pred_std = y_pred_inv_clean.std()
            desc = f"กราฟนี้แสดงการเปรียบเทียบระหว่างค่าจริงและค่าพยากรณ์ Fold {fold + 1} (30 วันแรก) ({file_name}) "
            desc += f"ค่าเฉลี่ย SO2 จจริง: {actual_mean:.2f} (±{actual_std:.2f}), ค่าพยากรณ์: {pred_mean:.2f} (±{pred_std:.2f}) "
            if pred_diff.max() > actual_std * 2:
                if max_diff_date:
                    desc += f"จุดที่มีความคลาดเคลื่อนสูงสุด ({pred_diff.max():.2f}) เกิดขึ้นวันที่ {max_diff_date} "
                else:
                    desc += f"จุดที่มีความคลาดเคลื่อนสูงสุด ({pred_diff.max():.2f}) อยู่นอกขอบเขตวันที่ที่มีข้อมูล "
                desc += (
                    "แสดงถึงความคลาดเคลื่อนที่สูงผิดปกติ อาจเกิดจาก noise หรือ outlier ในข้อมูล "
                )
            else:
                desc += "การพยากรณ์ส่วนใหญ่ใกล้เคียงกับค่าจริง "
            desc += f"วันที่เริ่มต้นของชุดทดสอบ: {df.index[test_start_idx + seq_length].strftime('%Y-%m-%d %H:%M')}, "
            desc += f"วันที่สิ้นสุด: {df.index[min(test_start_idx + seq_length + 719, len(df.index) - 1)].strftime('%Y-%m-%d %H:%M')}"
            add_figure(
                doc,
                self.output_dir / f"fold_{fold + 1}_pred_{english_file_name}.png",
                f"Fold {fold + 1} Predictions - {file_name}",
                desc,
            )

            doc.add_heading(f"Fold {fold + 1} Metrics - {file_name}", 2)
            doc.add_paragraph(f"MSE: {mse:.4f}\nMAE: {mae:.4f}\nR²: {r2:.4f}")

        # Add TSCV split size plot into the report

        # Add average metrics table into the report
        doc.add_heading(f"แบ่งชุดข้อมูล (Time Series Cross Validation) - {file_name}", 1)
        doc.add_paragraph(
            "กราฟนี้แสดงขนาดของชุดข้อมูลที่ถูกแบ่งเป็น Train และ Validation สำหรับแต่ละ Fold โดยแสดงจำนวนข้อมูลและเปอร์เซ็นต์การแบ่ง"
        )
        self.plot_tscv(dataset_size)
        add_figure(
            doc,
            "temp_fig.png",
            f"Time Series Cross Validation Split Sizes - {file_name}",
            "กราฟนี้แสดงขนาดของชุดข้อมูลในแต่ละ Fold โดย Train แสดงด้วยสีขาว และ Validation ด้วยสีดำ พร้อมตัวเลขและเปอร์เซ็นต์การแบ่ง",
        )

        if scores["mse"]:
            avg_metrics = pd.DataFrame(
                [
                    {"Metric": "MSE", "Value": np.mean(scores["mse"])},
                    {"Metric": "MAE", "Value": np.mean(scores["mae"])},
                    {"Metric": "R²", "Value": np.mean(scores["r2"])},
                ]
            )
            desc = f"ตารางนี้แสดงค่าเฉลี่ยของเมตริกทั้ง {len(scores['mse'])} folds ที่มีข้อมูลสมบูรณ์ ({file_name}) "
            r2_avg = avg_metrics["Value"][avg_metrics["Metric"] == "R²"].values[0]
            mse_avg = avg_metrics["Value"][avg_metrics["Metric"] == "MSE"].values[0]
            mae_avg = avg_metrics["Value"][avg_metrics["Metric"] == "MAE"].values[0]
            desc += f"R²: {r2_avg:.4f}, MSE: {mse_avg:.4f}, MAE: {mae_avg:.4f} "
            if r2_avg > 0.8:
                desc += "R² สูงกว่า 0.8 แสดงถึงความแม่นยำสูงมากของโมเดล "
            elif r2_avg < 0.5:
                desc += "R² ต่ำกว่า 0.5 แสดงถึงความแม่นยำที่อาจไม่เพียงพอ "
            else:
                desc += "R² อยู่ในระดับปานกลาง "
            desc += "MSE และ MAE สะท้อนระดับของความคลาดเคลื่อนในการพยากรณ์"
            add_table(
                doc, avg_metrics, f"Average Metrics Across Folds - {file_name}", desc
            )
        else:
            doc.add_paragraph(
                f"Cannot calculate average metrics due to no complete data in all folds ({file_name})"
            )

        print("Step 9: Computing SHAP Tree Importance with XGBoost...")
        doc.add_heading(
            f"Feature Importance with SHAP Tree via XGBoost (Post-Training) - {file_name}",
            1,
        )

        if features is not None:
            xgb_model = xgboost.XGBRegressor(
                objective="reg:squarederror",
                tree_method="hist",
                device="cuda",
                random_state=self.CONST.SEED_VALUE,
            )
            if np.any(np.isnan(X_full_flat)) or np.any(np.isnan(y_full_flat)):
                doc.add_paragraph(
                    f"Warning: Flat data for XGBoost contains NaN values, replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )
                X_full_flat = np.nan_to_num(X_full_flat, nan=0.0)
                y_full_flat = np.nan_to_num(y_full_flat, nan=0.0)
            xgb_model.fit(X_full_flat, y_full_flat)
            X_sample_flat = X_full_flat[-100:]
            explainer = shap.TreeExplainer(xgb_model)
            shap_values = explainer.shap_values(X_sample_flat)
            shap_explanation = shap.Explanation(
                values=shap_values,
                base_values=explainer.expected_value,
                data=X_sample_flat,
                feature_names=features,
            )
            shap_importance = np.abs(shap_values).mean(axis=0)
            feature_importance = pd.DataFrame(
                {
                    "Feature": features,
                    "Importance": shap_importance,
                    "Min_SHAP": np.min(shap_values, axis=0),
                    "Max_SHAP": np.max(shap_values, axis=0),
                    "Mean_SHAP": np.mean(shap_values, axis=0),
                    "Std_SHAP": np.std(shap_values, axis=0),
                }
            ).sort_values("Importance", ascending=False)

            plt.figure(figsize=(10, 6))
            shap.plots.bar(shap_explanation, max_display=len(features), show=False)
            plt.title(f"SHAP Feature Importance (Bar) - {english_file_name}")
            plt.savefig(
                self.output_dir / f"shap_bar_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
                format="png",
            )
            plt.close()

            plt.figure(figsize=(12, 8))
            shap.plots.beeswarm(shap_explanation, max_display=len(features), show=False)
            plt.title(f"SHAP Feature Importance (Beeswarm) - {english_file_name}")
            plt.savefig(
                self.output_dir / f"shap_beeswarm_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
                format="png",
            )
            plt.close()

            bar_plot_path = self.output_dir / f"shap_bar_{english_file_name}.png"
            beeswarm_plot_path = (
                self.output_dir / f"shap_beeswarm_{english_file_name}.png"
            )

            max_fi_feature = feature_importance.iloc[0]["Feature"]
            max_fi_value = feature_importance.iloc[0]["Importance"]
            max_shap_min = feature_importance.iloc[0]["Min_SHAP"]
            max_shap_max = feature_importance.iloc[0]["Max_SHAP"]
            desc = f"This graph shows the feature importance from XGBoost SHAP Tree Explainer ({file_name}) "
            desc += f"The most important feature is '{max_fi_feature}' ({max_fi_value:.4f}) "
            desc += f"SHAP values of '{max_fi_feature}' range from {max_shap_min:.4f} to {max_shap_max:.4f} "
            if max_fi_value > shap_importance.mean() * 2:
                desc += "indicating a significantly higher influence compared to other features "
            else:
                desc += "indicating an influence similar to other features "
            desc += "SHAP value distribution shows both positive and negative impacts on the prediction"
            add_figure(
                doc, bar_plot_path, f"SHAP Tree Summary Plot (Bar) - {file_name}", desc
            )
            add_figure(
                doc,
                beeswarm_plot_path,
                f"SHAP Tree Summary Plot (Beeswarm) - {file_name}",
                desc,
            )
            desc = f"This table shows the ranking of feature importance from SHAP Tree. Average Importance: {shap_importance.mean():.4f}, "
            desc += f"Std Importance: {shap_importance.std():.4f}, Min Importance: {shap_importance.min():.4f}, Max Importance: {shap_importance.max():.4f} ({file_name}) "
            if feature_importance["Importance"].std() > shap_importance.mean() * 0.5:
                desc += "the distribution of importance is highly variable "
            else:
                desc += "the distribution of importance is relatively consistent "
            add_table(
                doc,
                feature_importance,
                f"Feature Importance Ranking (SHAP Tree via XGBoost) - {file_name}",
                desc,
            )
        else:
            print(f"Features not defined. Skipping SHAP Tree Importance - {file_name}.")
            doc.add_paragraph(
                f"Features not defined due to empty or invalid data. Skipping XGBoost SHAP analysis - {file_name}."
            )
            doc.save(self.output_dir / f"report_{file_name}_error.docx")
            return

        print("Step 10: Generating Final Prediction with LSTM...")
        doc.add_heading(f"Final Prediction with LSTM - {file_name}", 1)
        if features is not None and len(X_full) > 0:
            last_data = X_full[-1:]
            if np.any(np.isnan(last_data)):
                doc.add_paragraph(
                    f"Warning: Data for final prediction contains NaN values, replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )
                last_data = np.nan_to_num(last_data, nan=0.0)
            pred_next_24 = model.predict(last_data)
            pred_next_24_inv = scaler_target.inverse_transform(pred_next_24)
            future_timestamps = pd.date_range(start=df.index[-1], periods=25, freq="H")[
                1:
            ]
            actual_last_24 = df["SO2"][-24:].values.reshape(-1, 1)
            if np.any(np.isnan(actual_last_24)):
                doc.add_paragraph(
                    f"Warning: Actual data for the last 24 hours contains NaN values, replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )
                actual_last_24 = np.nan_to_num(actual_last_24, nan=0.0)
            actual_last_24_inv = scaler_target.inverse_transform(actual_last_24)
            plt.figure(figsize=(12, 6))
            plt.plot(df.index[-24:], actual_last_24_inv, label="Actual")
            plt.plot(
                future_timestamps,
                pred_next_24_inv[0],
                label="Predicted",
                linestyle="--",
            )
            plt.title(f"Next 24-Hour SO2 Prediction (LSTM) - {english_file_name}")
            plt.legend()
            plt.savefig(
                self.output_dir / f"final_pred_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
            )
            plt.close()

            actual_last_mean = actual_last_24_inv.mean()
            actual_last_std = actual_last_24_inv.std()
            actual_last_min = actual_last_24_inv.min()
            actual_last_max = actual_last_24_inv.max()
            pred_mean = pred_next_24_inv[0].mean()
            pred_std = pred_next_24_inv[0].std()
            pred_min = pred_next_24_inv[0].min()
            pred_max = pred_next_24_inv[0].max()
            desc = f"กราฟนี้แสดงการพยากรณ์ SO2 24 ชั่วโมงสุดท้าย ({file_name}) "
            desc += f"ค่าเฉลี่ยจริง: {actual_last_mean:.2f} (±{actual_last_std:.2f}), ค่าพยากรณ์: {pred_mean:.2f} (±{pred_std:.2f} "
            desc += f"Min จริง: {actual_last_min:.2f}, Max จริง: {actual_last_max:.2f}, "
            desc += f"Min พยากรณ์: {pred_min:.2f}, Max พยากรณ์: {pred_max:.2f} "
            if abs(pred_mean - actual_last_mean) > actual_last_mean * 0.2:
                desc += "การพยากรณ์มีความคลาดเคลื่อนมากกว่า 20% จากค่าจริง "
            else:
                desc += "การพยากรณ์ใกล้เคียงกับค่าจริงในระดับที่น่าพอใจ "
            desc += f"วันที่เริ่มต้นพยากรณ์: {future_timestamps[0].strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {future_timestamps[-1].strftime('%Y-%m-%d %H:%M')}"
            add_figure(
                doc,
                self.output_dir / f"final_pred_{english_file_name}.png",
                f"Next 24-Hour Prediction (LSTM) - {file_name}",
                desc,
            )
        else:
            print(
                f"Features or X_full not defined. Skipping final prediction - {file_name}."
            )
            doc.add_paragraph(
                f"Features or sequences not defined due to empty or invalid data. Skipping final prediction - {file_name}."
            )
            set_thsarabupsk_font(doc)
            doc.save(self.output_dir / f"report_{file_name}_error.docx")
            return

        print("Step 11: Saving Report...")
        set_thsarabupsk_font(doc)
        doc.save(self.output_dir / f"report_{file_name}.docx")
        print(f"Step 12: Process Completed for {file_name}!")

        # Clean memory (free memory)
        del df, df_raw, doc, X_full, y_full, X_full_flat, y_full_flat, model
        gc.collect()

    def run(self):
        for csv_file in self.csv_files:
            try:
                self.process_file(csv_file)
            except Exception as e:
                print(f"Error processing {csv_file}: {e}")
            gc.collect()


if __name__ == "__main__":
    input_directory = (
        "/mnt/e/MayThesis2025/src/labs4/s02_v01_02_clean/output-680507-2319"
    )
    processor = DynamicInputDataset(input_directory)
    processor.run()
