# Import Libraries
print("Step 1: Importing Libraries...")
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.linear_model import LinearRegression  # Added for scatter plot regression
import tensorflow as tf
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout
import xgboost
import shap
from docx import Document
from docx.shared import Inches
import random
import os
from pathlib import Path
import gc

from docx.oxml.ns import qn


# Do not change these functions:
def set_thsarabupsk_font(doc, font_name="TH SarabunPSK"):
    """Sets the specified font for all text in a Word document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font._element.rPr.rFonts.set(qn("w:ascii"), font_name)
            font._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
            font._element.rPr.rFonts.set(qn("w:cs"), font_name)
            font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            font._element.rPr.rFonts.set(qn("w:complexScript"), font_name)
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    font._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    font._element.rPr.rFonts.set(qn("w:cs"), font_name)
    font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    font._element.rPr.rFonts.set(qn("w:complexScript"), font_name)


def add_table(doc, df, title, description=""):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style="Table Grid")
    table.style.border_width = 1
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
    if description:
        doc.add_paragraph(description)


def add_figure(doc, filename, title, description=""):
    if isinstance(filename, Path):
        filename = str(filename)
    if os.path.exists(filename):
        doc.add_heading(title, level=2)
        doc.add_picture(filename, width=Inches(6))
        if description:
            doc.add_paragraph(description)
    else:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Warning: Image file {filename} not found.")


def prepare_sequences(data, features, seq_length, pred_length):
    X, y = [], []
    for i in range(len(data) - seq_length - pred_length + 1):
        X.append(data[features].iloc[i : i + seq_length].values)
        y.append(data["SO2"].iloc[i + seq_length : i + seq_length + pred_length].values)
    return np.array(X), np.array(y)


# Wrap everything into a class with a constant configuration class.
class DynamicInputDataset:
    # Constant configuration class for easy to config
    class CONST:
        SEED_VALUE = 42
        BATCH_SIZE = 1024
        EPOCHS = 50
        N_SPLITS = 5
        SEQ_LENGTH = 24
        PRED_LENGTH = 24
        OUTPUT_FOLDER = "outputs"

    # File name mapping (unchanged)
    file_name_mapping = {
        "(37t)ศาลหลักเมือง(ปิดสถานี)": "San_Lak_Mueang_Closed_Station",
        "(37t)สถานีอุตุนิยมวิทยาลำปาง": "Lampang_Meteorological_Station",
        "(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด": "Ban_Sopad_Hospital",
        "(39t)โรงพยาบาลส่งเสริมสุขภาพตำบลท่าสี": "Tha_Si_Hospital",
        "(40t)การประปาส่วนภูมิภาคแม่เมาะ": "Mae_Mo_Waterworks",
    }

    def __init__(self, input_dir):
        self.input_dir = Path(input_dir)
        self.csv_files = list(self.input_dir.glob("*.csv"))
        # Create output directory
        self.output_dir = Path.cwd() / self.CONST.OUTPUT_FOLDER
        self.output_dir.mkdir(exist_ok=True)
        # Set random seeds for reproducibility
        random.seed(self.CONST.SEED_VALUE)
        np.random.seed(self.CONST.SEED_VALUE)
        tf.random.set_seed(self.CONST.SEED_VALUE)
        # Configure GPU memory growth for LSTM
        gpus = tf.config.experimental.list_physical_devices("GPU")
        if gpus:
            try:
                for gpu in gpus:
                    tf.config.experimental.set_memory_growth(gpu, True)
                print("GPU is set up and ready for LSTM!")
            except RuntimeError as e:
                print(e)
        print("Num GPUs Available: ", len(tf.config.list_physical_devices("GPU")))

    def plot_tscv(self, dataset_size):
        # plot hbar sorted by index ascending
        df = pd.DataFrame(dataset_size)
        df = df.sort_index(ascending=False)
        df.plot(kind="barh", stacked=True)

        # show value on the bar and show total value
        for i in range(len(df)):
            total_text = df.iloc[i]["train"] + df.iloc[i]["validation"]
            train_text = df.iloc[i]["train"]
            train_percentage = train_text / total_text * 100
            validation_text = df.iloc[i]["validation"]
            validation_percentage = validation_text / total_text * 100
            plt.text(
                total_text,
                i,
                f"{total_text}",
                ha="left",
                va="center",
                color="black",
                fontsize=7,
            )
            plt.text(
                train_text / 2,
                i,
                f"{train_text} ({train_percentage:.2f}%)",
                ha="center",
                va="center",
                color="white",
                fontsize=7,
            )
            plt.text(
                train_text + validation_text / 2,
                i,
                f"{validation_text} ({validation_percentage:.2f}%)",
                ha="center",
                va="center",
                color="black",
                fontsize=7,
            )

        plt.xlabel("Dataset Size")
        plt.ylabel("Split (Fold)")
        plt.title("Time Series Cross Validation")
        plt.legend(["Train", "Validation"], loc="upper right")
        plt.savefig(
            "temp_fig.png",
            dpi=300,
            bbox_inches="tight",
            facecolor="white",
            transparent=False,
        )
        plt.show()
        plt.close()

    def process_file(self, file_path):
        print(f"Processing file: {file_path}")
        # Step 4: Loading Data using dynamic input
        print("Step 4: Loading Data...")
        df = pd.read_csv(file_path, parse_dates=["Datetime"], index_col="Datetime")
        # remove dev if exist
        if "dev" in df.columns:
            df = df.drop(columns=["dev"])
        df_raw = df.copy()  # Keep raw data for analysis
        file_name = os.path.basename(file_path).replace(".csv", "")
        english_file_name = self.file_name_mapping.get(file_name, file_name)
        # Note: The Thai character check for filename might be overly restrictive
        # if english_file_name contains mixed characters or if file_name is already English.
        # For now, assuming it works as intended for the user's specific file naming.
        # if all([char >= "ก" and char <= "๙" for char in file_name]):
        #     raise ValueError("Invalid filename: Must be in Thai characters [ก-๙]") # This check might be problematic if filename is English

        doc = Document()
        doc.add_heading(
            f"LSTM Model Report for SO2 Prediction with XGBoost SHAP FI - {file_name}",
            0,
        )
        doc.add_paragraph(
            f"Dataset: {file_path}"
        )  # Removed (Limited to 1000 rows) as it might not be true anymore

        print("Step 5: Showing Data Before Any Processing...")
        doc.add_heading(f"Data Before Any Processing ({file_name})", 1)
        doc.add_paragraph(f"Shape Before: {df.shape}")
        doc.add_paragraph(f"Missing Values Total Before: {df.isnull().sum().sum()}")

        print("Step 5.1: Listing Columns (Before)...")
        doc.add_heading(f"Columns List (Before) - {file_name}", 2)
        doc.add_paragraph(
            f"List of columns in the dataset before processing ({file_name}):"
        )
        for col in df.columns:
            doc.add_paragraph(f"- {col}")

        print("Step 5.2: Showing Data Sample (Before)...")
        df_head_before = df.head(10)
        add_table(
            doc,
            df_head_before,
            f"Sample of First 10 Rows (Before) - {file_name}",
            f"ตารางนี้แสดงตัวอย่าง 10 แถวแรกของข้อมูลดิบก่อนดำเนินการ ({file_name})",
        )

        print("Step 5.3: Computing Data Information (Before)...")
        doc.add_heading(f"Data Information (Before) - {file_name}", 2)
        # df.info() prints to stdout, capture it if needed for doc
        # For now, let's assume this step is for console output.
        # To add to doc, you'd need to capture stdout.
        # Example:
        # import io
        # buffer = io.StringIO()
        # df.info(buf=buffer)
        # info_output_before = buffer.getvalue()
        # doc.add_paragraph(str(info_output_before))
        # For simplicity, keeping original behavior (prints to console)

        print("Step 6: Performing EDA (Before Any Processing)...")
        doc.add_heading(
            f"Exploratory Data Analysis (Before Any Processing) - {file_name}", 1
        )
        doc.add_paragraph(f"Shape: {df.shape}")
        doc.add_paragraph(f"Missing Values Total: {df.isnull().sum().sum()}")

        plt.figure(figsize=(12, 6))
        heatmap = sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(f"Missing Data Heatmap (Before Any Processing) - {english_file_name}")
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_before_all = df.isnull().sum().sum()
        nan_max_col_before_all = (
            df.isnull().sum().idxmax() if nan_count_before_all > 0 else "None"
        )
        nan_max_val_before_all = (
            df.isnull().sum().max() if nan_count_before_all > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) ก่อนดำเนินการใด ๆ รวมทั้งหมด {nan_count_before_all} ค่า ({file_name}) "
        if nan_count_before_all > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_before_all}' ({nan_max_val_before_all} ค่า) "
            if nan_max_val_before_all > df.shape[0] * 0.5:
                desc += (
                    " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่ "
                )
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น "
        else:
            desc += "ไม่มีข้อมูลขาดหาย "
        desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}"
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_before_all_{english_file_name}.png",
            f"Missing Data Heatmap (Before Any Processing) - {file_name}",
            desc,
        )

        corr_matrix_before = df.corr()
        if "SO2" in corr_matrix_before:
            correlation_df_before = pd.DataFrame(
                {
                    "feature": corr_matrix_before.index,
                    "correlation": corr_matrix_before["SO2"],
                    "abs_correlation": corr_matrix_before["SO2"].abs(),
                }
            )
            correlation_df_before = correlation_df_before.sort_values(
                by="abs_correlation", ascending=False
            )
            add_table(
                doc,
                correlation_df_before,
                f"Sorted Correlation Matrix with SO2 (Before) - {file_name}",
                f"ตารางนี้แสดงค่าสหสัมพันธ์กับ SO2 เรียงลำดับตามค่าสัมประสิทธิ์สัมพันธ์สัมบูรณ์ก่อนดำเนินการ ({file_name})",
            )

            corr_max_before = (
                df.corr()["SO2"].drop("SO2").abs().max()
                if len(df.corr()["SO2"].drop("SO2")) > 0
                else np.nan
            )
            corr_max_col_before = (
                df.corr()["SO2"].drop("SO2").abs().idxmax()
                if len(df.corr()["SO2"].drop("SO2")) > 0
                else "N/A"
            )
            corr_min_before = (
                df.corr()["SO2"].drop("SO2").abs().min()
                if len(df.corr()["SO2"].drop("SO2")) > 0
                else np.nan
            )
            corr_min_col_before = (
                df.corr()["SO2"].drop("SO2").abs().idxmin()
                if len(df.corr()["SO2"].drop("SO2")) > 0
                else "N/A"
            )
        else:
            (
                corr_max_before,
                corr_max_col_before,
                corr_min_before,
                corr_min_col_before,
            ) = np.nan, "N/A", np.nan, "N/A"
            doc.add_paragraph(
                f"SO2 column not found for correlation analysis (Before) - {file_name}"
            )

        plt.figure(figsize=(10, 8))
        heatmap = sns.heatmap(
            corr_matrix_before,
            annot=True,
            cmap="RdBu",
            center=0,
            vmin=-1,
            vmax=1,
            fmt=".2f",
        )
        plt.title(f"Correlation Matrix (Before Any Processing) - {english_file_name}")
        plt.xticks(rotation=45, ha="right")
        plt.yticks(rotation=0)
        plt.savefig(
            self.output_dir / f"eda_corr_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = f"กราฟนี้แสดงความสัมพันธ์ระหว่างฟีเจอร์ก่อนดำเนินการ ({file_name}). "
        if "SO2" in df.columns and not pd.isna(corr_max_before):
            desc += f"ค่าสหสัมพันธ์สูงสุดกับ SO2 คือ {corr_max_before:.2f} จาก '{corr_max_col_before}'. "
            desc += f"ค่าสหสัมพันธ์ต่ำสุดกับ SO2 คือ {corr_min_before:.2f} จาก '{corr_min_col_before}'. "
            if corr_max_before > 0.7:
                desc += "แสดงถึงความสัมพันธ์ที่สูงมาก. "
            elif corr_max_before > 0.3:
                desc += "แสดงถึงความสัมพันธ์ปานกลาง. "
            else:
                desc += "แสดงถึงความสัมพันธ์ที่ต่ำ. "
        else:
            desc += "ไม่สามารถคำนวณค่าสหสัมพันธ์กับ SO2 ได้. "
        desc += "การกระจายของค่าสหสัมพันธ์มีแนวโน้มสมมาตรตามแนวทแยง."
        add_figure(
            doc,
            self.output_dir / f"eda_corr_before_all_{english_file_name}.png",
            f"Correlation Matrix (Before Any Processing) - {file_name}",
            desc,
        )

        plt.figure(figsize=(15, 6))
        df.boxplot(figsize=(15, 6))  # Returns Axes object, store if needed
        plt.xticks(rotation=45, ha="right")
        plt.title(
            f"Feature Outlier Distribution (Before Any Processing) - {english_file_name}"
        )
        plt.savefig(
            self.output_dir / f"eda_boxplot_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        outliers_before = {}
        for col in df.columns:
            if pd.api.types.is_numeric_dtype(df[col]):
                q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
                iqr = q3 - q1
                if iqr == 0:  # Avoid division by zero or issues with constant columns
                    outlier_count = ((df[col] < q1) | (df[col] > q3)).sum()
                else:
                    outlier_count = (
                        (df[col] < (q1 - 1.5 * iqr)) | (df[col] > (q3 + 1.5 * iqr))
                    ).sum()
                outliers_before[col] = outlier_count
            else:
                outliers_before[col] = 0  # No outliers for non-numeric

        if outliers_before:  # Check if dictionary is not empty
            max_outlier_col_before = max(outliers_before, key=outliers_before.get)
            min_outlier_col_before = min(outliers_before, key=outliers_before.get)
            desc = f"กราฟนี้แสดงการกระจายและ outlier ของฟีเจอร์ก่อนดำเนินการ ({file_name}). คอลัมน์ที่มี outlier มากที่สุดคือ '{max_outlier_col_before}' ({outliers_before[max_outlier_col_before]} ค่า). "
            desc += f"คอลัมน์ที่มี outlier น้อยที่สุดคือ '{min_outlier_col_before}' ({outliers_before[min_outlier_col_before]} ค่า). "
            if outliers_before[max_outlier_col_before] > df.shape[0] * 0.1:
                desc += "ซึ่งเกิน 10% ของข้อมูล แสดงถึงการกระจายที่มีความผันผวนสูง. "
            else:
                desc += "ซึ่งอยู่ในระดับต่ำ ไม่มีผลกระทบมาก. "
            if pd.api.types.is_numeric_dtype(df[max_outlier_col_before]):
                desc += f"ค่าสถิติของ '{max_outlier_col_before}': Mean={df[max_outlier_col_before].mean():.2f}, Std={df[max_outlier_col_before].std():.2f}, "
                desc += f"Min={df[max_outlier_col_before].min():.2f}, Max={df[max_outlier_col_before].max():.2f}."
        else:
            desc = f"ไม่สามารถคำนวณ Outliers ได้ ({file_name})."

        add_figure(
            doc,
            self.output_dir / f"eda_boxplot_before_all_{english_file_name}.png",
            f"Feature Outlier Distribution (Before Any Processing) - {file_name}",
            desc,
        )

        n_cols = 4
        n_rows = (len(df.columns) + n_cols - 1) // n_cols
        plt.figure(figsize=(15, n_rows * 3))
        for i, col in enumerate(df.columns):
            plt.subplot(n_rows, n_cols, i + 1)
            data = df[col].dropna()
            if not data.empty and pd.api.types.is_numeric_dtype(data):
                plt.hist(data, bins=50, edgecolor="black")
            else:
                plt.text(0.5, 0.5, "No valid numeric data", ha="center", va="center")
            plt.title(col)
        plt.tight_layout()
        plt.savefig(
            self.output_dir / f"eda_histogram_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = (
            f"กราฟนี้แสดงการกระจายของแต่ละฟีเจอร์ก่อนดำเนินการในรูปแบบฮิสโตแกรม ({file_name}). "
        )
        numeric_df_before = df.select_dtypes(include=np.number)
        if not numeric_df_before.empty:
            skewness_before = numeric_df_before.skew().abs()
            if not skewness_before.empty:
                skew_max_before = skewness_before.max()
                skew_max_col_before = skewness_before.idxmax()
                skew_min_before = skewness_before.min()
                skew_min_col_before = skewness_before.idxmin()
                desc += f"ฟีเจอร์ที่มี skewness สูงสุดคือ '{skew_max_col_before}' ({skew_max_before:.2f}). "
                desc += f"ฟีเจอร์ที่มี skewness ต่ำสุดคือ '{skew_min_col_before}' ({skew_min_before:.2f}). "
                if skew_max_before > 1:
                    desc += "แสดงถึงการกระจายที่เบ้มาก (skewed distribution). "
                elif skew_max_before > 0.5:
                    desc += "แสดงถึงการกระจายที่เบ้ปานกลาง. "
                else:
                    desc += "แสดงถึงการกระจายที่ใกล้เคียงปกติ. "
                if pd.api.types.is_numeric_dtype(df[skew_max_col_before]):
                    desc += f"ค่าสถิติของ '{skew_max_col_before}': Min={df[skew_max_col_before].min():.2f}, Max={df[skew_max_col_before].max():.2f}, "
                    desc += f"Mean={df[skew_max_col_before].mean():.2f}, Std={df[skew_max_col_before].std():.2f}."
            else:
                desc += "ไม่สามารถคำนวณ Skewness ได้ (ไม่มีฟีเจอร์ตัวเลขที่เหลือ). "
        else:
            desc += "ไม่มีฟีเจอร์ตัวเลขสำหรับคำนวณ Skewness. "
        add_figure(
            doc,
            self.output_dir / f"eda_histogram_before_all_{english_file_name}.png",
            f"Feature Histograms (Before Any Processing) - {file_name}",
            desc,
        )

        print(
            "Step 6.3: Creating Timeseries Plot for All Features (Before Any Processing)..."
        )
        n_features = len(df.columns)
        plt.figure(figsize=(15, n_features * 4))
        for i, col in enumerate(df.columns):
            plt.subplot(n_features, 1, i + 1)
            data = df[col].dropna()
            if not data.empty:
                plt.plot(df.index[df[col].notna()], data, label=col)
            else:
                plt.text(0.5, 0.5, "No valid data (all NaN)", ha="center", va="center")
            plt.title(
                f"Timeseries of {col} - {english_file_name}"
            )  # Removed Yearly Labels for simplicity
            plt.legend()
        plt.tight_layout()
        plt.savefig(
            self.output_dir / f"timeseries_plot_before_all_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        desc = f"กราฟนี้แสดงข้อมูล Timeseries ของทุกฟีเจอร์ก่อนดำเนินการ ({file_name}). "
        desc += "ความสูงของภาพปรับตามจำนวนฟีเจอร์ (4 นิ้วต่อฟีเจอร์)."
        add_figure(
            doc,
            self.output_dir / f"timeseries_plot_before_all_{english_file_name}.png",
            f"Timeseries Plot for All Features (Before Any Processing) - {file_name}",
            desc,
        )

        print("Step 6.4: Removing Features with Missing > 50% (Excluding SO2)...")
        missing_pct = df.isnull().mean()
        features_to_drop = missing_pct[missing_pct > 0.5].index
        features_to_drop = [col for col in features_to_drop if col != "SO2"]
        if features_to_drop:
            df = df.drop(columns=features_to_drop)
            doc.add_paragraph(
                f"Dropped features (>50% missing, excluding SO2) - {file_name}: {features_to_drop}"
            )
        else:
            doc.add_paragraph(
                f"No features with missing > 50% (excluding SO2) to drop - {file_name}."
            )

        doc.add_heading(
            f"Exploratory Data Analysis (After Removing Missing > 50%, Before Trim SO2) - {file_name}",  # Simplified title
            1,
        )
        doc.add_paragraph(
            f"Shape after removing features with missing > 50% (excluding SO2) - {file_name}: {df.shape}"
        )
        doc.add_paragraph(
            f"Missing Values Total after removing features - {file_name}: {df.isnull().sum().sum()}"
        )

        plt.figure(figsize=(12, 6))
        sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(
            f"Missing Data Heatmap (After Removing Missing > 50%, Before Trim SO2) - {english_file_name}"
        )
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_after_remove_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_after_remove = df.isnull().sum().sum()
        nan_max_col_after_remove = (
            df.isnull().sum().idxmax() if nan_count_after_remove > 0 else "None"
        )
        nan_max_val_after_remove = (
            df.isnull().sum().max() if nan_count_after_remove > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) หลังลบฟีเจอร์ที่มี Missing > 50% (ไม่รวม SO2) แต่ก่อนตัดตาม SO2 รวมทั้งหมด {nan_count_after_remove} ค่า ({file_name}). "
        if nan_count_after_remove > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_after_remove}' ({nan_max_val_after_remove} ค่า). "
            if nan_max_val_after_remove > df.shape[0] * 0.5:
                desc += (
                    " ซึ่งมากกว่า 50% ของข้อมูล แสดงถึงการขาดหายในช่วงต้นหรือปลายข้อมูลเป็นส่วนใหญ่. "
                )
            else:
                desc += " ซึ่งกระจายในบางช่วงเวลาเท่านั้น. "
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วงนี้หลังลบฟีเจอร์. "
        if not df.empty:
            desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}."
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_after_remove_{english_file_name}.png",
            f"Missing Data Heatmap (After Removing Missing > 50%, Before Trim SO2) - {file_name}",
            desc,
        )

        # Correlation Matrix (After Removing Missing > 50%)
        if not df.empty:
            corr_matrix_after_remove = df.corr()
            if "SO2" in corr_matrix_after_remove:
                correlation_df_after_remove = pd.DataFrame(
                    {
                        "feature": corr_matrix_after_remove.index,
                        "correlation": corr_matrix_after_remove["SO2"],
                        "abs_correlation": corr_matrix_after_remove["SO2"].abs(),
                    }
                ).sort_values(by="abs_correlation", ascending=False)
                add_table(
                    doc,
                    correlation_df_after_remove,
                    f"Sorted Correlation Matrix with SO2 (After Removing Missing > 50%) - {file_name}",
                    f"ตารางนี้แสดงค่าสหสัมพันธ์กับ SO2 หลังลบฟีเจอร์ที่มี Missing > 50% ({file_name}).",
                )

                plt.figure(figsize=(10, 8))
                sns.heatmap(
                    corr_matrix_after_remove,
                    annot=True,
                    cmap="RdBu",
                    center=0,
                    vmin=-1,
                    vmax=1,
                    fmt=".2f",
                )
                plt.title(
                    f"Correlation Matrix (After Removing Missing > 50%) - {english_file_name}"
                )
                plt.xticks(rotation=45, ha="right")
                plt.yticks(rotation=0)
                corr_path_after_remove = (
                    self.output_dir / f"eda_corr_after_remove_{english_file_name}.png"
                )
                plt.savefig(corr_path_after_remove, dpi=300, bbox_inches="tight")
                plt.close()

                if not df.corr()["SO2"].drop("SO2", errors="ignore").empty:
                    corr_max_after = (
                        df.corr()["SO2"].drop("SO2", errors="ignore").abs().max()
                    )
                    corr_max_col_after = (
                        df.corr()["SO2"].drop("SO2", errors="ignore").abs().idxmax()
                    )
                else:
                    corr_max_after, corr_max_col_after = np.nan, "N/A"

                desc_corr_after_remove = (
                    f"กราฟนี้แสดงความสัมพันธ์ระหว่างฟีเจอร์หลังลบ Missing > 50% ({file_name}). "
                )
                if not pd.isna(corr_max_after):
                    desc_corr_after_remove += f"ค่าสหสัมพันธ์สูงสุดกับ SO2 คือ {corr_max_after:.2f} จาก '{corr_max_col_after}'. "
                add_figure(
                    doc,
                    corr_path_after_remove,
                    f"Correlation Matrix (After Removing Missing > 50%) - {file_name}",
                    desc_corr_after_remove,
                )
            else:
                doc.add_paragraph(
                    f"SO2 column not found for correlation after removing high missing features - {file_name}"
                )
        else:
            doc.add_paragraph(
                f"DataFrame is empty after removing high missing features, skipping correlation - {file_name}"
            )

        print("Step 6.5: Trimming Head/Tails Based on SO2...")
        if "SO2" in df.columns and df["SO2"].notna().any():
            first_valid_so2_index = df["SO2"].first_valid_index()
            last_valid_so2_index = df["SO2"].last_valid_index()
            if first_valid_so2_index is not None and last_valid_so2_index is not None:
                df = df.loc[first_valid_so2_index:last_valid_so2_index]
                doc.add_paragraph(
                    f"Trimmed head/tails based on non-NaN SO2 values - {file_name}. New shape: {df.shape}"
                )
            else:
                doc.add_paragraph(
                    f"No valid SO2 values range found for trimming - {file_name}. Dataframe shape: {df.shape}"
                )
        else:
            doc.add_paragraph(
                f"SO2 column not found or all NaN. Skipping trimming based on SO2 - {file_name}."
            )

        doc.add_heading(
            f"Exploratory Data Analysis (After Trimming Based on SO2) - {file_name}", 1
        )
        doc.add_paragraph(
            f"Shape after trimming based on SO2 - {file_name}: {df.shape}"
        )
        doc.add_paragraph(
            f"Missing Values Total after trimming - {file_name}: {df.isnull().sum().sum()}"
        )

        plt.figure(figsize=(12, 6))
        sns.heatmap(df.isnull(), cbar=True, cmap="Blues", annot=False)
        plt.title(
            f"Missing Data Heatmap (After Trimming Based on SO2) - {english_file_name}"
        )
        plt.xlabel("Features")
        plt.ylabel("Samples")
        plt.savefig(
            self.output_dir / f"eda_heatmap_after_trim_so2_{english_file_name}.png",
            dpi=300,
            bbox_inches="tight",
        )
        plt.close()

        nan_count_after_trim_so2 = df.isnull().sum().sum()
        nan_max_col_after_trim_so2 = (
            df.isnull().sum().idxmax() if nan_count_after_trim_so2 > 0 else "None"
        )
        nan_max_val_after_trim_so2 = (
            df.isnull().sum().max() if nan_count_after_trim_so2 > 0 else 0
        )
        desc = f"กราฟนี้แสดงการกระจายของข้อมูลที่ขาดหาย (NaN) หลังตัดทอนตาม SO2 รวมทั้งหมด {nan_count_after_trim_so2} ค่า ({file_name}). "
        if nan_count_after_trim_so2 > 0:
            desc += f"คอลัมน์ที่มี NaN มากที่สุดคือ '{nan_max_col_after_trim_so2}' ({nan_max_val_after_trim_so2} ค่า). "
            # ... (rest of description logic as before)
        else:
            desc += "ไม่มีข้อมูลขาดหายในช่วงนี้หลังตัดทอน. "
        if not df.empty:
            desc += f"วันที่เริ่มต้น: {df.index.min().strftime('%Y-%m-%d %H:%M')}, วันที่สิ้นสุด: {df.index.max().strftime('%Y-%m-%d %H:%M')}."
        add_figure(
            doc,
            self.output_dir / f"eda_heatmap_after_trim_so2_{english_file_name}.png",
            f"Missing Data Heatmap (After Trimming Based on SO2) - {file_name}",
            desc,
        )

        print(
            "Step 6.6: Showing Simple Data Sample (10 Rows) After Trimming/Remove Missing..."
        )
        df_head_after = df.head(10)
        add_table(
            doc,
            df_head_after,
            f"Sample of First 10 Rows (After Trimming/Remove Missing) - {file_name}",
            f"ตารางนี้แสดงตัวอย่าง 10 แถวแรกของข้อมูลหลังลบฟีเจอร์ที่มี Missing > 50% และตัดทอนตาม SO2 ({file_name})",
        )

        # ... (Statistical summary table part - assuming it's mostly correct from original)
        print("Step 6.7: Computing Data Information (After Trimming/Remove Missing)...")
        # ... (This part was for df.info() like table, keeping it simple)
        print(
            "Step 6.8: Creating Statistical Summary Table (After Trimming/Remove Missing)..."
        )
        if not df.empty:
            summary_stats_df = df.describe().transpose().reset_index()
            summary_stats_df.columns = [
                "Column",
                "Count",
                "Mean",
                "Std",
                "Min",
                "25%",
                "50%",
                "75%",
                "Max",
            ]
            add_table(
                doc,
                summary_stats_df,
                f"Statistical Summary Table (After Trimming/Remove Missing) - {file_name}",
                f"ตารางสรุปสถิติของข้อมูลหลังการประมวลผล ({file_name})",
            )
        else:
            doc.add_paragraph(
                f"DataFrame is empty, cannot generate statistical summary table - {file_name}"
            )

        print("Step 7: Preparing Data for LSTM...")
        doc.add_heading(f"Data Preprocessing for LSTM - {file_name}", 1)

        features = None
        scaler_features = None
        scaler_target = None

        if not df.empty and "SO2" in df.columns:
            # Fill remaining NaNs using interpolation after major cleaning
            # Reindex to ensure hourly frequency and interpolate
            if not df.index.is_monotonic_increasing:
                df = df.sort_index()

            all_dates = pd.date_range(
                start=df.index.min(), end=df.index.max(), freq="H"
            )
            df_reindexed = df.reindex(all_dates)

            # Interpolate all columns. For SO2, limit to avoid creating new issues if large gaps exist.
            # For features, linear interpolation is generally fine.
            feature_cols_to_interpolate = [
                col for col in df_reindexed.columns if col != "SO2"
            ]
            df_reindexed[feature_cols_to_interpolate] = df_reindexed[
                feature_cols_to_interpolate
            ].interpolate(method="linear", limit_direction="both")

            # For SO2, interpolate then fill remaining with 0 or a more robust method if needed
            df_reindexed["SO2"] = df_reindexed["SO2"].interpolate(
                method="linear", limit_direction="both"
            )

            # Drop rows where SO2 is still NaN (essential for target)
            df_processed = df_reindexed.dropna(subset=["SO2"])

            # For other features, if NaNs remain after interpolation, fill with 0 or mean/median
            # For simplicity, let's fill with 0 for remaining NaNs in features
            df_processed = df_processed.fillna(
                0
            )  # Fill any remaining NaNs in features with 0

            if not df_processed.empty:
                df = df_processed  # Update df to the processed version
                features = [col for col in df.columns if col != "SO2"]
                if not features:  # If only SO2 column remains
                    doc.add_paragraph(
                        f"No features available after processing for {file_name}. Cannot proceed with LSTM."
                    )
                    set_thsarabupsk_font(doc)
                    doc.save(
                        self.output_dir / f"report_{file_name}_error_no_features.docx"
                    )
                    return

                scaler_features = MinMaxScaler()
                scaler_target = MinMaxScaler()

                df[features] = scaler_features.fit_transform(df[features])
                df["SO2"] = scaler_target.fit_transform(df[["SO2"]])
                doc.add_paragraph(
                    f"Data scaled for LSTM. Shape: {df.shape} ({file_name})."
                )
            else:
                doc.add_paragraph(
                    f"Data is empty after processing NaNs for {file_name}. Cannot proceed."
                )
                set_thsarabupsk_font(doc)
                doc.save(self.output_dir / f"report_{file_name}_error_empty_data.docx")
                return
        else:
            doc.add_paragraph(
                f"DataFrame is empty or SO2 column missing before scaling for {file_name}. Cannot proceed."
            )
            set_thsarabupsk_font(doc)
            doc.save(self.output_dir / f"report_{file_name}_error_early_empty.docx")
            return

        print("Step 7.1: Preparing Sequences for LSTM...")
        seq_length = self.CONST.SEQ_LENGTH
        pred_length = self.CONST.PRED_LENGTH
        if features is not None and not df.empty:
            X_full, y_full = prepare_sequences(df, features, seq_length, pred_length)
            if X_full.shape[0] == 0:
                doc.add_paragraph(
                    f"Not enough data to create sequences for LSTM ({file_name}). X_full is empty."
                )
                set_thsarabupsk_font(doc)
                doc.save(
                    self.output_dir / f"report_{file_name}_error_empty_sequences.docx"
                )
                return
            print(f"X_full shape: {X_full.shape}, y_full shape: {y_full.shape}")
        else:
            print(
                f"Features not defined or df is empty. Cannot prepare sequences - {file_name}."
            )
            doc.add_paragraph(
                f"Features not defined or df is empty. Skipping sequence preparation - {file_name}."
            )
            set_thsarabupsk_font(doc)
            doc.save(self.output_dir / f"report_{file_name}_error_no_seq_prep.docx")
            return

        # Prepare flat data for XGBoost SHAP
        if features is not None and not df.empty:
            # Ensure no NaNs in data for XGBoost
            df_for_xgb = df.copy()  # Use the scaled df
            if (
                df_for_xgb[features].isnull().any().any()
                or df_for_xgb["SO2"].isnull().any()
            ):
                df_for_xgb[features] = df_for_xgb[features].fillna(0)
                df_for_xgb["SO2"] = df_for_xgb["SO2"].fillna(0)
                doc.add_paragraph(
                    f"Warning: Flat data for XGBoost (from scaled df) had NaNs, replaced with 0. ({file_name})"
                )
            X_full_flat = df_for_xgb[features].values
            y_full_flat = df_for_xgb["SO2"].values
            print(
                f"X_full_flat shape: {X_full_flat.shape}, y_full_flat shape: {y_full_flat.shape}"
            )
        else:
            print(
                f"Features not defined or df empty. Cannot prepare flat data for XGBoost SHAP - {file_name}."
            )
            # ... (error handling as before)
            set_thsarabupsk_font(doc)
            doc.save(self.output_dir / f"report_{file_name}_error_xgb_prep.docx")
            return

        print("Step 8: Training LSTM Model with All Features...")
        doc.add_heading(f"Model Training with LSTM - {file_name}", 1)
        doc.add_paragraph(
            f"Model: LSTM\nBatch Size: {self.CONST.BATCH_SIZE}\nEpochs: {self.CONST.EPOCHS}\nSplits: {self.CONST.N_SPLITS}\nFeatures Used: {list(features)} ({file_name})"
        )

        tscv = TimeSeriesSplit(n_splits=self.CONST.N_SPLITS)
        batch_size = self.CONST.BATCH_SIZE
        scores = {"mse": [], "mae": [], "r2": []}
        dataset_size = []

        # To store the model from the last fold for later use if needed (e.g. final prediction)
        last_fold_model = None
        last_fold_history = None

        for fold, (train_idx, test_idx) in enumerate(tscv.split(X_full)):
            print(f"Step 8.{fold + 1}: Training Fold {fold + 1} - {file_name}...")

            current_fold_size = {"train": len(train_idx), "validation": len(test_idx)}
            # This logic means only the last fold's training details are captured for dataset_size plot.
            # If you want all folds for plot_tscv, remove the 'if fold < ... continue' or adjust dataset_size append.
            if fold < self.CONST.N_SPLITS - 1:
                dataset_size.append(
                    current_fold_size
                )  # Still append for full TSCV plot
                continue

            dataset_size.append(current_fold_size)  # Append for the last fold too

            X_train, X_test = X_full[train_idx], X_full[test_idx]
            y_train, y_test = y_full[train_idx], y_full[test_idx]

            train_dataset = tf.data.Dataset.from_tensor_slices(
                (X_train, y_train)
            ).batch(batch_size)
            test_dataset = tf.data.Dataset.from_tensor_slices((X_test, y_test)).batch(
                batch_size
            )

            tf.keras.backend.clear_session()
            model = Sequential(
                [
                    LSTM(
                        128,
                        input_shape=(seq_length, len(features)),
                        return_sequences=True,
                    ),
                    Dropout(0.2),
                    LSTM(64),
                    Dropout(0.2),
                    Dense(32, activation="relu"),
                    Dense(pred_length),
                ]
            )
            model.compile(optimizer="adam", loss="mse")

            history = model.fit(
                train_dataset,
                epochs=self.CONST.EPOCHS,
                validation_data=test_dataset,
                verbose=1,  # Set to 0 or 2 for less output if preferred
            )
            last_fold_model = model  # Save the model from the last fold
            last_fold_history = history  # Save history from the last fold

            y_pred = model.predict(test_dataset)
            if np.any(np.isnan(y_test)) or np.any(np.isnan(y_pred)):
                doc.add_paragraph(
                    f"Warning: y_test or y_pred in Fold {fold + 1} contains NaN values found before inverse transform ({file_name})"
                )
                y_test = np.nan_to_num(y_test, nan=0.0)
                y_pred = np.nan_to_num(y_pred, nan=0.0)
                doc.add_paragraph(
                    f"NaN values in y_test and y_pred replaced with 0 to proceed. Results may be inaccurate ({file_name})"
                )

            y_test_inv = scaler_target.inverse_transform(y_test)
            y_pred_inv = scaler_target.inverse_transform(y_pred)

            # Clean NaNs for metrics calculation
            # Flatten and then remove NaNs pair-wise for metrics
            y_test_inv_flat_for_metrics = y_test_inv.flatten()
            y_pred_inv_flat_for_metrics = y_pred_inv.flatten()

            valid_metric_indices = ~np.isnan(y_test_inv_flat_for_metrics) & ~np.isnan(
                y_pred_inv_flat_for_metrics
            )
            y_test_inv_clean_flat = y_test_inv_flat_for_metrics[valid_metric_indices]
            y_pred_inv_clean_flat = y_pred_inv_flat_for_metrics[valid_metric_indices]

            if len(y_test_inv_clean_flat) > 0:
                mse = mean_squared_error(y_test_inv_clean_flat, y_pred_inv_clean_flat)
                mae = mean_absolute_error(y_test_inv_clean_flat, y_pred_inv_clean_flat)
                r2 = r2_score(y_test_inv_clean_flat, y_pred_inv_clean_flat)
                scores["mse"].append(mse)
                scores["mae"].append(mae)
                scores["r2"].append(r2)
            else:
                doc.add_paragraph(
                    f"Fold {fold + 1}: No valid data points left after cleaning NaNs for metrics calculation. Skipping metrics for this fold."
                )
                # Append NaN or skip for scores if desired, affects average calculation
                scores["mse"].append(np.nan)
                scores["mae"].append(np.nan)
                scores["r2"].append(np.nan)

            plt.figure(figsize=(10, 6))
            plt.plot(history.history["loss"], label="Training Loss")
            plt.plot(history.history["val_loss"], label="Validation Loss")
            plt.title(f"Fold {fold + 1} Loss - {english_file_name}")
            plt.legend()
            plt.savefig(
                self.output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png",
                dpi=300,
                bbox_inches="tight",
            )
            plt.close()
            # ... (Description for loss plot as before) ...
            # 
            #             add_figure(
                doc,
                self.output_dir / f"fold_{fold + 1}_loss_{english_file_name}.png",
                f"Fold {fold + 1} Loss - {file_name}",
                f"กราฟ Loss ของ Fold {fold + 1} ({file_name}). Training Loss: {np.mean(history.history['loss']):.4f}, Validation Loss: {np.mean(history.history['val_loss']):.4f}.",  # Simplified desc
            )

            # Scatter Plot for Actual vs Predicted (First 30 Days = 720 hours)
            plt.figure(figsize=(12, 6))

            # Prepare data for scatter plot (first 720 points from the cleaned flat arrays)
            num_points_for_scatter = min(720, len(y_test_inv_clean_flat))
            y_test_scatter = y_test_inv_clean_flat[:num_points_for_scatter]
            y_pred_scatter = y_pred_inv_clean_flat[:num_points_for_scatter]

            r2_scatter, slope_scatter, intercept_scatter = np.nan, np.nan, np.nan
            if (
                len(y_test_scatter) > 1 and len(y_pred_scatter) > 1
            ):  # Min 2 points for regression
                plt.scatter(
                    y_test_scatter, y_pred_scatter, label="Data points", alpha=0.7, s=10
                )  # Smaller points

                reg = LinearRegression().fit(
                    y_test_scatter.reshape(-1, 1), y_pred_scatter
                )
                r2_scatter = reg.score(y_test_scatter.reshape(-1, 1), y_pred_scatter)
                slope_scatter = reg.coef_[0]
                intercept_scatter = reg.intercept_

                line_x = np.array([min(y_test_scatter), max(y_test_scatter)])
                line_y = slope_scatter * line_x + intercept_scatter
                plt.plot(
                    line_x, line_y, color="red", label=f"LR Line (R²={r2_scatter:.2f})"
                )

                min_val = min(
                    min(y_test_scatter, default=0), min(y_pred_scatter, default=0)
                )
                max_val = max(
                    max(y_test_scatter, default=1), max(y_pred_scatter, default=1)
                )
                plt.plot(
                    [min_val, max_val], [min_val, max_val], "g--", label="Ideal (y=x)"
                )
            else:
                plt.text(
                    0.5,
                    0.5,
                    "Not enough data for scatter plot/regression",
                    ha="center",
                    va="center",
                    transform=plt.gca().transAxes,
                )

            plt.xlabel("Actual SO2")
            plt.ylabel("Predicted SO2")
            plt.title(
                f"Fold {fold + 1} Actual vs Predicted Scatter Plot (First {num_points_for_scatter // 24} Days) - {english_file_name}"
            )
            plt.legend()
            scatter_plot_filename = (
                self.output_dir
                / f"fold_{fold + 1}_scatter_pred_{english_file_name}.png"
            )
            plt.savefig(scatter_plot_filename, dpi=300, bbox_inches="tight")
            plt.close()

            desc_scatter = f"กราฟนี้แสดงการเปรียบเทียบระหว่างค่าจริงและค่าพยากรณ์ในรูปแบบ Scatter Plot สำหรับ Fold {fold + 1} ({num_points_for_scatter // 24} วันแรก) ({file_name}). "
            if not np.isnan(r2_scatter):
                desc_scatter += f"Linear Regression: R² = {r2_scatter:.3f}, Slope = {slope_scatter:.3f}, Intercept = {intercept_scatter:.3f}. "
                desc_scatter += f"Slope {'เข้าใกล้ 1' if abs(slope_scatter - 1) < 0.15 else 'ไม่เข้าใกล้ 1'}. "  # Adjusted threshold
                mean_y_test = y_test_scatter.mean() if len(y_test_scatter) > 0 else 0
                origin_threshold = (
                    0.15 * mean_y_test if mean_y_test != 0 else 0.15
                )  # Adjusted threshold
                desc_scatter += f"Origin (Intercept) {'เข้าใกล้ 0' if abs(intercept_scatter) < origin_threshold else 'ไม่เข้าใกล้ 0'}. "
            else:
                desc_scatter += "ไม่สามารถคำนวณ Linear Regression ได้ (ข้อมูลไม่เพียงพอ). "
            # ... (rest of original description logic if any)
            add_figure(
                doc,
                scatter_plot_filename,
                f"Fold {fold + 1} Actual vs Predicted Scatter Plot - {file_name}",
                desc_scatter,
            )

            doc.add_heading(f"Fold {fold + 1} Metrics - {file_name}", 2)
            if not np.isnan(mse):
                doc.add_paragraph(f"MSE: {mse:.4f}\nMAE: {mae:.4f}\nR²: {r2:.4f}")
            else:
                doc.add_paragraph(
                    f"Metrics could not be calculated for Fold {fold + 1}."
                )

        # Add TSCV split size plot into the report (now uses all fold sizes)
        doc.add_heading(f"แบ่งชุดข้อมูล (Time Series Cross Validation) - {file_name}", 1)
        # ... (plot_tscv and add_figure for TSCV as before)
        if dataset_size:
            self.plot_tscv(
                dataset_size
            )  # This will now plot for all folds if N_SPLITS > 1
            add_figure(
                doc,
                "temp_fig.png",
                f"Time Series Cross Validation Split Sizes - {file_name}",
                "กราฟนี้แสดงขนาดของชุดข้อมูลในแต่ละ Fold โดย Train และ Validation",
            )
        else:
            doc.add_paragraph("ไม่สามารถแสดงกราฟการแบ่งชุดข้อมูลได้")

        # Average Metrics Table
        # Filter out NaN scores before calculating mean
        valid_mse_scores = [s for s in scores["mse"] if not np.isnan(s)]
        valid_mae_scores = [s for s in scores["mae"] if not np.isnan(s)]
        valid_r2_scores = [s for s in scores["r2"] if not np.isnan(s)]

        if valid_mse_scores:  # Check if there's at least one valid score
            avg_metrics_data = []
            if valid_mse_scores:
                avg_metrics_data.append(
                    {"Metric": "MSE", "Value": np.mean(valid_mse_scores)}
                )
            if valid_mae_scores:
                avg_metrics_data.append(
                    {"Metric": "MAE", "Value": np.mean(valid_mae_scores)}
                )
            if valid_r2_scores:
                avg_metrics_data.append(
                    {"Metric": "R²", "Value": np.mean(valid_r2_scores)}
                )

            if avg_metrics_data:
                avg_metrics = pd.DataFrame(avg_metrics_data)
                desc = f"ตารางนี้แสดงค่าเฉลี่ยของเมตริกจาก {len(valid_mse_scores)} fold(s) ที่มีข้อมูลสมบูรณ์ ({file_name}). "
                # ... (rest of avg_metrics description logic as before) ...
                add_table(
                    doc,
                    avg_metrics,
                    f"Average Metrics Across Folds - {file_name}",
                    desc,
                )
        else:
            doc.add_paragraph(
                f"Cannot calculate average metrics due to no complete data in any folds ({file_name})"
            )

        # SHAP Analysis (using last_fold_model if only last fold was fully trained)
        # Or, if you want SHAP on a model trained on all X_full, retrain xgb_model here.
        # Current XGBoost SHAP uses X_full_flat, y_full_flat which is fine.
        print("Step 9: Computing SHAP Tree Importance with XGBoost...")
        # ... (SHAP analysis code as in original, ensure `feature_importance` DataFrame is created)
        doc.add_heading(
            f"Feature Importance with SHAP Tree via XGBoost (Post-Training) - {file_name}",
            1,
        )
        feature_importance = (
            pd.DataFrame()
        )  # Initialize to handle cases where SHAP might fail
        if features is not None and X_full_flat.shape[0] > 0:
            xgb_model = xgboost.XGBRegressor(
                objective="reg:squarederror",
                tree_method="hist",
                # device="cuda", # Enable if GPU is available and configured for XGBoost
                random_state=self.CONST.SEED_VALUE,
            )
            # Ensure no NaNs for XGBoost
            if np.any(np.isnan(X_full_flat)) or np.any(np.isnan(y_full_flat)):
                doc.add_paragraph(
                    f"Warning: Data for XGBoost SHAP contains NaNs. Imputing with 0. ({file_name})"
                )
                X_full_flat = np.nan_to_num(X_full_flat, nan=0.0)
                y_full_flat = np.nan_to_num(y_full_flat, nan=0.0)

            if X_full_flat.shape[0] > 1:  # Need at least 2 samples for XGBoost
                xgb_model.fit(X_full_flat, y_full_flat)

                # Sample for SHAP, ensure it's not too large if X_full_flat is huge.
                # Taking last 1000 or all if less than 1000.
                sample_size_shap = min(1000, X_full_flat.shape[0])
                X_sample_flat = (
                    X_full_flat[-sample_size_shap:]
                    if sample_size_shap > 0
                    else X_full_flat
                )  # handle if X_full_flat is tiny

                if X_sample_flat.shape[0] > 0:
                    explainer = shap.TreeExplainer(xgb_model)
                    shap_values = explainer.shap_values(
                        X_sample_flat
                    )  # Can fail if X_sample_flat is problematic

                    shap_explanation = shap.Explanation(
                        values=shap_values,
                        base_values=explainer.expected_value,
                        data=X_sample_flat,
                        feature_names=features,
                    )
                    shap_importance_vals = np.abs(shap_values).mean(axis=0)
                    feature_importance = pd.DataFrame(
                        {
                            "Feature": features,
                            "Importance": shap_importance_vals,
                            # Add other SHAP stats if needed
                        }
                    ).sort_values("Importance", ascending=False)

                    plt.figure(
                        figsize=(10, max(6, len(features) * 0.3))
                    )  # Adjust height for many features
                    shap.plots.bar(
                        shap_explanation, max_display=len(features), show=False
                    )
                    plt.title(f"SHAP Feature Importance (Bar) - {english_file_name}")
                    shap_bar_path = (
                        self.output_dir / f"shap_bar_{english_file_name}.png"
                    )
                    plt.savefig(shap_bar_path, dpi=300, bbox_inches="tight")
                    plt.close()
                    add_figure(
                        doc,
                        shap_bar_path,
                        f"SHAP Tree Summary Plot (Bar) - {file_name}",
                        "SHAP Bar plot showing feature importance.",
                    )

                    plt.figure(
                        figsize=(12, max(8, len(features) * 0.4))
                    )  # Adjust height
                    shap.plots.beeswarm(
                        shap_explanation, max_display=len(features), show=False
                    )
                    plt.title(
                        f"SHAP Feature Importance (Beeswarm) - {english_file_name}"
                    )
                    shap_beeswarm_path = (
                        self.output_dir / f"shap_beeswarm_{english_file_name}.png"
                    )
                    plt.savefig(shap_beeswarm_path, dpi=300, bbox_inches="tight")
                    plt.close()
                    add_figure(
                        doc,
                        shap_beeswarm_path,
                        f"SHAP Tree Summary Plot (Beeswarm) - {file_name}",
                        "SHAP Beeswarm plot.",
                    )

                    add_table(
                        doc,
                        feature_importance,
                        f"Feature Importance Ranking (SHAP) - {file_name}",
                        "Ranked feature importance from SHAP.",
                    )
                else:
                    doc.add_paragraph(
                        f"Not enough data for SHAP sample for {file_name}."
                    )
            else:
                doc.add_paragraph(
                    f"Not enough data to fit XGBoost model for SHAP for {file_name}."
                )
        else:
            doc.add_paragraph(
                f"Features not defined or X_full_flat empty. Skipping SHAP. ({file_name})"
            )

        print("Step 10: Generating Final Prediction with LSTM...")
        doc.add_heading(f"Final Prediction with LSTM - {file_name}", 1)
        # Use the model from the last fold if it was trained
        model_for_final_pred = (
            last_fold_model if last_fold_model is not None else model
        )  # Fallback if last_fold_model wasn't set

        if (
            model_for_final_pred
            and features is not None
            and len(X_full) > 0
            and scaler_target is not None
        ):
            last_data = X_full[-1:]  # Last sequence from the full dataset
            if np.any(np.isnan(last_data)):
                doc.add_paragraph(
                    f"Warning: Data for final prediction contains NaN values, replaced with 0. ({file_name})"
                )
                last_data = np.nan_to_num(last_data, nan=0.0)

            pred_next_24 = model_for_final_pred.predict(last_data)
            pred_next_24_inv = scaler_target.inverse_transform(pred_next_24)

            # Actual data for the last 24 hours (for context in the plot)
            # df here should be the scaled dataframe. We need original values for actual_last_24_inv
            # This requires careful handling of which 'df' is used.
            # Let's re-fetch last 24 actuals from the *unscaled* df_raw or a point before scaling.
            # For simplicity, using the scaled `df` and inverse transforming is okay if it's just for plot context.
            actual_last_24_scaled = (
                df["SO2"].iloc[-pred_length:].values.reshape(-1, 1)
            )  # Takes last PRED_LENGTH points
            actual_last_24_inv = scaler_target.inverse_transform(actual_last_24_scaled)

            last_actual_timestamps = df.index[-pred_length:]
            future_timestamps = pd.date_range(
                start=df.index[-1], periods=pred_length + 1, freq="H"
            )[1:]

            plt.figure(figsize=(12, 6))
            plt.plot(
                last_actual_timestamps,
                actual_last_24_inv.flatten(),
                label="Actual (Last 24H)",
                marker="o",
            )
            plt.plot(
                future_timestamps,
                pred_next_24_inv[0].flatten(),  # pred_next_24_inv is (1, pred_length)
                label="Predicted (Next 24H)",
                linestyle="--",
                marker="x",
            )
            plt.title(f"Next 24-Hour SO2 Prediction (LSTM) - {english_file_name}")
            plt.xlabel("Datetime")
            plt.ylabel("SO2")
            plt.legend()
            plt.xticks(rotation=45)
            plt.tight_layout()
            final_pred_path = self.output_dir / f"final_pred_{english_file_name}.png"
            plt.savefig(final_pred_path, dpi=300, bbox_inches="tight")
            plt.close()

            desc_final_pred = (
                f"กราฟนี้แสดงการพยากรณ์ SO2 สำหรับ {pred_length} ชั่วโมงข้างหน้า "
                f"เทียบกับข้อมูลจริง {pred_length} ชั่วโมงล่าสุด ({file_name}). "
                f"ค่าเฉลี่ยจริงล่าสุด: {actual_last_24_inv.mean():.2f}, "
                f"ค่าเฉลี่ยพยากรณ์: {pred_next_24_inv[0].mean():.2f}."
            )
            add_figure(
                doc,
                final_pred_path,
                f"Next 24-Hour Prediction (LSTM) - {file_name}",
                desc_final_pred,
            )
        else:
            doc.add_paragraph(
                f"Skipping final prediction due to missing model, features, data, or scaler. ({file_name})"
            )

        # Step 10.5: Retrain with Top 5 SHAP Features for specific stations
        stations_for_retrain = ["Ban_Sopad_Hospital", "Ban sadet station"]
        # ^ "Ban sadet station" should match the english_file_name derived from its CSV file.
        # Example: if file is "Ban sadet station.csv", english_file_name will be "Ban sadet station".

        if english_file_name in stations_for_retrain:
            print(
                f"Step 10.5: Retraining with Top 5 SHAP Features for {english_file_name}..."
            )
            doc.add_heading(f"Retraining with Top 5 SHAP Features - {file_name}", 1)

            if not feature_importance.empty and "Feature" in feature_importance.columns:
                top_5_features = feature_importance["Feature"].head(5).tolist()
                doc.add_paragraph(
                    f"Top 5 features selected based on SHAP values: {', '.join(top_5_features)}"
                )

                # df_for_retrain should be the scaled dataframe used for the main model.
                # It contains all original features (scaled) and scaled SO2.
                df_for_retrain = df.copy()  # df is already scaled here.

                if not all(f in df_for_retrain.columns for f in top_5_features):
                    doc.add_paragraph(
                        f"One or more top 5 SHAP features not found in the scaled DataFrame. Skipping retraining for {file_name}."
                    )
                else:
                    X_retrain_full, y_retrain_full = prepare_sequences(
                        df_for_retrain, top_5_features, seq_length, pred_length
                    )

                    if X_retrain_full.shape[0] > 0 and y_retrain_full.shape[0] > 0:
                        # Use the same last fold split logic for consistency with the main model's evaluation
                        # Or could do a simpler train/test split if preferred for retraining.
                        # For now, using last fold of TSCV.
                        retrain_tscv = TimeSeriesSplit(
                            n_splits=self.CONST.N_SPLITS
                        )  # Re-init or ensure tscv is general
                        retrain_train_idx, retrain_test_idx = list(
                            retrain_tscv.split(X_retrain_full)
                        )[-1]

                        X_train_rt, X_test_rt = (
                            X_retrain_full[retrain_train_idx],
                            X_retrain_full[retrain_test_idx],
                        )
                        y_train_rt, y_test_rt = (
                            y_retrain_full[retrain_train_idx],
                            y_retrain_full[retrain_test_idx],
                        )

                        if X_train_rt.shape[0] > 0 and X_test_rt.shape[0] > 0:
                            train_dataset_rt = tf.data.Dataset.from_tensor_slices(
                                (X_train_rt, y_train_rt)
                            ).batch(batch_size)
                            test_dataset_rt = tf.data.Dataset.from_tensor_slices(
                                (X_test_rt, y_test_rt)
                            ).batch(batch_size)

                            tf.keras.backend.clear_session()
                            model_retrain = Sequential(
                                [
                                    LSTM(
                                        128,
                                        input_shape=(seq_length, len(top_5_features)),
                                        return_sequences=True,
                                    ),
                                    Dropout(0.2),
                                    LSTM(64),
                                    Dropout(0.2),
                                    Dense(32, activation="relu"),
                                    Dense(pred_length),
                                ]
                            )
                            model_retrain.compile(optimizer="adam", loss="mse")
                            print(
                                f"Fitting retrained model for {english_file_name} with features: {top_5_features}"
                            )
                            model_retrain.fit(
                                train_dataset_rt,
                                epochs=self.CONST.EPOCHS,
                                validation_data=test_dataset_rt,
                                verbose=0,
                            )

                            y_pred_rt = model_retrain.predict(test_dataset_rt)
                            y_test_rt_inv = scaler_target.inverse_transform(y_test_rt)
                            y_pred_rt_inv = scaler_target.inverse_transform(y_pred_rt)

                            # Flatten and clean NaNs for scatter plot
                            y_test_rt_flat = y_test_rt_inv.flatten()
                            y_pred_rt_flat = y_pred_rt_inv.flatten()
                            valid_rt_indices = ~np.isnan(y_test_rt_flat) & ~np.isnan(
                                y_pred_rt_flat
                            )
                            y_test_rt_clean_flat = y_test_rt_flat[valid_rt_indices]
                            y_pred_rt_clean_flat = y_pred_rt_flat[valid_rt_indices]

                            r2_rt_scatter, slope_rt_scatter, intercept_rt_scatter = (
                                np.nan,
                                np.nan,
                                np.nan,
                            )
                            if len(y_test_rt_clean_flat) > 1:
                                plt.figure(figsize=(10, 6))
                                plt.scatter(
                                    y_test_rt_clean_flat,
                                    y_pred_rt_clean_flat,
                                    alpha=0.7,
                                    label="Data points",
                                    s=10,
                                )
                                plt.xlabel("Actual SO2 (Retrained with Top 5 Features)")
                                plt.ylabel(
                                    "Predicted SO2 (Retrained with Top 5 Features)"
                                )

                                reg_rt = LinearRegression().fit(
                                    y_test_rt_clean_flat.reshape(-1, 1),
                                    y_pred_rt_clean_flat,
                                )
                                r2_rt_scatter = reg_rt.score(
                                    y_test_rt_clean_flat.reshape(-1, 1),
                                    y_pred_rt_clean_flat,
                                )
                                slope_rt_scatter = reg_rt.coef_[0]
                                intercept_rt_scatter = reg_rt.intercept_

                                line_x_rt = np.array(
                                    [
                                        min(y_test_rt_clean_flat),
                                        max(y_test_rt_clean_flat),
                                    ]
                                )
                                line_y_rt = (
                                    slope_rt_scatter * line_x_rt + intercept_rt_scatter
                                )
                                plt.plot(
                                    line_x_rt,
                                    line_y_rt,
                                    color="red",
                                    label=f"LR Line (R²={r2_rt_scatter:.2f})",
                                )

                                min_val_rt = min(
                                    min(y_test_rt_clean_flat, default=0),
                                    min(y_pred_rt_clean_flat, default=0),
                                )
                                max_val_rt = max(
                                    max(y_test_rt_clean_flat, default=1),
                                    max(y_pred_rt_clean_flat, default=1),
                                )
                                plt.plot(
                                    [min_val_rt, max_val_rt],
                                    [min_val_rt, max_val_rt],
                                    "g--",
                                    label="Ideal (y=x)",
                                )
                                plt.legend()
                            else:
                                plt.text(
                                    0.5,
                                    0.5,
                                    "Not enough data for retrained scatter plot",
                                    ha="center",
                                    va="center",
                                    transform=plt.gca().transAxes,
                                )

                            plt.title(
                                f"Actual vs Predict (Top 5 SHAP Features) - {english_file_name}"
                            )
                            retrain_scatter_path = (
                                self.output_dir
                                / f"retrain_top5_scatter_{english_file_name}.png"
                            )
                            plt.savefig(
                                retrain_scatter_path, dpi=300, bbox_inches="tight"
                            )
                            plt.close()

                            desc_fig_retrain = (
                                f"Scatter plot of Actual vs Predicted SO2 after retraining with top 5 SHAP features: {', '.join(top_5_features)}. "
                                f"For station: {file_name}. "
                            )
                            if not np.isnan(r2_rt_scatter):
                                desc_fig_retrain += (
                                    f"Linear Regression: R² = {r2_rt_scatter:.3f}, Slope = {slope_rt_scatter:.3f}, Intercept = {intercept_rt_scatter:.3f}. "
                                    f"Slope {'เข้าใกล้ 1' if abs(slope_rt_scatter - 1) < 0.15 else 'ไม่เข้าใกล้ 1'}. "
                                    f"Origin (Intercept) {'เข้าใกล้ 0' if abs(intercept_rt_scatter) < (0.15 * (y_test_rt_clean_flat.mean() if len(y_test_rt_clean_flat) > 0 else 0.15)) else 'ไม่เข้าใกล้ 0'}."
                                )
                            else:
                                desc_fig_retrain += (
                                    "ไม่สามารถคำนวณ Linear Regression ได้ (ข้อมูลไม่เพียงพอ)."
                                )
                            add_figure(
                                doc,
                                retrain_scatter_path,
                                f"Actual vs Predict Scatter (Retrained Top 5) - {file_name}",
                                desc_fig_retrain,
                            )
                            del model_retrain
                        else:
                            doc.add_paragraph(
                                f"Not enough data in train/test splits for retraining for {file_name}."
                            )
                    else:
                        doc.add_paragraph(
                            f"Not enough data to create sequences for retraining model for {file_name} with top 5 features."
                        )
            else:
                doc.add_paragraph(
                    f"Skipping retraining for {file_name}: SHAP feature_importance not available or empty."
                )

        print("Step 11: Saving Report...")
        set_thsarabupsk_font(doc)
        doc.save(
            self.output_dir / f"report_{english_file_name}.docx"
        )  # Use english_file_name for consistency
        print(f"Step 12: Process Completed for {file_name}!")

        # Clean memory (free memory)
        del df, df_raw, X_full, y_full, X_full_flat, y_full_flat
        if "model" in locals():
            del model
        if "last_fold_model" in locals() and last_fold_model is not None:
            del last_fold_model
        if "xgb_model" in locals():
            del xgb_model
        if "doc" in locals():
            del doc  # Delete doc after saving
        gc.collect()

    def run(self):
        for csv_file in self.csv_files:
            try:
                self.process_file(csv_file)
            except Exception as e:
                print(f"Error processing {csv_file}: {e}")
                # Optionally, create a minimal error report for this file
                error_doc = Document()
                error_doc.add_heading(
                    f"Error Report for {os.path.basename(csv_file)}", 0
                )
                error_doc.add_paragraph(
                    f"An error occurred while processing the file: {csv_file}"
                )
                error_doc.add_paragraph(str(e))
                import traceback

                error_doc.add_paragraph("Traceback:")
                error_doc.add_paragraph(traceback.format_exc())
                set_thsarabupsk_font(error_doc)
                error_file_name = os.path.basename(csv_file).replace(".csv", "")
                english_error_file_name = self.file_name_mapping.get(
                    error_file_name, error_file_name
                )
                error_doc.save(
                    self.output_dir
                    / f"report_{english_error_file_name}_PROCESSING_ERROR.docx"
                )
            gc.collect()


if __name__ == "__main__":
    input_directory = "/mnt/e/MayThesis2025/src/labs4/s02_v01_02_clean/output-680507-2319"  # Example, replace with your actual path
    # Check if directory exists
    if not os.path.isdir(input_directory):
        print(f"Error: Input directory not found: {input_directory}")
        # Create a dummy directory with a sample CSV if you want to test without the exact path
        # For example:
        # Path("dummy_input").mkdir(exist_ok=True)
        # sample_data = pd.DataFrame({'Datetime': pd.to_datetime(['2023-01-01 00:00:00', '2023-01-01 01:00:00']), 'SO2': [1,2], 'Feature1': [3,4]})
        # sample_data.to_csv("dummy_input/(38t)โรงพยาบาลส่งเสริมสุขภาพตำลบ้านสบป้าด.csv", index=False)
        # input_directory = "dummy_input"
        # print(f"Using dummy input directory: {input_directory}")

    processor = DynamicInputDataset(input_directory)
    processor.run()
